# -*- coding:UTF-8 -*-
import sys

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Cm
# from pptx.chart.data import CategoryChartData
# from pptx.enum.chart import XL_CHART_TYPE, XL_LEGEND_POSITION, XL_DATA_LABEL_POSITION
# from pyecharts import options as opts
from pyecharts.charts import Line, Bar, Bar3D, Page, Pie
from pyecharts.globals import ChartType
# from pyecharts import Bar3D,Line,Bar
from pyecharts import options as opts
# -*- coding: UTF-8 -*-
from pyecharts.render import make_snapshot
from pyecharts.globals import CurrentConfig
CurrentConfig.ONLINE_HOST = ""
from snapshot_phantomjs import snapshot
# sys.path.append(".")
extract_template = {
    'dw': [],
    'xmly': [],
    '项目': [],
    '起止时间': [],
    '分组方式': [],
    '项目状态': [],
    '统计类型': [],
    '': [],  # 第三种，非科研....问一下归类
}

'''统计类模板'''
{
    ''
}


'''风险类模板'''
risk_template = {
   '项目状态': {
       '在研延期': '20%'
   },
    '非科研类':'2'
}


def generate_bar_chart(document):


    attr = ["JT公司", "战略支援部", "战区", "公安部", "国家科工局", "国家自然基金委"]
    v1 = [5, 20, 36, 10, 75, 90]
    v2 = [10, 25, 8, 60, 20, 80]
    bar = Bar()
    bar.add_yaxis("2022年", [5, 20, 36, 10, 75, 90], category_gap="50%")
    bar.add_yaxis("2021年", [10, 25, 8, 60, 20, 80], category_gap="50%")
    bar.add_xaxis(attr)

    # bar.add("2021年", attr, v1, is_stack=True)
    # bar.add("2022年", attr, v2, is_stack=True)


    # snapshot.PHANTOMJS_EXEC = r"D:\Data\tool\phantomjs-2.1.1-windows\bin\phantomjs.exe"
    make_snapshot(snapshot, bar.render(), "bar_chart.png")
    paragraph = document.add_paragraph()
    run = paragraph.add_run('201所2021年2022年部门经费对比')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('bar_chart.png', width=Cm(10), height=Cm(5))
    document.paragraphs[2].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def generate_line_chart(document):
    # 生成一个简单的折线图
    def create_line_chart():
        line = (
            Line()
            .add_xaxis(
                ["一月", "二月", "三月", "四月", "五月", "六月", "七月", "八月", "九月", "十月", "十一月", "十二月"])
            .add_yaxis("2022年经费", [5, 20, 36, 10, 75, 90, 35, 60, 80, 40, 50, 60])
            .add_yaxis("2021年经费", [10, 25, 8, 60, 20, 80, 15, 40, 50, 30, 40, 50])
            # .set_global_opts(title_opts=opts.TitleOpts(title="按时间统计201所经费折线图"))
        )
        return line

    # 创建图表并导出为 JPG
    line_chart = create_line_chart()
    # snapshot.PHANTOMJS_EXEC = r"D:\Data\tool\phantomjs-2.1.1-windows\bin\phantomjs.exe"
    make_snapshot(snapshot, line_chart.render(), "line_chart.png")

    paragraph = document.add_paragraph()
    # paragraph.add_run('201所2021年总计共支出428万元，2022总计共支出561万元，')')
    run = paragraph.add_run('201所2021、2022年度经费报告')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('line_chart.png', width=Cm(10), height=Cm(5))
    document.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def generate_3d_chart(document):
    bar3d = Bar3D()
    x_axis = [
        "12a", "1a", "2a", "3a", "4a", "5a", "6a", "7a", "8a", "9a", "10a", "11a",
        "12p", "1p", "2p", "3p", "4p", "5p", "6p", "7p", "8p", "9p", "10p", "11p"
    ]
    y_axis = [
        "Saturday", "Friday", "Thursday", "Wednesday", "Tuesday", "Monday", "Sunday"
    ]
    data = [
        [0, 0, 5], [0, 1, 1], [0, 2, 0], [0, 3, 0], [0, 4, 0], [0, 5, 0],
        [0, 6, 0], [0, 7, 0], [0, 8, 0], [0, 9, 0], [0, 10, 0], [0, 11, 2],
        [0, 12, 4], [0, 13, 1], [0, 14, 1], [0, 15, 3], [0, 16, 4], [0, 17, 6],
        [0, 18, 4], [0, 19, 4], [0, 20, 3], [0, 21, 3], [0, 22, 2], [0, 23, 5],
        [1, 0, 7], [1, 1, 0], [1, 2, 0], [1, 3, 0], [1, 4, 0], [1, 5, 0],
        [1, 6, 0], [1, 7, 0], [1, 8, 0], [1, 9, 0], [1, 10, 5], [1, 11, 2],
        [1, 12, 2], [1, 13, 6], [1, 14, 9], [1, 15, 11], [1, 16, 6], [1, 17, 7],
        [1, 18, 8], [1, 19, 12], [1, 20, 5], [1, 21, 5], [1, 22, 7], [1, 23, 2],
        [2, 0, 1], [2, 1, 1], [2, 2, 0], [2, 3, 0], [2, 4, 0], [2, 5, 0],
        [2, 6, 0], [2, 7, 0], [2, 8, 0], [2, 9, 0], [2, 10, 3], [2, 11, 2],
        [2, 12, 1], [2, 13, 9], [2, 14, 8], [2, 15, 10], [2, 16, 6], [2, 17, 5],
        [2, 18, 5], [2, 19, 5], [2, 20, 7], [2, 21, 4], [2, 22, 2], [2, 23, 4],
        [3, 0, 7], [3, 1, 3], [3, 2, 0], [3, 3, 0], [3, 4, 0], [3, 5, 0],
        [3, 6, 0], [3, 7, 0], [3, 8, 1], [3, 9, 0], [3, 10, 5], [3, 11, 4],
        [3, 12, 7], [3, 13, 14], [3, 14, 13], [3, 15, 12], [3, 16, 9], [3, 17, 5],
        [3, 18, 5], [3, 19, 10], [3, 20, 6], [3, 21, 4], [3, 22, 4], [3, 23, 1],
        [4, 0, 1], [4, 1, 3], [4, 2, 0], [4, 3, 0], [4, 4, 0], [4, 5, 1],
        [4, 6, 0], [4, 7, 0], [4, 8, 0], [4, 9, 2], [4, 10, 4], [4, 11, 4],
        [4, 12, 2], [4, 13, 4], [4, 14, 4], [4, 15, 14], [4, 16, 12], [4, 17, 1],
        [4, 18, 8], [4, 19, 5], [4, 20, 3], [4, 21, 7], [4, 22, 3], [4, 23, 0],
        [5, 0, 2], [5, 1, 1], [5, 2, 0], [5, 3, 3], [5, 4, 0], [5, 5, 0],
        [5, 6, 0], [5, 7, 0], [5, 8, 2], [5, 9, 0], [5, 10, 4], [5, 11, 1],
        [5, 12, 5], [5, 13, 10], [5, 14, 5], [5, 15, 7], [5, 16, 11], [5, 17, 6],
        [5, 18, 0], [5, 19, 5], [5, 20, 3], [5, 21, 4], [5, 22, 2], [5, 23, 0],
        [6, 0, 1], [6, 1, 0], [6, 2, 0], [6, 3, 0], [6, 4, 0], [6, 5, 0],
        [6, 6, 0], [6, 7, 0], [6, 8, 0], [6, 9, 0], [6, 10, 1], [6, 11, 0],
        [6, 12, 2], [6, 13, 1], [6, 14, 3], [6, 15, 4], [6, 16, 0], [6, 17, 0],
        [6, 18, 0], [6, 19, 0], [6, 20, 1], [6, 21, 2], [6, 22, 2], [6, 23, 6]
    ]
    range_color = ['#313695', '#4575b4', '#74add1', '#abd9e9', '#e0f3f8', '#ffffbf',
                   '#fee090', '#fdae61', '#f46d43', '#d73027', '#a50026']
    bar3d.add(
        "",
        [[d[1], d[0], d[2]] for d in data],
    )

    make_snapshot(snapshot, bar3d.render(), "bar3d_chart.png")

    paragraph = document.add_paragraph()
    run = paragraph.add_run('201所3D柱状图示例')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('bar3d_chart.png', width=Cm(10), height=Cm(5))
    document.paragraphs[3].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def to_local(document, geolines,path_name,chart_name):
    make_snapshot(snapshot, geolines.render(), path_name)

    paragraph = document.add_paragraph()
    run = paragraph.add_run(chart_name)
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture(path_name, width=Cm(10), height=Cm(5))
    document.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def generate_geo_chart(document):
    from pyecharts.charts import Geo,Line

    city_pairs = [("北京", "上海"), ("北京", "广州"), ("北京", "拉萨"), ("北京", "乌鲁木齐"), ("北京", "西安")]

    # 创建一个 Geo 图
    geo = (
        Geo()
        .add_schema(maptype="china")
    )

    # 添加城市数据到 Geo 图
    for start, end in city_pairs:
        geo.add("城市", [(start, 1), (end, 1)], type_='effectScatter')

        # 添加连线
    geo.add(
        "线路",
        city_pairs,
        type_=ChartType.LINES,
        effect_opts=opts.EffectOpts(
            # 设置线条的颜色和变化
            scale=6,
            period=15,
            color='blue',
            symbol_size=6,
        ),
        linestyle_opts=opts.LineStyleOpts(
            width=2,  # 设置线条宽度
            type_="solid",  # 设置线条样式
            color="lightblue",  # 设置线条颜色
            curve=0.4  # 设置线条的弯曲度 (0: 直线, 1: 最大弯曲)
        ),
        label_opts=opts.LabelOpts(
            is_show=True,  # 显示标签
            position="end",  # 标签显示在线条的末尾
            font_size=8,  # 标签字体大小
            formatter="{b}",  # 标签内容为地名
        ),
    )

    # 设置全局图表选项
    geo.set_global_opts(title_opts=opts.TitleOpts(title="GeoLine 示例"))
    to_local(document,geo,"geo_chart.png",'201所业务分布图示例')


def generate_pie_chart(document):
    pie = (
        Pie()
        .add(
            series_name="202所项目数量图",
            data_pair=[
                ("非科研类", 40),
                ("军工所", 30),
                ("军科院", 20),
                ("本部", 10),
            ],
        )
        .set_series_opts(
            label_opts=opts.LabelOpts(
                is_show=True,  # 显示标签
                formatter="{b}: {c} ({d}%)"  # 显示格式为: 名称: 数值 (百分比%)
            )
        )
        .set_global_opts(
            title_opts=opts.TitleOpts(title="201所项目数量饼图")  # 设置标题
        )
    )
    to_local(document,pie,"pie_chart.png","201所项目数量图")


def generate_docx():

    document = Document()
    #document.add_picture('chart.png', width=Cm(15.2), height=Cm(22.9))


 
    snapshot.PHANTOMJS_EXEC = 'phantomjs-2.1.1-linux-x86_64/bin/phantomjs'
    # generate_line_chart(document)
    # generate_bar_chart(document)
    # # generate_3d_chart(document)
    # generate_geo_chart(document)
    # generate_pie_chart(document)
    paragraph = document.add_paragraph()
    run1 = paragraph.add_run('总体概况',)
    run1.bold = True
    paragraph.add_run('22123123123131')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph2 = document.add_paragraph()

    run2 = paragraph2.add_run('重点项目执行情况', )
    run2.bold = True

    paragraph3 = document.add_paragraph()
    run3 = paragraph3.add_run('项目来源情况', )
    run3.bold = True
    generate_bar_chart(document)


    document.save('test3.docx')


if __name__ == '__main__':
    generate_docx()
