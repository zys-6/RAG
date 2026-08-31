import copy
import datetime
import hashlib
import json
import logging
import os
import re
from collections import defaultdict, Counter
from typing import Union, List, Dict

import openai
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from pyecharts.charts import Line, Bar
from pyecharts.render import make_snapshot
from pymilvus import MilvusClient
from snapshot_phantomjs import snapshot
#snapshot.PHANTOMJS_EXEC = "rag/static/phantomjs-2.1.1-linux-x86_64/bin/phantomjs"
from rag.configs import app_config,prompt_config
from rag.utils.mapping import get_organization_mapping, get_normal_mapping, get_project_date_mapping, PROJECT_SOURCE, \
    SUGGEST_DEPARTMENT, TYPE, TYPE_ALIAS, PROJECT_SOURCE_ALIAS, LABEL, SIGNIFICANCE, BIDDING, STATUS, MAPPING_URL, \
    get_undertaking_unit
from rag.utils.request_llm import get_question_expend, generate_question_by_history, get_llm_response, \
    get_processing_count
from pyecharts.globals import CurrentConfig
CurrentConfig.ONLINE_HOST = 'rag/'
now_year = datetime.datetime.now()
logger = logging.getLogger(__name__)

# MILVUS_COLLECTION = os.environ['COLLECTION_NAME']
# MILVUS_URI = os.environ['MILVUS_URI']
# LLM_URL = os.environ['API_BASE']
# MODEL_NAME = os.environ['MODEL_NAME']
# os.environ['OPENAI_API_KEY'] = os.environ['API_KEY']
# os.environ['OPENAI_API_BASE'] = os.environ['API_BASE']
# openai_client = openai.AsyncClient(api_key=os.environ['API_KEY'],
#                                    base_url=os.environ['API_BASE'])
# RERANK_URL = os.environ['RERANK_URL']

MILVUS_COLLECTION = app_config['MILVUS_COLLECTION']
MILVUS_URI = app_config['MILVUS_URI']
LLM_URL = app_config['LLM_URL']
RERANK_URL = app_config['RERANK_URL']
EMBEDDING_URL = app_config['EMBEDDING_URL']
MODEL_NAME = app_config['MODEL_NAME']
API_KEY = app_config['API_KEY']
API_BASE_URL = app_config['API_BASE_URL']
TOKENIZE_URL = app_config['TOKENIZE_URL']
MAX_TOKENS = app_config['MAX_TOKENS']
GET_TOKEN_URL = app_config['GET_TOKEN_URL']
MAX_NEW_TOKENS = app_config['MAX_NEW_TOKENS']

INTENT_PROMPT = prompt_config['intent_prompt']
INTENT_EXAMPLE = prompt_config['intent_example']
SECOND_SYSTEM_TEMPLATE_ = prompt_config['second_system_template_']
SECOND_HUMAN_TEMPLATE = prompt_config['second_human_template']
ATTRIBUTE_PROMPT = prompt_config['attribute_prompt']
ATTRIBUTE_EXAMPLE = prompt_config['attribute_example']
ATTRIBUTE_HUMAN_TEMPLATE = prompt_config['attribute_human_prompt']
LLM_ANSWER_SYSTEM_TEMPLATE_ = prompt_config['llm_answer_system_template_']
LLM_ANSWER_HUMAN_TEMPLATE_ = prompt_config['llm_answer_human_template_']
QUESTION_CLASSIFICATION_PROMPT = prompt_config['question_classification_prompt']
QUESTION_CLASSIFICATION_EXAMPLE = prompt_config['question_classification_example']

milvus_client = MilvusClient(MILVUS_URI)
openai_client = openai.AsyncClient(api_key=API_KEY,
                                   base_url=API_BASE_URL)


def get_num_tokens(texts: List[str]) -> List[int]:
    num_tokens = []
    for text in texts:
        resp = requests.post(TOKENIZE_URL, json={
            "model": MODEL_NAME,
            "prompt": text,
            "add_special_tokens": True
        })
        num_tokens.append(len(resp.json()['tokens']))
    return num_tokens


def get_document_rerank(data):
    url = RERANK_URL
    resp = requests.post(url, json=data)
    resp = resp.json()
    return resp['scores']


def get_vectors(texts: List[str]) -> List[List[float]]:
    url = EMBEDDING_URL
    response = requests.post(url, json={"input": texts})
    vectors = []
    for item in response.json()['data']:
        vectors.append(item['embedding'])
    return vectors


def get_length(text: str):
    token_count = int(len('inf'))
    return 0
    try:
        resp = requests.post(GET_TOKEN_URL, json={
            "text": text,
            "model": MODEL_NAME
        },timeout=(2,30))
        token_count = json.loads(json.loads(resp.text)).get('tokencount', 0)
    except Exception as e:
        logger.error(e)
    return token_count


async def search(
        text: Union[str, List],
        filter: str = "",
        limit: int = 10) -> List[Dict]:
    output_fields = ["id", "page_content", "pages", "coordinates", "outline", "parent_id", "index", "type",
                     "document_id"]

    if isinstance(text, str):
        query_embed = get_vectors([text])[0]
    else:
        query_embed = text
    result = milvus_client.search(MILVUS_COLLECTION,
                                  data=[query_embed],
                                  limit=limit,
                                  filter=filter,
                                  output_fields=output_fields)[0]

    docs = []
    for item in result:
        page_content = item['entity']['page_content']
        del item['entity']['page_content']
        metadata = item['entity']
        metadata['id'] = item['id']
        docs.append({'page_content': page_content,
                     'metadata': metadata})
    return docs


def query(filter: str = "") -> List[Dict]:
    output_fields = ["id", "page_content", "pages", "coordinates", "outline", "document_id",
                     "parent_id", "index", "type"]
    result = milvus_client.query(MILVUS_COLLECTION, filter=filter, output_fields=output_fields)
    ret = []
    remove_context = set()
    for item in result:
        page_content = item['page_content']
        if page_content in remove_context:
            continue
        remove_context.add(page_content)
        del item['page_content']
        metadata = item
        ret.append({'page_content': page_content,
                    'metadata': metadata})
    ret = sorted(ret, key=lambda x: x['metadata']['index'])
    return ret


def remove_abstract_documents(documents: List[Dict]) -> List[Dict]:
    abstract_parent_id = None
    for _doc in documents:
        if _doc['page_content'].strip().startswith("摘要") or _doc['page_content'].strip().lower().startswith(
                "abstract"):
            abstract_parent_id = _doc['metadata']['parent_id']
    if abstract_parent_id:
        documents = [_doc for _doc in documents if _doc['metadata']['parent_id'] != abstract_parent_id]
    return documents


async def rerank_documents(query: str,
                           documents: List[Dict],
                           merge_special: bool = False):
    raw_documents = None
    if merge_special:
        """将documents中的文档中的特殊列表进行组装"""
        sorted_documents = sorted(documents, key=lambda x: x['metadata']['index'])
        """用parent_id进行分组"""
        parent_id2documents = defaultdict(list)
        for _document in sorted_documents:
            parent_id2documents[_document['metadata']['parent_id']].append(_document)

        _documents = []
        for parent_id, _sorted_documents in parent_id2documents.items():
            idx = 0
            while idx < len(_sorted_documents):
                to_merge = []
                if re.search("[：:]$", _sorted_documents[idx]['page_content'].strip()):
                    jdx = idx + 1
                    while jdx < len(_sorted_documents) and \
                            re.search("^[(（\[【]?[0-9一二三四五六七八九十]+[)）\]】]",
                                      _sorted_documents[jdx]['page_content'].strip()):
                        to_merge.append(_sorted_documents[jdx])
                        jdx += 1
                    if to_merge:
                        to_merge.insert(0, _sorted_documents[idx])
                        _metadata = {"ids": [_doc['metadata']['id'] for _doc in to_merge]}
                        _documents.append({
                            "page_content": "\n".join([_doc['page_content'] for _doc in to_merge]),
                            "metadata": _metadata
                        })
                    idx = jdx - 1
                else:
                    _documents.append(_sorted_documents[idx])
                idx += 1
        raw_documents = documents
        documents = _documents

    scores = get_document_rerank({"query": query, "texts": [document['page_content'] for document in documents]})

    """根据scores的正负值区分"""
    ret = []
    for idx in range(len(documents)):
        if scores[idx] > 0:
            ret.append((scores[idx], documents[idx]))
    ret = [x[1] for x in sorted(ret, key=lambda x: x[0], reverse=True)]
    if raw_documents:
        _ret = []
        for doc in ret:
            if 'ids' in doc['metadata']:
                _ret.extend([_doc for _doc in raw_documents if _doc['metadata']['id'] in doc['metadata']['ids']])
            else:
                _ret.append(doc)
        ret = _ret
    return ret


async def search_with_same_outline(
        text: Union[str, List],
        with_rerank: bool = True) -> List[Dict]:

    docs = await search(text, limit=10)
    like_docs = await search(text,filter="page_content LIKE '% %'", limit=3)
    docs.extend(like_docs)
    parent_ids = list(set([item['metadata']['parent_id']
                           for item in docs
                           if item['metadata']['parent_id'] != 'None']))
    if parent_ids:

        docs = query("({}) and type == 'text'".format(
            " or ".join(["parent_id LIKE '%{}%'".format(_parent_id) for _parent_id in parent_ids])
        ))
        print('query docs len:', str(len(docs)))

    return docs


async def get_user_intent(question: str) -> str:
    intent_template = INTENT_PROMPT + '\t示例:' +INTENT_EXAMPLE
    messages = [{
            "role": "system",
            "content": intent_template
        },
        {
            "role": "user",
            "content": question
    }]
    llm_response = await openai_client.chat.completions.create(messages=messages,
                                                               model=MODEL_NAME,
                                                               stream=False, temperature=0.2)
    try:
        resp_json = llm_response.choices[0].message.content.strip()
        intent = eval(resp_json).get('type','文档')
    except Exception as e:
        logger.error(e)
        intent = '文档'
    return intent


async def get_docs_subscription(query):

    llm_response = await openai_client.chat.completions.create(messages=query['messages'],model=query['model'],stream=True,temperature=0.2,max_tokens=2048)
    print(llm_response)
    return llm_response


async def get_question_classification_from_question(question: str):

    system_template = QUESTION_CLASSIFICATION_PROMPT.format(date=now_year.strftime('%Y年%m月%d日')) + QUESTION_CLASSIFICATION_EXAMPLE.format(date=now_year.year)
    messages = [{"role": "system",
                 "content":system_template},
                {"role": "user", "content": question}]

    try:
        subscription = await openai_client.chat.completions.create(messages=messages,
                                                                   # model=os.environ['MODEL_NAME'],
                                                                   model=MODEL_NAME,
                                                                   stream=False, max_tokens=2048, temperature=0.2)
        # print(subscription)
        subscription_ = subscription.choices[0].message.content.strip().replace("['']","[]").replace("organization\":None","organization\":[]")
    except Exception as e:
        logger.error(e)
        subscription_ = None
    return eval(subscription_)


def get_data_from_database(sql_query,question) -> str:

    resp = requests.post(os.environ['STATISTIC_URL'], json=sql_query,cookies={'access_token':os.environ['TOKEN']}).json()
    table_name = question + '表'
    ret = {'table':[],'table_name':table_name}
    print(resp)
    value = 0
    try:
        for data in resp['data'][0]['data']:
            ret['table'].append({'name': data['name'],'value': data['value']})
            value += data['value']
        ret['table'].append({'name':'合计','value':value})


    except Exception as e:
        ret = []
        logger.error(e)
    print('result',str(ret))
    return str(ret)





async def mapping_project_info(project_info_params):
    normal_mapping = []
    organizations = project_info_params.get('organization')
    for key,value in project_info_params.items():
        if key in ('type','label','significance','status','bidding'):
            if value:
                normal_mapping.append(value)
    expend = project_info_params.get('expend')
    project_date = project_info_params.get('project_date')
    if organizations:
        project_info_params['organization'] = get_organization_mapping(organizations)
    """项目这里项目类型和项目标签、项目重要性、竞标状态等一起判断"""
    if normal_mapping:
        mapping_ret = get_normal_mapping(normal_mapping)
        project_info_params.update(mapping_ret)

    """项目经费"""
    if expend:
        project_info_params['expend'] = await get_question_expend(expend)

    """立项时间"""
    if project_date:
        project_info_params['project_date'] = get_project_date_mapping(project_date)


def get_llm_query(temperature: float, max_tokens: int, model_name: str,
                    system_template: str, user_template: str):
    query = {
        'model': model_name,
        'temperature': temperature,
        'messages': [
            {
                'role': 'system',
                'content': system_template
            },
            {
                'role': 'user',
                'content': user_template
            }
        ],
        'max_tokens': max_tokens
    }

    return query


def process_qa_history(history, query, question):
    messages = []
    max_tokens = MAX_TOKENS - MAX_NEW_TOKENS

    _history = []
    _length = 0
    history_zero_flag = False
    for idx in list(range(len(history)))[::-1]:
        _length += len(history[idx])
        _history.append(history[idx])
        if _length >= max_tokens and len(_history) > 0:
            _history.pop(-1)
            _length -= leng(history[idx])
            if len(_history) % 2 != 0:  # 去除system回答
                _history.pop(-1)
                _length -= len(history[idx])
            break    # 为什么去除一次历史就退出了？？？
        elif _length > max_tokens and len(_history) == 0:
            history_zero_flag = True
            break

    if history_zero_flag:
        history, _history_length = [], 0
    else:
        history, _history_length = _history[::-1],max(_length + 5 * len(_history), 0)  # 此处为什么要加5*len(_history)？？

    for idx in range(len(history)):
        messages.append({
            'role': 'user',
            'content': history[idx]
        })
    query['message'] = messages
    query['max_tokens'] = max(MAX_NEW_TOKENS, max_tokens - _history_length)
    return query


def generate_sql_query(sql_query, source, query_key, operateType):
    sql_query['queryList'].append({
      "queryKey": query_key,
      "queryValue": source,
      "operateType": operateType
    })

def generate_sql_group_by(sql_query, group_by, type, ext):
    sql_query['groupRule'] = {
        "classKey": "projectInformation",
        "propertyName": group_by,
        "groupType": type,
        "typeExt": ext
    }



def generate_query(project_info_params):
    ret = {
        'project_source': [],
        'charge_org': [],
        'admin_dept': [],
        'type': [],
        'label': [],
        'significance': [],
        'status': [],
        'bidding': [],
        'expend': {},
    }
    sql_query = {"classKey": "projectInformation", "queryList": [], "groupRule": {
        "classKey": "projectInformation",
        "propertyName": "project_source",
        "groupType": "syscategory",
        "typeExt": "{\"categoryKey\" : \"GP.ProjectFrom\"}"
    }, "statisticItemList": []}
    organizations = project_info_params.get('organization')
    if organizations:
        for organization in organizations:
            if PROJECT_SOURCE.get(organization) or PROJECT_SOURCE_ALIAS.get(organization):
                organization_index = PROJECT_SOURCE.get(organization) if PROJECT_SOURCE.get(organization) else PROJECT_SOURCE_ALIAS.get(organization)
                ret['project_source'].append(organization_index)
            elif organization in SUGGEST_DEPARTMENT:
                """TODO：根据api获取主管部门索引"""
                index = get_undertaking_unit(organization,'4')
                ret['admin_dept'].append(index)
            else:
                """TODO：根据api获取部门索引"""
                index = get_undertaking_unit(organization,'3')
                ret['charge_org'].append(index)
    if ret.get('project_source'):
        for source in ret['project_source']:
            generate_sql_query(sql_query, source, 'project_source', 'in')
    if ret.get('charge_org'):
        for org in ret['charge_org']:
            generate_sql_query(sql_query, org, 'charge_org', 'in')
    if ret.get('admin_dept'):
        for dept in ret['admin_dept']:
            generate_sql_query(sql_query, dept, 'suggest_admin_dept', 'in')
    if project_info_params.get('type'):
        for _type in project_info_params['type']:
            _type_index = TYPE.get(_type) if TYPE.get(_type) else TYPE_ALIAS.get(_type)
            generate_sql_query(sql_query, _type_index, 'project_type', 'in')
    if project_info_params.get('label'):
        for _label in project_info_params['label']:
            _label_index = LABEL.get(_label)
            generate_sql_query(sql_query, _label_index, 'project_flag', 'in')
    if project_info_params.get('significance'):
        for _significance in project_info_params['significance']:
            _significance_index = SIGNIFICANCE.get(_significance)
            generate_sql_query(sql_query, _significance_index, 'project_level', 'equals')
    if project_info_params.get('status'):
        for _status in project_info_params['status']:
            _status_index = STATUS.get(_status)
            generate_sql_query(sql_query, _status_index, 'project_status', 'equals')
    if project_info_params.get('bidding'):
        for _bidding in project_info_params['bidding']:
            _bidding_index = BIDDING.get(_bidding)
            generate_sql_query(sql_query, _bidding_index, 'projectBidding_bidding_status', 'in')
    if project_info_params.get('expend'):
        if project_info_params['expend'].get('compare_word') == 'less':
            generate_sql_query(sql_query, project_info_params['expend']['amount'], 'total_amount', 'lessThanOrEquals')
        elif project_info_params['expend'].get('compare_word') == 'greater':
            generate_sql_query(sql_query, project_info_params['expend']['amount'], 'total_amount', 'greaterThanOrEquals')
        elif project_info_params['expend'].get('compare_word') == 'equals':
            generate_sql_query(sql_query, project_info_params['expend']['amount'], 'total_amount','equals')
    if project_info_params.get('project_date'):
        generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')

    '''分组维度'''
    group_by = project_info_params.get('group_by')
    if group_by:
        if group_by == 'project_date':
            generate_sql_group_by(sql_query, "confirm_date", 'datetime', "{\"format\" : \"year\"}")
        else:
            if group_by == 'project_source':
                type = 'syscategory'
                ext =  "{\"categoryKey\" : \"GP.ProjectFrom\"}"
            elif group_by == 'project_type':
                type = 'syscategory'
                ext = "{\"categoryKey\" : \"GP.ProjectType\"}"
            elif group_by == 'project_level':
                type = 'sysdic'
                ext = "{\"dicKey\" : \"gpProjectStar\"}"
            elif group_by == 'project_status':
                type = 'sysdic'
                ext = "{\"dicKey\" : \"prjStatus\"}"
            else:
                type = 'datetime'
                ext = "{\"format\" : \"year\"}"
            generate_sql_group_by(sql_query, group_by, type, ext)

    result_form = project_info_params.get('result_form')
    if result_form:
        if result_form == 'count':
            sql_query['statisticItemList'] = [{
                "classKey": "projectInformation",
                "statisticType": "count"
            }]
        elif result_form == 'amount':
            sql_query['statisticItemList'] = [{
              "classKey": "projectInformation",
              "propertyName": "total_amount",
              "statisticType": "sum"
            }]
    print(sql_query)
    return sql_query



async def qa_stream(question: str, history: List[str]):
    
    """intent recognition"""
    intent_type = await get_user_intent(question)
    print(intent_type)
    if intent_type == '文档':

       prompt_tokens = sum(get_num_tokens([SECOND_SYSTEM_TEMPLATE_]))

       """获取关联docs"""
       relevant_documents = await search_with_same_outline(question)
       print('relvant')
       docs = defaultdict(lambda: {"name": "", "info": "", "contexts": []})

       for chunk in relevant_documents:
           docs[chunk['metadata']['document_id']]['contexts'].append(chunk)

       to_del = set()

       for doc_id in docs:
           contexts = []
           async for _batch_documents in auto_batch(prompt_tokens, docs[doc_id]['contexts'], ""):
               contexts.append(_batch_documents)
           docs[doc_id]['contexts'] = contexts

       for _doc_id in list(to_del):
           docs.pop(_doc_id)
       print(docs.keys())
       examples = []
       for val in docs.values():
           for _batch_documents in val['contexts']:
               page_content = "\n".join([_doc['page_content'] for _doc in _batch_documents])
               page_content = "Document Contexts: {}".format(page_content)
               examples.append(
                   {'page_content': page_content, 'metadata': {"name": val['name']}})

       async for batch_documents in auto_batch(prompt_tokens, examples):

           context = "\n\n".join([_doc['page_content'] for _doc in batch_documents])
           #print(context)

           user_answer_template = SECOND_HUMAN_TEMPLATE.format(question=question, context=context)
           print('111')
           max_tokens = MAX_TOKENS
           print('44') 
           query = get_llm_query(0.6, max_tokens, MODEL_NAME, SECOND_SYSTEM_TEMPLATE_, user_answer_template)
           print('222')
           query = process_qa_history(history, query, question)
           
           """获取query文档关联结果"""
           subscription = await get_docs_subscription(query)

           return subscription, rel_docs2refs(relevant_documents),[]

    elif intent_type == '数据库':
        if history:
            """根据history重新生成question"""
            question = await generate_question_by_history(question, history)

            """提取文本相关字段"""
        project_info_params = await get_question_classification_from_question(question)
        """映射项目数据"""
        await mapping_project_info(project_info_params)

        """生成查询"""
        sql_query = generate_query(project_info_params)

        print('sql:',sql_query)
        print(project_info_params)

        """查询api获取数据"""
        try:
           data = get_data_from_database(sql_query,question)
        except Exception as e:
           print(e)
           data = ''
        """根据数据生成返回结果"""
        print(data)
        llm_resp = await get_llm_response(question, data,is_text=False)
        return llm_resp, [], data
    elif intent_type == '生成报告':
        print('generate report')
        generate_data, data = await generate_report(question,history)
        total_llm_resp = await get_llm_response(question,generate_data,is_text=False)
        return total_llm_resp, [], data
    else:
        if history:
           question = await generate_question_by_history(question,history)
        total_llm_resp = await get_llm_response(question,{},is_text=False)
        return total_llm_resp,[],{}

def rel_docs2refs(rel_docs):
    ret = []
    refs = {}
    for index, document in enumerate(rel_docs):
        doc_id = document['metadata']['document_id']
        pages = document['metadata']['pages']
        coordinates = document['metadata'].get("coordinates", {})
        if len(coordinates) and isinstance(coordinates[0], List):
            coordinates = {0: {"points": coordinates, "page": 0, "width": 0, "height": 0}}
        elif isinstance(coordinates, List):
            coor = defaultdict(lambda: {"points": [], "width": 0, "height": 0, "page": 0})
            for _coor in coordinates:
                page = _coor.get("page", 0)
                coor[page]['points'].append(_coor['points'])
                coor[page]['width'] = _coor.get("width", 0)
                coor[page]['height'] = _coor.get("height", 0)
                coor[page]['page'] = _coor.get("page", 0)
            coordinates = coor
        if doc_id not in refs:
            try:
                doc_name = document['metadata'].get("source", '')
            except Exception as e:
                continue
            refs[doc_id] = {
                "doc_name": doc_name,
                "doc_id": doc_id,
                "pages": pages,
                "coordinates": coordinates,
                "content": [document['page_content']],
                "index": index
            }
        else:
            refs[doc_id]['pages'] = list(set(refs[doc_id]['pages'] + pages))
            for page, val in coordinates.items():
                coor = refs[doc_id]['coordinates']
                if page in coor:
                    points = coor[page].get("points", [])
                    points += val.get("points", [])
                    coor[page]['points'] = points
                else:
                    coor[page] = {"points": val.get("points", [])}
                    coor[page]['width'] = val.get('width', 0)
                    coor[page]['height'] = val.get("height", 0)
                    coor[page]['page'] = val.get("page", 0)
            refs[doc_id]['content'] += [document['page_content']]
    refs = sorted(list(refs.values()), key=lambda x: x['index'])
    for ref in refs:
        ret.append(ref)
    return ret


async def auto_batch(prompt_tokens: int, documents: List[Dict], prefix: str = None):
    """对documents根据num_tokens进行自动化的batch"""
    batch_tokens = 16000 - prompt_tokens
    if prefix:
        batch_tokens -= get_num_tokens([prefix])[0]
    """严格小于batch_tokens"""
    batch = []
    num_tokens = 0
    idx = 0
    while idx < len(documents):
        batch.append(documents[idx])
        document_num_tokens = documents[idx]['metadata'].get("num_tokens",
                                                             get_num_tokens([documents[idx]['page_content']])[
                                                                 0])
        num_tokens += document_num_tokens
        if num_tokens >= batch_tokens:
            if len(batch) <= 1:
                """强制进行truncate"""
                logger.warning(
                    "当前batch_size==1, 但其num_tokens ({}) > {}, 进行truncate".format(num_tokens, batch_tokens))
                batch[0]['page_content'] = batch[0]['page_content'][:int(batch_tokens * 0.9)]
            else:
                num_tokens -= document_num_tokens
                batch.pop(-1)
                idx -= 1
            yield batch
            batch = []
            num_tokens = 0
        idx += 1
    if batch:
        yield batch


def generate_report_data(project_info_params):
    print('project_info_params',str(project_info_params))
    data = {}
    ret = {
        'project_source': [],
        'charge_org': [],
        'admin_dept': [],
        'type': [],
        'label': [],
        'significance': [],
        'status': [],
        'bidding': [],
        'expend': {},
    }
    sql_query = {"classKey": "projectInformation", "queryList": [], "statisticItemList": [{
        "classKey": "projectInformation",
        "statisticType": "count"
    }]}

    if project_info_params.get('project_date'):
        generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    """获取正在进行的项目数量"""
    status_processing_sql = copy.deepcopy(sql_query)
    generate_sql_query(status_processing_sql, '1', 'project_status', 'equals')
    print('status_sql',str(status_processing_sql))
    processing_count = get_processing_count(status_processing_sql)
    """获取型号项目数量"""
    type_model_sql = copy.deepcopy(sql_query)
    generate_sql_query(type_model_sql, ["1402_14027"], 'project_type', 'in')
    print('type_sql:',str(type_model_sql))
    model_count = get_processing_count(type_model_sql)
    """获取涉火项目数量"""
    is_fire_sql = copy.deepcopy(sql_query)
    generate_sql_query(is_fire_sql, '1', 'is_fire', 'equals')
    print('is_fire_sql:',str(is_fire_sql))
    is_fire_count = get_processing_count(is_fire_sql)
    """获取重点跟踪数量"""
    importance_sql = copy.deepcopy(sql_query)
    generate_sql_query(importance_sql, ['ZDGC'], 'project_flag', 'in')
    print('import_sql:',str(importance_sql))
    importance_count = get_processing_count(importance_sql)
    """各个项目来源的数量"""
    projects_sql = copy.deepcopy(sql_query)
    generate_sql_group_by(sql_query, "project_source", 'syscategory', "{\"categoryKey\" : \"GP.ProjectFrom\"}")
    print('projects_sql:',str(projects_sql))
    projects_count = get_processing_count(projects_sql, False)
    """新增项目"""
    new_project_sql = copy.deepcopy(sql_query)
    if project_info_params.get('project_date'):
        generate_sql_query(new_project_sql, project_info_params['project_date'], 'confirm_date', 'between')
    else:
        generate_sql_query(new_project_sql, ['2024-01-01 00:00:00', '2024-12-31 23:59:59'], 'create_time', 'between')
    print('new_sql:',str(new_project_sql))
    new_project_count = get_processing_count(new_project_sql)
    """拖期项目"""
    delay_project_sql = copy.deepcopy(sql_query)
    generate_sql_query(delay_project_sql, ['TQXM'], 'project_flag', 'in')
    print('delay_sql:',str(delay_project_sql))
    delay_project_count = get_processing_count(delay_project_sql)

    """责任单位任务承担 dept - > count"""
    admin_dept_sql = copy.deepcopy(sql_query)
    generate_sql_query(admin_dept_sql, '1', 'project_admin_dept', 'in')
    print('admin_dept_sql:',str(admin_dept_sql))
    admin_dept_count = get_processing_count(admin_dept_sql,False)

    """重点项目执行情况"""

    """项目数量情况"""
    projects_source_count_sql = copy.deepcopy(sql_query)
    projects_source_count_sql['groupRule'] = {
        "classKey": "projectInformation",
        "propertyName": "project_source",
        "groupType": "syscategory",
        "typeExt": "{\"categoryKey\" : \"GP.ProjectFrom\"}"
    }
    print('project_source_count:',str(projects_source_count_sql))
    projects_source_count = get_processing_count(projects_source_count_sql, False)

    """项目经费执行情况"""
    projects_source_amount_sql = copy.deepcopy(sql_query)
    projects_source_amount_sql['groupRule'] = {
        "classKey": "projectInformation",
        "propertyName": "project_source",
        "groupType": "syscategory",
        "typeExt": "{\"categoryKey\" : \"GP.ProjectFrom\"}"
    }
    print('projects_amount_sql:',str(projects_source_amount_sql))
    projects_source_amount_sql['statisticItemList'] = [
        {
            "classKey": "projectInformation",
            "propertyName": "total_amount",
            "statisticType": "sum"
        }]
    projects_source_amount = get_processing_count(projects_source_amount_sql,False)
    data = {
        'processing_count': processing_count,
        'model_count': model_count,
        'is_fire_count': is_fire_count,
        'importance_count': importance_count,
        'projects_count': projects_count,       # project -> count
        'new_project_count': new_project_count,
        'delay_project_count': delay_project_count,
        'admin_dept_count': admin_dept_count,
        'projects_source_count': projects_source_count,
        'projects_source_amount': projects_source_amount
    }
    print(data)
    return data

def generate_bar_chart(document):


    attr = ["JT", "LJZYB", "ZQ", "GAD", "CTD", "CAD"]
    bar = Bar()
    bar.add_yaxis("YEAR: 2022", [5, 20, 36, 10, 75, 90], category_gap="50%")
    # bar.add_yaxis("2021年", [10, 25, 8, 60, 20, 80], category_gap="50%")
    bar.add_xaxis(attr)

    # bar.add("2021年", attr, v1, is_stack=True)
    # bar.add("2022年", attr, v2, is_stack=True)


    snapshot.PHANTOMJS_EXEC = r"rag/phantomjs-2.1.1-linux-x86_64/bin/phantomjs"
    try:
       make_snapshot(snapshot, bar.render(), "rag/static/bar_chart.png")
    except Exception as e:
       print(e)
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run('项目来源情况柱状图')
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('rag/static/bar_chart.png', width=Cm(15), height=Cm(8))

def get_random_md5():
    random_byts = os.urandom(16)
    md5 = hashlib.md5()
    md5.update(random_byts)
    return md5.hexdigest()

async def generate_document(question,generate_data):
    totol_desc_resp = await get_llm_response('根据现有数据，生成总体概览报告分析',str(generate_data['total_description']),False)
    document = Document()
    document.add_picture('rag/static/chart.png',width=Cm(15.2), height=Cm(22.9))
    paragraph = document.add_paragraph()
    total_desc_run = paragraph.add_run('总体概况:')
    total_desc_run.bold = True
    # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run(totol_desc_resp)

    projects_source_run = paragraph.add_run('项目来源情况')
    projects_source_run.bold = True
    '''生成摘要'''
    projects_source_abstract = await get_llm_response('总结项目来源情况',generate_data['project_sources'],stream=False)
    projects_source_run.add_text(projects_source_abstract)
    generate_bar_chart(document)

    admin_dept_run = paragraph.add_run('责任单位任务承担情况')
    admin_dept_run.bold = True
    '''生成摘要'''
    admin_dept_abstract = await get_llm_response('总结责任单位任务承担情况',generate_data['admin_dept'],stream=False)
    admin_dept_run.add_text(admin_dept_abstract)
    generate_bar_chart(document)

    '''项目经费情况'''
    project_amount_run = paragraph.add_run('项目经费执行情况')
    '''项目经费只要生成'''
    amount_abstract = await get_llm_response('总结项目经费执行情况',generate_data['project_amount'],stream=False)
    project_amount_run.bold = True
    project_amount_run.add_text(amount_abstract)


    md5 = get_random_md5()

    doc_static_path = r'rag/static/docx'
    if not os.path.exists(doc_static_path):
        os.mkdir(doc_static_path)
    file_path = os.path.join(doc_static_path,str(md5) + '.docx')
    print(file_path)
    try:
       document.save(file_path)
    except Exception as e:
       print(e)
    print('save done')
    file_path = file_path.split('rag/')[1]
    return {'generate_docx':file_path}


def transfor_report_data2json(report_data):
    ret = {}
    """总体概况"""
    ret['total_description'] = ('项目总体概况：项目状态为正在执行的数量：' + str(report_data['processing_count'])
                                + '，型号研制项目数量：' + str(report_data['model_count']) + '，涉火项目数量：' +
                                str(report_data['is_fire_count']) + '，重点跟踪数量：' + str(report_data['importance_count'])
                                + '，新增项目数量：' + str(report_data['new_project_count']) + '，拖期项目数量：' +
                                str(report_data['delay_project_count']) + '。' + str(report_data['projects_count']))
    """各个项目来源的数量"""
    ret['project_sources'] = report_data['projects_count']

    '''责任单位任务承担'''
    ret['admin_dept'] = report_data['admin_dept_count']

    '''项目经费情况'''
    ret['project_amount'] = report_data['projects_source_amount']

    return ret


async def generate_report(question: str, history: List[str]):
    if history:
        """根据history重新生成question"""
        question = await generate_question_by_history(question, history)

    """提取文本相关字段"""
    project_info_params = await get_question_classification_from_question(question)

    await mapping_project_info(project_info_params)

    """生成查询"""
    report_data = generate_report_data(project_info_params)

    """查询结果生成数据"""
    generate_data = transfor_report_data2json(report_data)

    """data生成图表，生成document"""
    file_path = await generate_document(question,generate_data)

    return generate_data, file_path
