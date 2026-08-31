import asyncio
import json
import pathlib
import random
import uuid
from functools import wraps
from typing import Dict

import markdown
import requests
from docx import Document
from starlette import status
from starlette.responses import JSONResponse

from rag.configs import app_config
from rag.mappers.task import Dialogue
from rag.utils.request_llm import get_api_example_input, get_question_intent_api
API_URL = app_config['API_URL']
api_config_path = pathlib.Path(__file__).parent.parent / 'configs' / 'api_config.json'
api_data_path = pathlib.Path(__file__).parent.parent / 'configs' / 'api_data.json'


def requests_upload_file(package_id, ftp_url, user_id):
    # TODO:改为解析服务地址
    resp = requests.post(API_URL + "/zip/upload_ftp", params={'package_id': package_id, 'ftp_url': ftp_url, 'user_id': user_id})
    print(resp.text)


def async_wrap(func):
    """将同步函数包装为异步执行的装饰器"""
    @wraps(func)
    async def wrapper(*args, **kwargs):
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            None,
            lambda: func(*args, **kwargs)
        )

    return wrapper


def construct_response(data, detail='成功', status_code=status.HTTP_200_OK, headers=None):
    return {'data': data, 'detail': detail, 'status_code': status_code}


def create_uuid():
    return str(uuid.uuid5(uuid.NAMESPACE_DNS, str(uuid.uuid1()) + str(random.random())))


def transform_markdown2docx(doc: Document, markdown_text: str):

    html = markdown.markdown(markdown_text)

    # 将 HTML 解析为 DOCX
    for line in html.splitlines():
        if line.startswith('<h1>'):
            doc.add_heading(line[4:-5], level=1)  # 添加一级标题
        elif line.startswith('<h2>'):
            doc.add_heading(line[4:-5], level=2)  # 添加二级标题
        elif line.startswith('<p>'):
            p_line = line[3:-4]
            p_paragraph = doc.add_paragraph()
            if p_line.startswith('<strong>'):
                p_paragraph.add_run(p_line[8:-9]).bold = True  # 添加加粗文本
            else:
                p_paragraph.add_run(p_line[3:-4])  # 添加段落
        elif line.startswith('<ul>'):
            continue  # 跳过无用的列表标签
        elif line.startswith('<li>'):
            doc.add_paragraph(line[4:-5], style='ListBullet')  # 添加列表项
        elif line.startswith('<strong>'):
            doc.add_paragraph(line[8:-9], style='IntenseQuote')  # 添加加粗文本

def get_api_config():
    with open(api_config_path, 'r', encoding='utf-8') as file:
        config = json.load(file)
    return config

def save_api_config(config: Dict):
    with open(api_config_path, 'w', encoding='utf-8') as file:
        json.dump(config, file, ensure_ascii=False, indent=4)



def insert_api_field(api_name,type, field_data):
    configs = get_api_config()
    api_info = configs.get(api_name)
    if api_info:
        insert_flag = True
        for field in api_info['fields'][type]:
            if field_data['name'] == field['name']:
                insert_flag = False
                break
        if insert_flag:
            api_info['fields'][type].append(field_data)
            save_api_config(configs)
    api_data = get_api_data_by_name(api_name)
    if api_data:
        api_data['fields'].update({field_data['name']:[field_data['example']]})
        api_data['alias2name'].update({field_data['example']: field_data['example']})
        save_api_data(api_name, api_data)




def get_api_field_desc(api_name: str):
    field_list = []
    config = get_api_config()
    api_info = config.get(api_name)
    if api_info:
        for name, fields in api_info['fields'].items():
            if name == 'search':
                for field in fields:
                    if field.get('is_time') or isinstance(field.get('type'),list):
                        continue
                    field_list.append(field)
    return field_list



def get_api_field_name_example(name: str):
    filed_list = {}
    config = get_api_config()
    api_info = config.get(name)
    if api_info:
        for name, fields in api_info['fields'].items():
            if name == 'search':
                for field in fields:
                    filed_list.update({field['name']: [field['example']]})
    return filed_list


def get_api_config_by_name(api_name: str):
    configs = get_api_config()
    api_info = configs.get(api_name)
    return api_info

def get_api_field(name: str):
    filed_dict = {}
    config = get_api_config()
    api_info = config.get(name)
    if api_info:
        for name, fields in api_info['fields'].items():
            if name =='search':
                for field in fields:
                    filed_dict.update({field['zh_']:field})
            elif name == 'group':
                filed_dict['分组'] = {'example':[] ,'name': []}
                for field in fields:
                    filed_dict['分组']['example'].append(field['example'])
                    filed_dict['分组']['name'].append(field['name'])
            elif name == 'order':
                filed_dict['结果形式'] = {'example':[] ,'name': []}
                for field in fields:
                    filed_dict['结果形式']['example'].append(field['example'])
                    name = field['name'] if field.get('name') else ''
                    filed_dict['结果形式']['name'].append(name)

    return filed_dict

def get_api_field_name2zh(name: str):
    filed_dict = {}
    config = get_api_config()
    api_info = config.get(name)
    if api_info:
        for name, fields in api_info['fields'].items():
            if name == 'search':
                for field in fields:
                    filed_dict.update({field['name']: field['zh_']})

    return filed_dict


def get_api_field_name2zh_time_and_type(name: str):
    filed_dict = {}
    config = get_api_config()
    api_info = config.get(name)
    if api_info:
        for name, fields in api_info['fields'].items():
            if name == 'search':
                for field in fields:
                    if field.get('is_time') or len(field.get('type')) > 1:
                        filed_dict.update({field['zh_']: field['example']})

    return filed_dict


def get_prompt_example(prompt_example, fields, is_search=False, example_num=2):
    field_list = []
    is_time_field = []
    for item in fields:
        if not item.get('is_time') and isinstance(item.get('type'),str):
            field_list.append({item['zh_']: item['example']})
        else:
            is_time_field.append({item['zh_']:item['example']})
    # field_list = [{item['zh_']: item['example']} for item in fields if not item['is_time']]
    if is_search:
        num2select = random.randint(1, min(3,len(fields) - 1))
        for i in range(example_num):
            prompt_example[i] = random.sample(field_list, num2select -1)
            if is_time_field:
                prompt_example[i].extend(is_time_field)
    else:
        for i in range(example_num):
            prompt_example[i].extend(random.sample(field_list, 1))
    return prompt_example



def generate_api_example(name: str):
    prompt_example = ''
    config = get_api_config()
    api_info = config.get(name)
    if api_info:
        example_num = 2
        prompt_example_input = [None] * example_num
        for name, fields in api_info['fields'].items():
            if name == 'classKey':
                continue
            if name == 'search':
                get_prompt_example(prompt_example_input, fields, is_search=True, example_num=example_num)
            else:
                get_prompt_example(prompt_example_input, fields, is_search=False, example_num=example_num)

        for example in prompt_example_input:
            prompt_example += '\t输出: ' + str(example)
        # '''生成example输入'''
        # prompt_example_output = await get_api_example_input(prompt_example_input)
        # prompt_example = combination_example(prompt_example_output, prompt_example_input)

    #else:
        '''未找到API信息'''

    return prompt_example


async def get_api_name_by_intent(question):
    api_infos = get_api_config()
    api_list = [{name: api_info.get('info',{}).get('desc','') for name, api_info in api_infos.items()}]
    api_name = await get_question_intent_api(question, api_list)
    return api_name


def generate_alias2name(field_list):
    ret = {}
    for name,data in field_list.items():
        ret.update({
            data[0]:data[0]
        })

    return ret

def get_api_data():
    with open(api_data_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data

def save_api_data(api_name, data):
    api_data = get_api_data()
    api_data[api_name] = data
    with open(api_data_path, 'w', encoding='utf-8') as file:
        json.dump(api_data, file, ensure_ascii=False, indent=4)

def get_api_data_by_name(api_name):
    api_data = get_api_data()
    data = api_data.get(api_name)
    return data


def transfrom_alias2name(value,api_name):
    data = get_api_data_by_name(api_name)
    if data:
        alias2name = data.get('alias2name')
        if isinstance(value, list):
            ret = []
            for item in value:
                ret.append(alias2name.get(item,""))
            return ret
        else:
            name = alias2name.get(value)
            if name:
                return name 
            else:
                return value
    return ''


def get_category_schema(api_name: str):
    res = {}
    name2zh = get_api_field_name2zh(api_name)
    data = get_api_data_by_name(api_name)
    #print('api_schema_data:',data)
    if data:
        for name,_info in data['fields'].items():
            _name = name2zh.get(name)
            if _name:
                res.update({_name:_info})
    #print('api_schema_data:',res)
    return res


def generate_question_classification_example(api_name: str):
    ret = ''
    res = get_category_schema(api_name)
    if res:
        for name, info in res.items():
            if info:
                # ret += name + "：从以下枚举值中抽取：" + '、'.join(info) + '\n'
                ret += name + "：" + '、'.join(info) + '\n'
            else:
                ret += name + '：'
    return ret



def get_api_data_fields(api_name, field):
    ret = {}
    data = get_api_data_by_name(api_name)
    if data:
        ret = data.get(field)
    return ret

def get_api_updatable_fields(api_name):
    ret = []
    api_config = get_api_config_by_name(api_name)
    if api_config:
        if api_config.get('fields',{}).get('search',{}):
            for field in api_config['fields']['search']:
                if not field.get('is_time') and not isinstance(field.get('type'),list):
                    ret.append({
                        'name': field['name'],
                        'zh_': field['zh_']
                    })
        else:
            return []
    return ret


def save_api_data_fields(api_name, field, data):
    api_data = get_api_data_by_name(api_name)
    if api_data:
        api_data.update({field: data})
    save_api_data(api_name, api_data)


def get_api_data_name2index(api_name):
    ret = {}
    data = get_api_data_by_name(api_name)
    if data:
        ret = data.get('name2index')
    return ret


def get_data_union(api_data_field, alias_data, name_data):
    api_data_field = list(set(api_data_field).union(set(alias_data), [name_data]))
    return api_data_field

def insert_api_data(api_name, field_name, field_data):
    try:
        name_data = field_data.get('name')
        alias_data = field_data.get('alias')
        index_data = field_data.get('index')
        api_data_fields = get_api_data_fields(api_name, 'fields')
        api_data_field = api_data_fields.get(field_name)
        '''更新到api_fields'''
        if api_data_field:
            api_data_fields[field_name] = get_data_union(api_data_field, alias_data, name_data)
        else:
            _data = list(set(alias_data).union([name_data]))
            api_data_fields.update({field_name: _data})
        '''更新到api_alias2name'''
        api_data_alias2name = get_api_data_fields(api_name, 'alias2name')
        if alias_data:
            for alias in alias_data:
                api_data_alias2name.update({alias: name_data})
        api_data_alias2name.update({name_data: name_data})

        '''name2index'''
        api_data_name2index = get_api_data_fields(api_name, 'name2index')
        if api_data_alias2name:
            api_data_name2index.update({name_data: index_data})
        save_api_data_fields(api_name, 'alias2name', api_data_alias2name)
        save_api_data_fields(api_name, 'fields', api_data_fields)
        save_api_data_fields(api_name, 'name2index', api_data_name2index)
    except Exception as e:
        print(e)
        print('插入api字段数据失败')

def api_data_update(api_name, field_name, field_data):
    org_data = field_data.get('org_data')
    new_data = field_data.get('new_data')
    name_data = field_data.get('name')
    org_alias_data = org_data.get('alias')
    org_index_data = org_data.get('index')
    new_alias_data = new_data.get('alias')
    new_index_data = new_data.get('index')


    greater_alias_data = [item for item in new_alias_data if item not in org_alias_data]
    lower_alias_data = [item for item in org_alias_data if item not in new_alias_data]
    greater_index_data = [item for item in new_index_data if item not in org_index_data]
    lower_index_data = [item for item in org_index_data if item not in new_index_data]

    api_data_fields = get_api_data_fields(api_name, 'fields')
    api_name2index = get_api_data_fields(api_name, 'name2index')
    api_alias2name= get_api_data_fields(api_name, 'alias2name')
    api_data_field = api_data_fields.get(field_name)
    """先拿到比原先少的数据"""
    if greater_alias_data:
        api_data_field.extend(greater_alias_data)
        api_data_fields[field_name] = api_data_field
        for alias in greater_alias_data:
            api_alias2name[alias] = name_data
    if lower_alias_data:
        api_data_fields[field_name] = [item for item in api_data_field if item not in lower_alias_data]
        for alias in lower_alias_data:
            if api_alias2name.get(alias):
                api_alias2name.pop(alias)
    if api_name2index.get(name_data):
        if greater_index_data:
            api_name2index[name_data].extend(greater_index_data)
        if lower_index_data:
            api_name2index[name_data] = [item for item in api_name2index.get(name_data) if item not in lower_index_data]

    save_api_data_fields(api_name, 'alias2name', api_alias2name)
    save_api_data_fields(api_name, 'fields', api_data_fields)
    save_api_data_fields(api_name, 'name2index', api_name2index)



def delete_api_data_by_fields(api_name, field_name,data_list):
    api_data_fields = get_api_data_fields(api_name, 'fields')
    if api_data_fields:
        api_data_field = api_data_fields.get(field_name)
        if api_data_field and len(api_data_field) > len(data_list):
            api_data_fields[field_name] = [item for item in api_data_field if item not in data_list]
        else:
            return
    api_data_alias2name = get_api_data_fields(api_name, 'alias2name')
    new_api_data_fields = {}
    if api_data_alias2name:
        for alias, name in api_data_alias2name.items():
            if name not in data_list:
                new_api_data_fields.update({alias: name})
    api_data_name2index = get_api_data_fields(api_name, 'name2index')
    new_api_data_name2index = {}
    if api_data_name2index:
        for name, index in api_data_name2index.items():
            if name not in data_list:
                new_api_data_name2index.update({name: index})
    save_api_data_fields(api_name, 'alias2name', new_api_data_fields)
    save_api_data_fields(api_name, 'fields', api_data_fields)
    save_api_data_fields(api_name, 'name2index', new_api_data_name2index)


def get_api_data_list(page_on, page_size, api_name, field_name):
    ret = []
    name2alias = {}
    # index_type_field = [item['name'] for item in get_api_field_desc(api_name) if item.get('index_type')]
    name2index = get_api_data_fields(api_name, 'name2index')
    alias2name = get_api_data_fields(api_name, 'alias2name')
    for alias, name in alias2name.items():
        if name2alias.get(name):
            name2alias[name].append(alias)
        else:
            name2alias.update({name: [alias]})
    api_data_fields = get_api_data_fields(api_name, 'fields')
    if api_data_fields:
        api_data_field = api_data_fields.get(field_name)
        for item in api_data_field:
            # if item in index_type_field:
            #     continue
            if name2alias.get(item):
                ret.append({
                    'name': item,
                    'alias': name2alias.get(item),
                    'index': name2index.get(item)
                })
    start_index = page_on * page_size - page_size
    end_index = page_on * page_size
    api_data_list = ret[start_index:end_index]
    ret = {
        'api_data_list': api_data_list,
        'total_count': len(api_data_list),
        'page_on': page_on,
        'page_size': page_size
    }

    return ret


import datetime
task_status_path = pathlib.Path(__file__).parent.parent / 'static' / 'log' / 'task_status.json'

def get_task_status():
    with open(task_status_path, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data

def save_task_status(data):
    with open(task_status_path, 'w', encoding='utf-8') as file:
        json.dump(data, file, ensure_ascii=False, indent=4)


def get_task_by_id(task_id):
    task_status = get_task_status()
    if task_status:
        task_status = task_status.get(task_id)
        return task_status
    else:
        return {}


def get_task_status_by_id(task_id):
    _task = Dialogue.get_single_dialogue(id=task_id)
    if _task:
        return _task.status
    else:
        return ''


def save_task_status_by_id(task_id, data):
    task_status = get_task_status()
    task_status.update({task_id: data})
    save_task_status(task_status)


def add_task_status(task_id, status_name):
    _task = Dialogue.get_single_dialogue(id=task_id)
    if _task:
        _task.status = status_name
        _task.update()

    # if status_name == 'create':
    #     task_status = {
    #         'task_id': task_id,
    #         'create_time': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
    #         'end_time': '',
    #         'task_status': "intent",
    #     }
    # else:
    #     task_status = get_task_by_id(task_id)
    #     if task_status:
    #         task_status.update({'task_status': status_name})
    #
    # save_task_status_by_id(task_id, task_status)
