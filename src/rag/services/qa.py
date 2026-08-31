import copy
import datetime
import hashlib
import json
import logging
import os
import re
from collections import defaultdict
from importlib.resources import is_resource
from sys import flags
from typing import Union, List, Dict
import traceback
import openai
import pathlib
import requests
import  uuid
import asyncio
import pathlib
import markdown
import pandas as pd
from bs4 import BeautifulSoup
from openai import organization
from requests.auth import HTTPBasicAuth
from openpyxl.cell.cell import MergedCell

from openpyxl.styles import  Border, Side, PatternFill, Font, Alignment

from  datetime import datetime,timedelta
import random

import json
import requests
from openpyxl import load_workbook
from openpyxl.styles import Border, Side, PatternFill
from openpyxl.utils import get_column_letter


from concurrent.futures import ThreadPoolExecutor, as_completed

from fastapi import APIRouter, Body, HTTPException,File, UploadFile
import  tempfile
from docx import Document, oxml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, RGBColor, Pt, Inches
from pyecharts.charts import Bar
from pyecharts.render import make_snapshot
from pymilvus import MilvusClient
from snapshot_phantomjs import snapshot
from pyecharts.globals import CurrentConfig


from rag.mappers.knowledge import File
from rag.mappers.task import Dialogue
from rag.mappers.team_info import TeamInfo
from rag.services.api_manage import get_api_url

# 设置为本地min.js
CurrentConfig.ONLINE_HOST = ""
from rag.configs import app_config, prompt_config
from rag.utils.mapping import get_project_date_mapping, rerank
from rag.utils.request_llm import generate_question_by_history, get_llm_other_response, generate_ocr_result, \
    get_processing_count, get_week_processing_count, get_week_problem_processing_count, get_week_risk_processing_count, \
    get_week_delay_processing_count, get_type_mapping, get_index_by_bq_api, get_sql_llm_response, get_maybe_questions, \
    get_task_divide_from_question, get_table_info, get_year_by_question, generate_title_analysis, get_docs_subscription, \
    generate_question_analysis, generate_delay_analysis, generate_normal_result, generate_markdown_normal_result, \
    generate_graph_normal_result, generate_ocr_org_result, generate_jira_normal_result, generate_request_normal_result, \
    generate_jira_normal_result_report, generate_stream_result, generate_contract_result, \
    generate_contract_end_json_result,generate_ocr_team_result,generate_mermaid_result
from rag.utils.utils import generate_api_example, get_api_field, \
    get_api_config_by_name, get_api_name_by_intent, \
    generate_question_classification_example, transfrom_alias2name, get_api_data_name2index, add_task_status, \
    get_category_schema

from rag.services.user_config import get_user_config


now_year = datetime.now()
logger = logging.getLogger(__name__)

OCR_URL = app_config['OCR_URL']
JIRA_URL = app_config['JIRA_URL']
JIRA_USER = app_config['JIRA_USER']
JIRA_TOKEN = app_config['JIRA_TOKEN']
JIRA_BOARD = app_config['JIRA_BOARD']
MILVUS_COLLECTION = app_config['MILVUS_COLLECTION']
MILVUS_URI = app_config['MILVUS_URI']
RERANK_URL = app_config['RERANK_URL']
EMBEDDING_URL = app_config['EMBEDDING_URL']
MODEL_NAME = app_config['MODEL_NAME']
API_KEY = app_config['API_KEY']
API_BASE_URL = app_config['API_BASE_URL']
TOKENIZE_URL = app_config['TOKENIZE_URL']
MAX_TOKENS = app_config['MAX_TOKENS']
GET_TOKEN_URL = app_config['GET_TOKEN_URL']
MAX_NEW_TOKENS = app_config['MAX_NEW_TOKENS']
TOKEN = app_config['TOKEN']
REQUEST_URL = app_config['REQUEST_URL']

INTENT_PROMPT = prompt_config['intent_prompt']
DB_AGENT_INTENT_PROMPT = prompt_config['db_agent_intent_prompt']
SECOND_SYSTEM_TEMPLATE_ = prompt_config['second_system_template_']
# SECOND_HUMAN_TEMPLATE = prompt_config['second_human_template']
QUESTION_CLASSIFICATION_PROMPT = prompt_config['question_classification_prompt']

milvus_client = MilvusClient(MILVUS_URI)
openai_client = openai.AsyncClient(api_key=API_KEY,
                                   base_url=API_BASE_URL)

RERANK_POOL_MAX = 50


MODEL_NAME2 = app_config['MODEL_NAME2']
API_KEY2 = app_config['API_KEY2']
API_BASE_URL2 = app_config['API_BASE_URL2']

openai_client2 = openai.AsyncClient(api_key=API_KEY2,
                                   base_url=API_BASE_URL2)

BQ_DATA_DESC_URL = app_config['BQ_DATA_DESC_URL']


def get_num_tokens(texts: List[str]) -> List[int]:
    num_tokens = []
    for text in texts:
        num_tokens.append(len(text))
        continue
        resp = requests.post(TOKENIZE_URL, json={
            "model": MODEL_NAME,
            "prompt": text,
            "add_special_tokens": True
        })
        num_tokens.append(len(resp.json()['tokens']))
    return num_tokens


def get_document_rerank(data):
    url = RERANK_URL
    resp = requests.post(url, json=data)
    resp = resp.json()
    return resp['scores']


def get_vectors(texts: List[str]) -> List[List[float]]:
    url = EMBEDDING_URL
    response = requests.post(url, json={"input": texts})
    vectors = []
    for item in response.json()['data']:
        vectors.append(item['embedding'])
    return vectors


def get_length(text: str):
    return len(text)
    token_count = float('inf')
    try:
        resp = requests.post(GET_TOKEN_URL, json={
            "text": text,
            "model": MODEL_NAME
        }, timeout=(2, 30))
        token_count = json.loads(json.loads(resp.text)).get('tokencount', 0)
    except Exception as e:
        logger.error(e)
    return token_count


async def search(
        text: Union[str, List],
        filter: str = "",
        limit: int = 10) -> List[Dict]:
    output_fields = ["id", "page_content", "pages", "coordinates", "outline", "parent_id", "index", "type", 'file_name',
                     "document_id"]

    if isinstance(text, str):
        query_embed = get_vectors([text])[0]
    else:
        query_embed = text
    result = milvus_client.search(MILVUS_COLLECTION,
                                  data=[query_embed],
                                  limit=limit,
                                  filter=filter,
                                  output_fields=output_fields)[0]

    docs = []
    for item in result:
        page_content = item['entity']['page_content']
        del item['entity']['page_content']
        metadata = item['entity']
        metadata['id'] = item['id']
        docs.append({'page_content': page_content,
                     'metadata': metadata})
    return docs


def query(filter: str = "") -> List[Dict]:
    output_fields = ["id", "page_content", "pages", "coordinates", "outline", "document_id", "file_name",
                     "parent_id", "index", "type"]
    result = milvus_client.query(MILVUS_COLLECTION, filter=filter, output_fields=output_fields)
    ret = []
    remove_context = set()
    for item in result:
        page_content = item['page_content']
        if page_content in remove_context:
            continue
        remove_context.add(page_content)
        del item['page_content']
        metadata = item
        ret.append({'page_content': page_content,
                    'metadata': metadata})
    ret = sorted(ret, key=lambda x: x['metadata']['index'])
    return ret


def merge_documents(primary: List[Dict], secondary: List[Dict]) -> List[Dict]:
    """Merge document lists; keep primary order, dedupe by fragment id or content."""
    seen = set()
    merged = []
    for doc in primary + secondary:
        key = doc['metadata'].get('id') or doc['page_content']
        if key in seen:
            continue
        seen.add(key)
        merged.append(doc)
    return merged


def cap_for_rerank(docs: List[Dict], max_pool: int = RERANK_POOL_MAX) -> List[Dict]:
    if len(docs) <= max_pool:
        return docs
    return docs[:max_pool]


def remove_abstract_documents(documents: List[Dict]) -> List[Dict]:
    abstract_parent_id = None
    for _doc in documents:
        if _doc['page_content'].strip().startswith("摘要") or _doc['page_content'].strip().lower().startswith(
                "abstract"):
            abstract_parent_id = _doc['metadata']['parent_id']
    if abstract_parent_id:
        documents = [_doc for _doc in documents if _doc['metadata']['parent_id'] != abstract_parent_id]
    return documents


async def rerank_documents(query: str,
                           documents: List[Dict],
                           merge_special: bool = False):
    raw_documents = None
    if merge_special:
        """将documents中的文档中的特殊列表进行组装"""
        sorted_documents = sorted(documents, key=lambda x: x['metadata']['index'])
        """用parent_id进行分组"""
        parent_id2documents = defaultdict(list)
        for _document in sorted_documents:
            parent_id2documents[_document['metadata']['parent_id']].append(_document)

        _documents = []
        for parent_id, _sorted_documents in parent_id2documents.items():
            idx = 0
            while idx < len(_sorted_documents):
                to_merge = []
                if re.search("[：:]$", _sorted_documents[idx]['page_content'].strip()):
                    jdx = idx + 1
                    while jdx < len(_sorted_documents) and \
                            re.search("^[(（\[【]?[0-9一二三四五六七八九十]+[)）\]】]",
                                      _sorted_documents[jdx]['page_content'].strip()):
                        to_merge.append(_sorted_documents[jdx])
                        jdx += 1
                    if to_merge:
                        to_merge.insert(0, _sorted_documents[idx])
                        _metadata = {"ids": [_doc['metadata']['id'] for _doc in to_merge]}
                        _documents.append({
                            "page_content": "\n".join([_doc['page_content'] for _doc in to_merge]),
                            "metadata": _metadata
                        })
                    idx = jdx - 1
                else:
                    _documents.append(_sorted_documents[idx])
                idx += 1
        raw_documents = documents
        documents = _documents

    scores = get_document_rerank({"query": query, "texts": [document['page_content'] for document in documents]})

    """根据scores的正负值区分"""
    ret = []
    for idx in range(len(documents)):
        if scores[idx] > 0:
            ret.append((scores[idx], documents[idx]))
    ret = [x[1] for x in sorted(ret, key=lambda x: x[0], reverse=True)]
    if raw_documents:
        _ret = []
        for doc in ret:
            if 'ids' in doc['metadata']:
                _ret.extend([_doc for _doc in raw_documents if _doc['metadata']['id'] in doc['metadata']['ids']])
            else:
                _ret.append(doc)
        ret = _ret
    return ret


async def search_with_same_outline(
        text: Union[str, List], _filter: str = '',
        with_rerank: bool = True, top_k: int = 8) -> List[Dict]:
    vector_docs = await search(text, filter=_filter, limit=30)
    docs = vector_docs

    parent_ids = list({
        item['metadata']['parent_id']
        for item in vector_docs
        if item['metadata'].get('parent_id') not in (None, 'None', '')
    })
    if parent_ids:
        sibling_expr = "({}) and type == 'text'".format(
            " or ".join(["parent_id LIKE '%{}%'".format(_pid) for _pid in parent_ids])
        )
        if _filter:
            sibling_expr = "({}) and ({})".format(_filter, sibling_expr)
        sibling_docs = query(sibling_expr)
        docs = merge_documents(vector_docs, sibling_docs)
        print('merged docs len:', len(docs),
              '(vector:', len(vector_docs), 'siblings:', len(sibling_docs), ')')

    if with_rerank and docs and isinstance(text, str):
        pool = cap_for_rerank(docs)
        if len(pool) < len(docs):
            print('rerank pool capped:', len(docs), '->', len(pool))
        reranked = await rerank_documents(text, pool)
        if reranked:
            docs = reranked[:top_k] if top_k else reranked
        elif top_k:
            docs = vector_docs[:top_k]

    return docs


def docs_to_retrieval_hits(docs: List[Dict]) -> List[Dict]:
    hits = []
    for rank, doc in enumerate(docs, start=1):
        meta = doc['metadata']
        hits.append({
            'rank': rank,
            'id': meta.get('id'),
            'document_id': meta.get('document_id'),
            'file_name': meta.get('file_name'),
            'index': meta.get('index'),
            'type': meta.get('type'),
            'outline': meta.get('outline'),
            'parent_id': meta.get('parent_id'),
            'distance': meta.get('distance'),
            'page_content': doc.get('page_content', ''),
        })
    return hits


async def retrieval_search(
        query: str,
        package_id: str = "",
        document_ids: List[str] = None,
        mode: str = "pipeline",
        limit: int = 30,
        top_k: int = 8) -> Dict:
    """Run retrieval without LLM. Modes: vector, outline, pipeline."""
    if package_id:
        file_ids = [_file.file_id for _file in File.get_by(package_id=package_id)]
        if not file_ids:
            return {'query': query, 'mode': mode, 'count': 0, 'hits': []}
        _filter = "document_id in {}".format(file_ids)
    elif document_ids:
        _filter = "document_id in {}".format(document_ids)
    else:
        _filter = ""

    mode = mode.lower()
    if mode == "vector":
        docs = await search(query, filter=_filter, limit=limit)
    elif mode == "outline":
        docs = await search_with_same_outline(query, _filter, with_rerank=False)
    elif mode == "pipeline":
        docs = await search_with_same_outline(
            query, _filter, with_rerank=True, top_k=top_k)
    else:
        raise ValueError("mode must be vector, outline, or pipeline")

    return {
        'query': query,
        'mode': mode,
        'package_id': package_id or None,
        'document_ids': document_ids,
        'count': len(docs),
        'hits': docs_to_retrieval_hits(docs),
    }


async def retrieval_search_by_file_ids(
        query: str,
        ids: List[str],
        mode: str = "pipeline",
        limit: int = 30,
        top_k: int = 8) -> Dict:
    file_ids = []
    for _file in File.get_by(id=ids):
        if _file.file_id and _file.file_id not in file_ids:
            file_ids.append(_file.file_id)
    return await retrieval_search(
        query=query,
        document_ids=file_ids,
        mode=mode,
        limit=limit,
        top_k=top_k
    )


async def get_user_intent(question: str, history: List[str]) -> str:
    intent_template = INTENT_PROMPT.format(question=question, history=history)
    messages = [{
        "role": "system",
        "content": intent_template
    },
        {
            "role": "user",
            "content": question
        }]
    llm_response = await openai_client.chat.completions.create(messages=messages,
                                                               model=MODEL_NAME,
                                                               stream=False, temperature=0.2)
    try:
        resp_json = llm_response.choices[0].message.content.strip()
        print("intent raw:", repr(resp_json), flush=True)

        if "</think>" in resp_json:
            resp_json = resp_json.split("</think>", 1)[1].strip()

        resp_json = resp_json.replace("```json", "").replace("```", "").strip()
        intent = eval(resp_json).get('type', '文档')
    except Exception as e:
        logger.error(e)
        intent = '其他'
    return intent


async def get_user_intent_db_agent(question: str, history: List[str]) -> str:
    intent_template = DB_AGENT_INTENT_PROMPT.format(question=question, history=history)
    messages = [{
        "role": "system",
        "content": intent_template
    },
        {
            "role": "user",
            "content": question
        }]
    llm_response = await openai_client.chat.completions.create(messages=messages,
                                                               model=MODEL_NAME,
                                                               stream=False, temperature=0.2)
    try:
        resp_json = llm_response.choices[0].message.content.strip()
        intent = eval(resp_json).get('type', '数据库')
    except Exception as e:
        logger.error(e)
        intent = '数据库'
    return intent


'''提取命名实体'''


async def get_question_classification_from_question(question: str, api_name: str):
    category_schema = generate_question_classification_example(api_name)

    # print('category_schema:',category_schema)
    # category = get_api_filed_zh(api_name)
    # category_example = get_field_zh_example(api_name)
    prompt_example = generate_api_example(api_name)

    # system_template = QUESTION_CLASSIFICATION_PROMPT.format(date=now_year.strftime('%Y年%m月%d日'),category=category) + '\t示例: ' + prompt_example
    system_template = QUESTION_CLASSIFICATION_PROMPT.format(date=now_year.strftime('%Y年%m月%d日'),
                                                            category_schema=category_schema) + '\t示例: ' + prompt_example
    messages = [{"role": "system",
                 "content": system_template},
                {"role": "user", "content": question}]

    # '''抽取项目时间和经费'''
    # time_and_type_schema = generate_question_time_and_type_example(api_name)
    # time_and_type_template = QUESTION_CLASSIFICATION_PROMPT.format(date=now_year.strftime('%Y年%m月%d日'),category_schema=time_and_type_schema) + '\t示例: ' + prompt_example
    # time_and_type_messages = [{"role": "system",
    #              "content": time_and_type_template},
    #             {"role": "user", "content": question}]
    # try:
    #     time_and_type_subscription = await openai_client.chat.completions.create(messages=time_and_type_messages,
    #                                                                # model=os.environ['MODEL_NAME'],
    #                                                                model=MODEL_NAME,
    #                                                                stream=False, max_tokens=2048, temperature=0)
    #     # print(subscription)
    #     time_and_type_subscription_ = time_and_type_subscription.choices[0].message.content.strip().replace("```","").replace("json","").replace("\n","")
    #     time_and_type_subscription__ = eval(time_and_type_subscription_)
    # except Exception as e:
    #     logger.error(e)
    #     time_and_type_subscription__ = None

    try:
        subscription = await openai_client.chat.completions.create(messages=messages,
                                                                   model=MODEL_NAME,
                                                                   stream=False, max_tokens=4080, temperature=0)
        print("年度报告大模型返回结果:", subscription)
        subscription_ = subscription.choices[0].message.content.strip().replace("```", "").replace("json", "").replace(
            "\n", "")
        subscription_ = eval(subscription_)
    except Exception as e:
        logger.error(e)
        print(e)
        traceback.print_exc()
        subscription_ = []
    print('抽取结果：', str(subscription_))
    return subscription_


def get_data_from_database(_dialogue: Dialogue, api_name, sql_query, ret_data, index=None):
    try:
        result = []
        value = 0
        print('api_name:', api_name)
        api_url = get_api_url(api_name)
        print('api_url:', api_url)
        print('请求接口:', sql_query)
        # print('TOKEN:', TOKEN)
        resp = requests.post(api_url, json=sql_query, cookies={'access_token': TOKEN}).json()
        print('接口返回状态:', resp)
        for data in resp['data'][0]['data']:
            if index:
                result.append({'name': data['name'], 'value': data['value']})
            else:
                result.append({'name': str(data['name']), 'value': data['value']})
            value += data['value']
        if index:
            result.append({'name': '合计', 'value': value})
        else:
            result.append({'name': '合计', 'value': value})
        index = index if index else 1
        ret_data['table'].append({
            'group': '结果' + str(index),
            'data': result
        })
    except Exception as e:
        logger.error(e)
        _dialogue.update_status('failed')
    # print('result', str(ret_data))


async def mapping_project_info(project_info_params, api_name):
    mapping_result = {'search': {}, 'group': {}, 'order': {}}
    organize_is_exist = ''
    '''字段转换'''
    api_config = get_api_field(api_name)
    for project_info_param in project_info_params:
        key, value = list(project_info_param.keys())[0], list(project_info_param.values())[0]

        try:
            field = api_config.get(key)
            if not field:
                _key, index = rerank(key, list(api_config.keys()), again_flag=True)
                print('rerank:', _key)
                field = api_config.get(_key)
                if field:
                    continue
        except Exception as e:
            print(e)
            continue
        type_len = len(field.get('type')) if isinstance(field.get('type'), list) else 1
        if key in ('分组', '结果形式'):
            _type = 'group' if key == '分组' else 'order'
            _value, max_index = rerank(value, api_config[key]['example'], again_flag=True)
            print('key: {key}, _value:{_value}'.format(key=key, _value=_value))
            if _value:
                name = api_config[key]['name'][max_index]
                # mapping_result[name] = _value
                mapping_result[_type].update({name: _value})
            elif key == '分组':
                _value, max_index = rerank("按时间统计", api_config[key]['example'], again_flag=True)
                if _value:
                    name = api_config[key]['name'][max_index]
                    # mapping_result[name] = _value
                    mapping_result[_type].update({name: _value})
        elif type_len > 1:
            '''类型'''
            mapping_result['search'].update({field['name']: await get_type_mapping(value, field['type'])})
        elif field and not field.get('is_time'):
            _value = transfrom_alias2name(value, api_name)
            '''处理一下BQ中心特殊索引'''
            index_type = field.get('index_type')
            if index_type:
                _index = get_index_by_bq_api(_value, index_type)
                # print('bq api index',_index)
                if _index:
                    mapping_result['search'].update({field['name']: _index})
                # else:
                #    organize_is_exist = _value
            else:
                name2index = get_api_data_name2index(api_name)
                if isinstance(_value, list):
                    for __value in _value:
                        _index = name2index.get(__value)
                        mapping_result['search'].update({field['name']: _index})
                elif isinstance(_value, str):
                    _index = name2index.get(_value)
                    mapping_result['search'].update({field['name']: _index})
                else:
                    print('抽取结果类型错误')

        elif field and field.get('is_time'):
            '''时间字段，需要转换value'''
            mapping_result['search'].update({field['name']: get_project_date_mapping(value)})

    return mapping_result, organize_is_exist


def process_qa_history(history, query, question):
    messages = []
    max_tokens = MAX_TOKENS - get_length(question) - MAX_NEW_TOKENS

    _history = []
    _length = 0
    history_zero_flag = False
    for idx in list(range(len(history)))[::-1]:
        _length += get_length(history[idx])
        _history.append(history[idx])
        if _length >= max_tokens and len(_history) > 0:
            _history.pop(-1)
            _length -= get_length(history[idx])
            if len(_history) % 2 != 0:  # 去除system回答
                _history.pop(-1)
                _length -= get_length(history[idx])
            break  # 为什么去除一次历史就退出了？？？
        elif _length > max_tokens and len(_history) == 0:
            history_zero_flag = True
            break

    if history_zero_flag:
        history, _history_length = [], 0
    else:
        history, _history_length = _history[::-1], max(_length + 5 * len(_history), 0)  # 此处为什么要加5*len(_history)？？

    for idx in range(len(history)):
        messages.append({
            'role': 'user',
            'content': history[idx]
        })
    query['messages'] = messages
    query['max_tokens'] = max(MAX_NEW_TOKENS, max_tokens - _history_length)
    return query


def generate_sql_query(sql_query, source, query_key, operateType):
    sql_query['queryList'].append({
        "queryKey": query_key,
        "queryValue": source,
        "operateType": operateType
    })


def generate_sql_group_by(sql_query, group_by, type, ext):
    sql_query['groupRule'] = {
        "classKey": "projectInformation",
        "propertyName": group_by,
        "groupType": type,
        "typeExt": ext
    }


def generate_query(mapping_result, api_name):
    ret = {}
    api_config = get_api_config_by_name(api_name)['fields']
    ret.update({'classKey': api_config['classKey']})
    search_config = api_config.get('search')
    group_config = api_config.get('group')
    order_config = api_config.get('order')

    """先处理search"""
    for search_field in search_config:
        name = search_field['name']
        if mapping_result['search'].get(name):
            ret['queryList'] = ret['queryList'] if ret.get('queryList') else []
            if isinstance(search_field.get('type'), list):
                value, _type = mapping_result['search'][name]['compare_word'], mapping_result['search'][name]['value']
                ret['queryList'].append({
                    "queryKey": name,
                    "queryValue": value,
                    "operatorType": _type
                })
            else:
                if search_field.get('type') == 'equals':
                    ret['queryList'].append({
                        "queryKey": name,
                        "queryValue": mapping_result['search'][name][0],
                        "operateType": search_field.get('type')
                    })
                else:
                    ret['queryList'].append({
                        "queryKey": name,
                        "queryValue": mapping_result['search'][name],
                        "operateType": search_field.get('type')
                    })

    if mapping_result['group']:
        for group_field in group_config:
            name = group_field['name']
            if mapping_result['group'].get(name):
                ret['groupRule'] = ret['groupRule'] if ret.get('groupRule') else {}
                ret['groupRule'] = {
                    "classKey": group_field.get('key'),
                    "propertyName": name,
                    "groupType": group_field.get('type'),
                    "typeExt": group_field.get('extr')
                }
                break
    else:
        ret['groupRule'] = {
            "classKey": "projectInformation",
            "propertyName": "confirm_date",
            "groupType": "datetime",
            "typeExt": "{\"format\" : \"year\"}"
        }

    ret['statisticItemList'] = []
    if mapping_result['order']:
        name, desc = list(mapping_result['order'].keys())[0], list(mapping_result['order'].values())[0]
        for order_field in order_config:
            if desc == order_field['example']:
                ret['statisticItemList'].append(
                    {'classKey': order_field.get('key'), 'propertyName': order_field.get('name', ''),
                     'statisticType': order_field.get('type')})
                break
    else:
        for order_field in order_config:
            ret['statisticItemList'].append(
                {'classKey': order_field.get('key'), 'propertyName': order_field.get('name', ''),
                 'statisticType': order_field.get('type')})
            break
    return ret


async def process_qa_sql_query(_dialogue: Dialogue, question: str, api_name: str, compare_flag=False):
    """提取文本相关字段"""
    project_info_params = await get_question_classification_from_question(question, api_name=api_name)
    if not compare_flag:
        _dialogue.update_status('mapping')
    """映射项目数据"""
    mapping_result, organize_is_exist = await mapping_project_info(project_info_params, api_name=api_name)
    if not compare_flag:
        _dialogue.update_status('acquire')
    """生成查询"""
    sql_query = generate_query(mapping_result, api_name=api_name)

    print(sql_query)
    print(project_info_params)

    return sql_query, mapping_result


async def qa_stream(task_id: str, user_id: str, question: str, history: List[str], thing_pattern: bool):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    """intent recognition"""
    intent_type = await get_user_intent(question, history)
    print('intent_type:', intent_type,flush=True)
    if intent_type == '文档':
        _dialogue.update_status('extract')
        """获取关联docs"""
        relevant_documents = await search_with_same_outline(question)

        reference_source = {}
        index = 0
        for document in relevant_documents:
            document_id = document['metadata']['document_id']
            file_name = document['metadata']['file_name']
            if reference_source.get(document_id):
                reference_source[document_id]['content'] += document['page_content']
            else:
                reference_source[document_id] = {
                    'file_name': f'来源{index}:' + file_name,
                    'content': document['page_content']
                }

        document_template = SECOND_SYSTEM_TEMPLATE_.format(sources=list(reference_source.values()), text=question)[:10240]

        _dialogue.update_status('acquire')
        """获取query文档关联结果"""
        subscription = await get_docs_subscription(0.6, document_template, thing_pattern=thing_pattern)
        _dialogue.update_status('success')
        return subscription, rel_docs2refs(relevant_documents), []

    elif '数据库' in intent_type:

        if history:
            """根据history重新生成question"""
            question = await generate_question_by_history(question, history)

        print('question:', question)
        api_name = await get_api_name_by_intent(question)

        _dialogue.update_status('extract')

        """查询api获取数据"""
        sql_query = {}
        mapping_result = {}
        ret_data = {'table': [], 'table_name': "数据统计表"}
        if intent_type == "数据库对比":
            questions = await get_task_divide_from_question(question)
            for index, _question in enumerate(questions):
                sql_query, mapping_result = await process_qa_sql_query(_dialogue, _question, api_name,
                                                                       compare_flag=True)
                get_data_from_database(_dialogue, api_name, sql_query, ret_data, index + 1)

        else:
            sql_query, mapping_result = await process_qa_sql_query(_dialogue, question, api_name)
            get_data_from_database(_dialogue, api_name, sql_query, ret_data)

        # print('返回结果:', ret_data)
        # ret_data['sql_query'] = sql_query   # 回传sql_query给前端
        """获取可能会问"""
        category_schema = get_category_schema(api_name)
        ret_data['maybe_questions'] = await get_maybe_questions(question, api_name, category_schema)

        llm_resp = await get_sql_llm_response(question, ret_data, mapping_result, thing_pattern=thing_pattern,
                                              is_text=False)
        print('ret_data:', ret_data)
        _dialogue.update_status('success')
        return llm_resp, [], ret_data

    elif intent_type == '生成报告':
        _dialogue.update_status('extract')
        generate_data, data, year = await generate_report(task_id, "project_api", question, history)
        total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
        _dialogue.update_status('success')
        return total_llm_resp, [], data

    else:
        """other意图"""
        if history:
            """根据history重新生成question"""
            question = await generate_question_by_history(question, history)
        total_llm_resp = await get_llm_other_response(question, thing_pattern)
        _dialogue.update_status('success')
        return total_llm_resp, [], {}


# 单段最大字符长度（≈ 等效于 8000 tokens，适配 GPT-4-32k）
MAX_CHARS_PER_REQUEST = 8000


# 🔹 分段函数：按字符长度粗略分段
def split_text_by_length(text: str, max_length: int = MAX_CHARS_PER_REQUEST) -> List[str]:
    return [text[i:i + max_length] for i in range(0, len(text), max_length)]


# 🔹 调用 OpenAI 处理单段文本
async def call_llm_on_chunk(chunk_text: str, thing_pattern: bool):
    messages = [{
        "role": "system",
        "content": "请读取text后的文本，并处理成人能识别的markdown的形式方便人观看：" + chunk_text
    }]

    if thing_pattern:
        response = await openai_client2.chat.completions.create(
            messages=messages, model=MODEL_NAME2,
            stream=False, temperature=0.3
        )
    else:
        response = await openai_client.chat.completions.create(
            messages=messages, model=MODEL_NAME,
            stream=False, temperature=0.3
        )
    return response.choices[0].message.content


#🔹 主处理函数：OCR 文本分段 + 并发调用 LLM  如果需要进行大文本传输就打开这个，可以处理   但是controller  返回的时候会有问题
# async def generate_ocr_result(data_info: str, thing_pattern=False):
#     chunks = split_text_by_length(data_info, MAX_CHARS_PER_REQUEST)
#     tasks = [call_llm_on_chunk(chunk, thing_pattern) for chunk in chunks]
#     results = await asyncio.gather(*tasks)
#     return "\n\n".join(results)


# 🔹 主入口函数：处理文件 OCR 上传、调用大模型
async def qa_ocr_stream(
    task_id: str,
    user_id: str,
    question: str,
    history: List[str],
    thing_pattern: bool,
    file: UploadFile
):
    print("来到 service", flush=True)
    url = OCR_URL

    file_bytes = file.file.read()
    files = {
        "file": (file.filename, file_bytes, file.content_type)
    }

    # 🔸 调用本地 OCR 接口
    response = requests.post(url, files=files)
    # ocr_text = response.text
    ocr_text = response.json().get('text')
    # ocr_text = bytes(ocr_text, 'utf-8').decode('unicode_escape')
    print("接口拿回数据:", ocr_text, flush=True)

    print("接口拿回数据（字符长度）:", len(ocr_text), flush=True)

    # 🔸 并发调用大模型，处理文本
    total_llm_resp = await generate_ocr_org_result(ocr_text, thing_pattern=thing_pattern)

    print("返回数据:", total_llm_resp, flush=True)

    return total_llm_resp, [], {}



async def qa_ocr_file2txt(
    task_id: str,
    user_id: str,
    question: str,
    history: List[str],
    thing_pattern: bool,
    file: UploadFile,

):
    print("来到 准备开始调用ocr接口", flush=True)
    # url = "http://192.168.1.223:5000/ocr-chat"
    url = OCR_URL
    file_bytes = file.file.read()
    files = {
        "file": (file.filename, file_bytes, file.content_type)
    }

    # 🔸 调用本地 OCR 接口
    response = requests.post(url, files=files)
    ocr_text = response.json().get("text",str)
    return ocr_text


#将文件和 teamid  传入，输出json 可以存到数据库中，返回状态即可
async def qa_ocr_team(
    team_id: str,
    file: UploadFile,
    creator_guid: str

):
    print("来到 service", flush=True)
    records = []
    result = []
    # url = "http://192.168.1.223:5000/ocr-chat"
    url = OCR_URL
    file_bytes = file.file.read()
    files = {
        "file": (file.filename, file_bytes, file.content_type)
    }

    # 🔸 调用本地 OCR 接口
    response = requests.post(url, files=files)
    ocr_text = response.json().get("text",str)
    # ocr_text = bytes(ocr_text, 'utf-8').decode('unicode_escape')
    print("ocr_text"+ocr_text, flush=True)
    #判断问题  是否包含预览，如果包含预览，就直接调用ocr的返回值，如果不包含预览，包含团队，就调用团队prompt，否则prompt嵌入
    print("来到团队", flush=True)
    # 调用团队prompt
    total_llm_resp_all = await generate_ocr_team_result(ocr_text)
    content = total_llm_resp_all.choices[0].message.content
    cleand = content.strip()
    print("cleand", cleand,flush=True)
    pattern = r"```json\s*.*?```"
    match = re.search(pattern, cleand, re.DOTALL)
    cleand = match.group(0) if match else ""
    print("处理完以后的样式",cleand,flush=True)
    cleand = re.sub(r"```(?:json)?\n", "", cleand)
    cleand = re.sub(r"```", "", cleand)
    json_cleand = json.loads(cleand)
    for item in json_cleand:
        item.setdefault("team_guid", team_id)
        item.setdefault("creator_guid",creator_guid)
    print("团队处理后的数据", json_cleand,flush=True)
    if isinstance(json_cleand,list):
        #首先删除所有当前team_guid 对应的所有记录
        TeamInfo.delete_by(team_guid=team_id)
        items = [TeamInfo(**item) for item in json_cleand]
        records = TeamInfo.updates(items)
        result=TeamInfo.get_by(team_guid=team_id)
        # result = json.dumps([r.model_dump() for r in result],ensure_ascii=False)
        print("data from database",result)
    return result


async def qa_ocr_stream_org(
    task_id: str,
    user_id: str,
    question: str,
    history: List[str],
    thing_pattern: bool,
    file: UploadFile,

):
    print("来到 service", flush=True)
    # url = "http://192.168.1.223:5000/ocr-chat"
    url = OCR_URL
    file_bytes = file.file.read()
    files = {
        "file": (file.filename, file_bytes, file.content_type)
    }

    # 🔸 调用本地 OCR 接口
    response = requests.post(url, files=files)
    ocr_text = response.json().get("text",str)
    # ocr_text = bytes(ocr_text, 'utf-8').decode('unicode_escape')
    print("ocr_text"+ocr_text, flush=True)
    file_path_obj = {}
    total_llm_resp = ''
    #判断问题  是否包含预览，如果包含预览，就直接调用ocr的返回值，如果不包含预览，包含团队，就调用团队prompt，否则prompt嵌入
    if '预览' in question:
        # 调用OCR的返回值
        total_llm_resp = await generate_ocr_result(ocr_text, thing_pattern=thing_pattern)
    elif '团队' in question:
        print("来到团队", flush=True)
        # 调用团队prompt
        # total_llm_resp = await generate_ocr_org_result(ocr_text, thing_pattern=thing_pattern)
        total_llm_resp_all = await generate_ocr_org_result(ocr_text, thing_pattern=thing_pattern)
        content = total_llm_resp_all.choices[0].message.content
        cleand = content.strip()
        if cleand.startswith("```"):
            cleand = re.sub(r"```(?:json)?\n", "", cleand)
            cleand = re.sub(r"```", "", cleand)
            json_cleand = json.loads(cleand)
        if isinstance(json_cleand,list):
            print("是list", flush=True)
            df = pd.DataFrame(json_cleand)
            excel_static_path = pathlib.Path(__file__).parent.parent / 'static'
            xlsx_path = excel_static_path / 'xlsx'
            os.makedirs(xlsx_path, exist_ok=True)

            document_name = "团队模板" + datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + ".xlsx"
            excel_path = excel_static_path / 'xlsx'/document_name
            file_path = os.path.join(xlsx_path, document_name)
            df.to_excel(excel_path, index=False)

            file_path = file_path.split('rag/')[1] if 'rag/' in file_path else file_path
            file_path_obj = {'generate_docx': file_path}
        total_llm_resp = await generate_ocr_result(json_cleand, thing_pattern=thing_pattern)
    else:
        # prompt嵌入处理
        total_llm_resp = await generate_markdown_normal_result(ocr_text,question, thing_pattern=thing_pattern)
    # 🔸 并发调用大模型，处理文本
    # total_llm_resp = await generate_ocr_org_result(ocr_text, thing_pattern=thing_pattern)


    print("返回数据:", total_llm_resp, flush=True)

    return total_llm_resp, [], file_path_obj






# async def qa_ocr_stream( task_id: str, user_id: str, question: str, history: List[str],thing_pattern: bool,file: UploadFile):
#
#     #首先拿到通过调用接口拿到返回值，然后将返回值作为入参调用当前的接口
#     print("来到 service")
#     url = "http://192.168.1.223:5000/ocr-chat"
#
#     file_bytes = file.file.read()
#     files = {
#         "file":(file.filename, file_bytes,file.content_type)
#     }
#     response = requests.post(url, files=files)
#     print("接口拿回数据"+response.text,flush=True)
#     total_llm_resp = await generate_ocr_result(response.text, thing_pattern=thing_pattern)
#     return total_llm_resp, [], {}
#


#
# def split_text(text, max_length=1024):
#     return [text[i:i+max_length] for i in range(0, len(text), max_length)]
#
# async def process_large_response(response_text, thing_pattern):
#     print("来到 线程池")
#     # 根据需求调整最大长度
#     text_segments = split_text(response_text)
#
#     total_llm_resp_parts = []
#     with ThreadPoolExecutor(max_workers=5) as executor:  # 设置线程池大小
#         future_to_segment = {executor.submit(generate_ocr_result, segment, thing_pattern): segment for segment in
#                              text_segments}
#
#         for future in as_completed(future_to_segment):
#             try:
#                 part_resp = future.result()
#                 total_llm_resp_parts.append(part_resp)
#             except Exception as e:
#                 print(f"Error processing segment: {e}")
#
#     return ''.join(total_llm_resp_parts), [], {}

def rel_docs2refs(rel_docs):
    ret = []
    refs = {}
    document_ids = list({
        document['metadata'].get('document_id')
        for document in rel_docs
        if document['metadata'].get('document_id')
    })
    file_id_map = {
        _file.file_id: _file.id
        for _file in File.get_by(file_id=document_ids)
        if _file.file_id
    } if document_ids else {}

    for index, document in enumerate(rel_docs):
        doc_id = document['metadata']['document_id']
        if doc_id not in refs:
            try:
                doc_name = document['metadata'].get("file_name", '')
                print('file_name:', doc_name)
                type = '.' + doc_name.split('.')[-1] if doc_name else '.docx'
            except Exception as e:
                continue
            refs[doc_id] = {
                "doc_name": doc_name if doc_name else doc_id + type,
                "doc_id": file_id_map.get(doc_id, doc_id),
                "doc_path": '/static/file/' + doc_id + type,
                "content": [document['page_content']],
            }
        else:
            refs[doc_id]['content'] += [document['page_content']]
    for key,value in refs.items():
        ret.append(value)
    return ret


async def auto_batch(prompt_tokens: int, documents: List[Dict], prefix: str = None):
    """对documents根据num_tokens进行自动化的batch"""
    batch_tokens = 16000 - prompt_tokens
    if prefix:
        batch_tokens -= get_num_tokens([prefix])[0]
    """严格小于batch_tokens"""
    batch = []
    num_tokens = 0
    idx = 0
    while idx < len(documents):
        batch.append(documents[idx])
        document_num_tokens = documents[idx]['metadata'].get("num_tokens",
                                                             get_num_tokens([documents[idx]['page_content']])[
                                                                 0])
        num_tokens += document_num_tokens
        if num_tokens >= batch_tokens:
            if len(batch) <= 1:
                """强制进行truncate"""
                logger.warning(
                    "当前batch_size==1, 但其num_tokens ({}) > {}, 进行truncate".format(num_tokens, batch_tokens))
                batch[0]['page_content'] = batch[0]['page_content'][:int(batch_tokens * 0.9)]
            else:
                num_tokens -= document_num_tokens
                batch.pop(-1)
                idx -= 1
            yield batch
            batch = []
            num_tokens = 0
        idx += 1
    if batch:
        yield batch


async def db_agent_stream(task_id, user_id, question, history, thing_pattern):
    intent_type = await get_user_intent_db_agent(question, history)

    if history:
        """根据history重新生成question"""
        question = await generate_question_by_history(question, history)

    print('question:', question)
    api_name = await get_api_name_by_intent(question)

    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, api_name=api_name)
    _dialogue.update()

    """查询api获取数据"""
    mapping_result = {}
    ret_data = {'table': [], 'table_name': "数据统计表"}

    if intent_type == '数据库对比':
        questions = await get_task_divide_from_question(question)
        for index, _question in enumerate(questions):
            sql_query, mapping_result = await process_qa_sql_query(_dialogue, _question, api_name,
                                                                   compare_flag=True)
            get_data_from_database(_dialogue, api_name, sql_query, ret_data, index + 1)
    else:
        sql_query, mapping_result = await process_qa_sql_query(_dialogue, question, api_name)
        get_data_from_database(_dialogue, api_name, sql_query, ret_data)

    """获取可能会问"""
    category_schema = get_category_schema(api_name)
    ret_data['maybe_questions'] = await get_maybe_questions(question, api_name, category_schema)

    llm_resp = await get_sql_llm_response(question, ret_data, mapping_result, thing_pattern=thing_pattern,
                                          is_text=False)
    print('ret_data:', ret_data)
    _dialogue.maybe_query = ret_data['maybe_questions']
    _dialogue.status = 'success'
    _dialogue.update()
    return llm_resp, [], ret_data


async def knowledge_agent_stream_by_document_ids(document_ids, task_id, user_id, question, history, thing_pattern):
    if history:
        """根据history重新生成question"""
        question = await generate_question_by_history(question, history)

    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    """intent recognition"""

    prompt_tokens = sum(get_num_tokens([SECOND_SYSTEM_TEMPLATE_]))
    _dialogue.update_status('extract')
    file_ids = []
    for document_id in document_ids:
        if document_id not in file_ids:
            file_ids.append(document_id)
    _filter = "document_id in {}".format(file_ids)

    print('package _filter:', _filter)

    """获取关联docs"""
    relevant_documents = await search_with_same_outline(question, _filter)

    reference_source = {}
    index = 0
    for document in relevant_documents:
        document_id = document['metadata']['document_id']
        file_name = document['metadata']['file_name']
        if reference_source.get(document_id):
            reference_source[document_id]['content'] += document['page_content']
        else:
            reference_source[document_id] = {
                'file_name': f'来源{index}:' + file_name,
                'content': document['page_content']
            }

    document_template = SECOND_SYSTEM_TEMPLATE_.format(sources=list(reference_source.values()), text=question)[:10240]

    _dialogue.update_status('acquire')
    """获取query文档关联结果"""
    subscription = await get_docs_subscription(0.6, document_template, thing_pattern=thing_pattern)
    _dialogue.update_status('success')
    return subscription, rel_docs2refs(relevant_documents), []


async def knowledge_agent_stream(package_id, task_id, user_id, question, history, thing_pattern):
    return await knowledge_agent_stream_by_document_ids(
        [_file.file_id for _file in File.get_by(package_id=package_id)],
        task_id,
        user_id,
        question,
        history,
        thing_pattern
    )


async def knowledge_file_agent_stream(ids, task_id, user_id, question, history, thing_pattern):
    return await knowledge_agent_stream_by_document_ids(
        [_file.file_id for _file in File.get_by(id=ids) if _file.file_id],
        task_id,
        user_id,
        question,
        history,
        thing_pattern
    )


async def report_agent_stream(task_id, user_id, question, thing_pattern):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    generate_data, data, year = await generate_report(task_id, "project_api", question, [''])
    total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    _dialogue.update_status('success')
    return total_llm_resp, [], data

async def week_agent_stream(task_id, user_id, question, thing_pattern):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    generate_data, data, year = await generate_week_report(task_id, "project_api", question, [''])
    total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    _dialogue.update_status('success')
    return total_llm_resp, [], data


async def template_agent_stream(task_id, user_id, question, thing_pattern,config_id):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    generate_data, data, year = await generate_template_report(task_id, "project_api", question, [''],config_id)
    generate_data_result = await generate_ocr_result(generate_data)
    # total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    _dialogue.update_status('success')
    # return total_llm_resp, [], data
    return generate_data_result, [], data


async def jira_week_agent_stream(task_id, user_id, question, thing_pattern,config_id):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    generate_data, data, year = await generate_template_report(task_id, "project_api", question, [''],config_id)
    generate_data_result = await generate_ocr_result(generate_data)
    # total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    _dialogue.update_status('success')
    # return total_llm_resp, [], data
    return generate_data_result, [], data



#需求管理
async def request_agent_stream(task_id, user_id, question, thing_pattern,document_id ,token):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    await generate_request_report(task_id,question,document_id ,token)
    # generate_data_result = await generate_ocr_result(generate_data)
    # # total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    # _dialogue.update_status('success')
    # # return total_llm_resp, [], data
    # return generate_data_result, [], data


async def mermaid_agent_stream(task_id, user_id, question, thing_pattern):
    _dialogue = Dialogue(query=question, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    generate_data_result = await generate_mermaid_report(question,thing_pattern)

    # generate_data_result = await generate_ocr_result(generate_data)
    # # total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    # _dialogue.update_status('success')
    # # return total_llm_resp, [], data
    return generate_data_result, [], {}


#合同管理
async def contract_agent_stream(task_id, user_id, query, history, thing_pattern,fileinput):
    _dialogue = Dialogue(query=query, user_id=user_id, id=task_id, thinking_pattern=thing_pattern, status='create')
    _dialogue.update()
    _dialogue.update_status('extract')
    json_result  = await generate_contract_report(task_id,query,fileinput)
    # generate_data_result = await generate_ocr_result(generate_data)
    # # total_llm_resp = await generate_question_analysis(generate_data, year, thing_pattern=thing_pattern)
    # _dialogue.update_status('success')
    # # return total_llm_resp, [], data
    return json_result





def generate_report_data(project_info_params):

    group_rule = {
        "classKey": "projectInformation",
        "propertyName": "confirm_date",
        "groupType": "datetime",
        "typeExt": "{\"format\" : \"year\"}"
    }

    sql_query = {"classKey": "projectInformation", "queryList": [], "statisticItemList": [{
        "classKey": "projectInformation",
        "statisticType": "count"
    }]}
    # if project_info_params.get('project_date'):
    #    generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    """获取正在进行的项目数量"""
    print("获取正在进行的项目数量")
    status_processing_sql = copy.deepcopy(sql_query)
    generate_sql_query(status_processing_sql, '1', 'project_status', 'equals')
    processing_count = get_processing_count(status_processing_sql)
    """获取型号项目数量"""
    print("获取型号项目数量")
    type_model_sql = copy.deepcopy(sql_query)
    generate_sql_query(type_model_sql, ["10"], 'project_type', 'in')
    model_count = get_processing_count(type_model_sql)
    """获取涉火项目数量"""
    print("获取涉火项目数量")
    is_fire_sql = copy.deepcopy(sql_query)
    generate_sql_query(is_fire_sql, "1", 'is_fire', 'equals')
    is_fire_count = get_processing_count(is_fire_sql)
    """获取重点跟踪数量"""
    print("获取重点跟踪数量")
    importance_sql = copy.deepcopy(sql_query)
    generate_sql_query(importance_sql, ['ZDGC'], 'project_flag', 'in')
    importance_count = get_processing_count(importance_sql)
    """各个项目来源的数量"""
    print("各个项目来源的数量")
    projects_sql = copy.deepcopy(sql_query)
    generate_sql_group_by(sql_query, "project_source", 'syscategory', "{\"categoryKey\" : \"GP.ProjectFrom\"}")
    projects_count = get_processing_count(projects_sql, False)
    """新增项目"""
    print("新增项目")
    new_project_sql = copy.deepcopy(sql_query)
    if project_info_params.get('project_date'):
        generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    else:
        generate_sql_query(sql_query, ['2024-01-01 00:00:00', '2024-12-31 23:59:59'], 'create_time', 'between')
    new_project_count = get_processing_count(new_project_sql)
    """拖期项目"""
    print("拖期项目")
    delay_project_sql = copy.deepcopy(sql_query)
    generate_sql_query(delay_project_sql, ['TQXM'], 'project_flag', 'in')
    delay_project_count = get_processing_count(delay_project_sql)

    """责任单位任务承担 dept - > count"""
    print("责任单位任务承担")
    admin_dept_sql = copy.deepcopy(sql_query)
    generate_sql_query(admin_dept_sql, 'de83', 'charge_org', 'in')
    admin_dept_count = get_processing_count(admin_dept_sql, False)

    """重点项目执行情况"""

    """项目数量情况"""
    print("项目数量情况")
    projects_source_count_sql = copy.deepcopy(sql_query)
    projects_source_count_sql['groupRule'] = {
        "classKey": "projectInformation",
        "propertyName": "project_source",
        "groupType": "syscategory",
        "typeExt": "{\"categoryKey\" : \"GP.ProjectFrom\"}"
    }
    projects_source_count = get_processing_count(projects_source_count_sql, False)

    """项目经费执行情况"""
    print("项目经费执行情况")
    projects_source_amount_sql = copy.deepcopy(sql_query)
    projects_source_amount_sql['groupRule'] = {
        "classKey": "projectInformation",
        "propertyName": "project_source",
        "groupType": "syscategory",
        "typeExt": "{\"categoryKey\" : \"GP.ProjectFrom\"}"
    }
    projects_source_amount_sql['statisticItemList'] = [
        {
            "classKey": "projectInformation",
            "propertyName": "total_amount",
            "statisticType": "sum"
        }]
    projects_source_amount = get_processing_count(projects_source_amount_sql, False)
    data = {
        'processing_count': processing_count,
        'model_count': model_count,
        'is_fire_count': is_fire_count,
        'importance_count': importance_count,
        'projects_count': projects_count,  # project -> count
        'new_project_count': new_project_count,
        'delay_project_count': delay_project_count,
        'admin_dept_count': admin_dept_count,
        'projects_source_count': projects_source_count,
        'projects_source_amount': projects_source_amount
    }
    return data


def get_week_nomal_processing_count(sql_query: Dict, total_status: bool = True) -> Dict:

    # resp = requests.post("http://192.168.1.172/gateway/ebpAppSomRuntime/api/som-mgr-objects/query", json=sql_query, cookies={'access_token': TOKEN}).json()
    resp = requests.post(BQ_DATA_DESC_URL, json=sql_query, cookies={'access_token': TOKEN}).json()

    return resp



#生成周报的方法
def generate_week_process_report_data():



    sql_query = {"viewKey": "prjMonthlyManage","page":1,"pageSize":15,"query":{ "is_research_prj":"1","project_type":[],"project_manage_type":[],"project_source":[],
    "project_status":["0","1","8"],"project_level":["4","5"],"project_data_status":"approved", "project_research_type":"overall","sort":[]

    }}
    # if project_info_params.get('project_date'):
    #    generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    print("拿到周报")
    status_processing_sql = copy.deepcopy(sql_query)
    data = get_week_processing_count(status_processing_sql)

    return data

#生成周报的方法
def generate_week_problem_report_data():

    sql_query = {"viewKey": "project_problem_list","page":1,"pageSize":15,"query":{ "problem_level":"","problem_source":"projectProblem","project_stage":"","project_status":"",
    "problem_type": "","problem_yuantou":"projectProblem"
    }}
    # if project_info_params.get('project_date'):
    #    generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    print("拿到问题报告")
    status_processing_sql = copy.deepcopy(sql_query)
    data = get_week_problem_processing_count(status_processing_sql)

    return data


#生成风险报告的方法
def generate_week_risk_report_data():

    sql_query = {"viewKey": "project_risk_manage_list","page":1,"pageSize":15,"query":{ "access_bool":"","risk_level":"","risk_num":None,"risk_possibility":[],
    "risk_seriously": "","risk_source":"projectRisk","risk_status":[],"risk_type":"","risk_yuantou":"projectRisk","solve_bool":""},"sort":[]}
    # if project_info_params.get('project_date'):
    #    generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    print("拿到风险报告")
    status_processing_sql = copy.deepcopy(sql_query)
    data = get_week_risk_processing_count(status_processing_sql)

    return data

#生成延期项目报告的方法
def generate_week_delay_report_data():

    sql_query = {"viewKey": "prjDelayManage","page":1,"pageSize":15,"query":{"project_source": [],"project_type": [],"project_status": [],"project_level": [],
    "project_data_status": ["approved"],"projectLabel_code": "TQXM","is_research_prj": "1"},"sort":[]}
    # if project_info_params.get('project_date'):
    #    generate_sql_query(sql_query, project_info_params['project_date'], 'confirm_date', 'between')
    print("拿到项目拖期报告")
    status_processing_sql = copy.deepcopy(sql_query)
    data = get_week_delay_processing_count(status_processing_sql)

    return data




def generate_bar_chart(document, year, bar_data, title):
    bar = Bar()
    bar.add_yaxis(str(year), list(bar_data.values()), category_gap="50%")
    bar.add_xaxis(list(bar_data.keys()))

    # snapshot.PHANTOMJS_EXEC = r"rag\phantomjs-2.1.1-windows\bin\phantomjs.exe"
    snapshot.PHANTOMJS_EXEC = r"phantomjs"
    # snapshot.PHANTOMJS_EXEC = r"src\rag\phantomjs-2.1.1-linux-x86_64\bin\phantomjs"
    # snapshot.PHANTOMJS_EXEC = pathlib.Path(__file__).parent.parent / 'phantomjs-2.1.1-linux-x86_64' / 'bin' / 'phantomjs'
    # snapshot.PHANTOMJS_EXEC = pathlib.Path(__file__).parent.parent / 'static' / 'phantomjs-2.1.1-windows' / 'bin' / 'phantomjs.exe'
    make_snapshot(snapshot, bar.render(), "bar_chart.png")
    paragraph = document.add_paragraph()
    run = paragraph.add_run('{}柱状图'.format(title))
    run.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_picture('bar_chart.png', width=Cm(10), height=Cm(5))

    # 设置图片居中
    picture_paragraph = document.paragraphs[-1]
    picture_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def get_random_md5():
    random_byts = os.urandom(16)
    md5 = hashlib.md5()
    md5.update(random_byts)
    return md5.hexdigest()


def _add_heading(document, _text, _level, is_center=False):
    title = document.add_heading('', level=_level)
    if is_center:
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_before = Pt(0.5)  # 设置段前 0.5 磅
    title.paragraph_format.space_after = Pt(0.5)  # 设置段后 0.5 磅
    title.paragraph_format.line_spacing = 1.5  # 设置行间距为 1.5
    run = title.add_run(_text)
    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置标题颜色为黑色
    run.font.name = u'宋体'
    run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')


def _add_llm_response_to_document(document, llm_resp):
    is_first_line = True
    for line in llm_resp.split('\n'):
        if line:
            if is_first_line:
                _paragraph = document.add_paragraph(line)
                _paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                _paragraph.paragraph_format.first_line_indent = Pt(20)
            else:
                document.add_paragraph(line).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def generate_report_home_page(doc):
    # ==================== 设置页面格式 ====================
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ==================== 添加页眉 ====================
    header = section.header
    header_para = header.paragraphs[0]
    header_run = header_para.add_run("中国兵器工业信息中心")
    header_run.font.name = "黑体"
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # ==================== 首页主体内容 ====================
    # 添加垂直间距
    doc.add_paragraph().paragraph_format.space_after = Pt(88)

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("项目数据分析简报")
    title_run.font.name = u"宋体"
    title_run.font.size = Pt(34)
    title_run.font.bold = True
    title_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')

    # 添加装饰线
    line = doc.add_paragraph()
    line.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    line_run = line.add_run("―" * 36)  # 使用长划线装饰
    line_run.font.color.rgb = RGBColor(79, 129, 189)
    line.paragraph_format.space_before = Pt(4)
    line.paragraph_format.space_after = Pt(24)

    doc.add_paragraph().paragraph_format.space_after = Pt(298)
    # 机构名称
    org = doc.add_paragraph()
    org.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    org_run = org.add_run("中国兵器工业信息中心")
    org_run.font.name = u"宋体"
    org_run.font.size = Pt(12)
    org_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')
    # org.paragraph_format.space_after = Pt(36)

    # 日期
    # date_str = datetime.datetime.today().strftime("%Y年%m月%d日")
    date_str = datetime.today().strftime("%Y年%m月%d日")
    date = doc.add_paragraph()
    date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_run = date.add_run(date_str)
    date_run.font.name = u"宋体"
    date_run.font.size = Pt(12)
    date_run._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')
    doc.add_paragraph().paragraph_format.space_after = Pt(158)

    # ==================== 添加页脚 ====================
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_run = footer_para.add_run("资料 · 注意保管")
    footer_run.font.name = "宋体"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(192, 0, 0)
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


async def generate_document(question, generate_data, year):
    # totol_desc_resp = await get_llm_response("# Option 分析某单位年度数据并生成总体概况，采用系统性、结构化的方法，重点突出关键指标和业务价值，同时规避数据缺失的影响。 # Rules 1. 只对存在的数据进行分析，不存在的数据忽略2. 生成支持word的文本格式，禁止使用markdown",str(generate_data['total_description']),False)
    totol_desc_resp = await generate_title_analysis("总体概况", str(generate_data['total_description']), year)
    document = Document()

    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')

    '''报告首页'''
    generate_report_home_page(document)

    # report_picture_path = pathlib.Path(__file__).parent.parent / 'static' / 'chart.png'
    # document.add_picture(str(report_picture_path),width=Cm(15.2), height=Cm(22.9))

    _add_heading(document, '总体概况', 1)
    _add_llm_response_to_document(document, totol_desc_resp)

    # transform_markdown2docx(document, totol_desc_resp)

    _add_heading(document, '项目来源情况', 1)

    '''生成摘要'''
    # projects_source_abstract = await get_llm_response('分析项目来源情况',generate_data['project_sources'],False)
    projects_source_abstract = await generate_title_analysis("项目来源情况", str(generate_data['project_sources']),
                                                             year)
    # projects_source_run.add_text(projects_source_abstract)
    _add_llm_response_to_document(document, projects_source_abstract)
    generate_bar_chart(document, year, generate_data['project_sources'], '项目来源')

    _add_heading(document, '责任单位任务承担情况', 1)

    '''生成摘要'''
    # admin_dept_abstract = await get_llm_response('总结责任单位任务承担情况',generate_data['admin_dept'],False)
    admin_dept_abstract = await generate_title_analysis("责任单位任务承担情况", str(generate_data['admin_dept']), year)
    # admin_dept_run.add_text(admin_dept_abstract)
    # transform_markdown2docx(document, admin_dept_abstract)
    _add_llm_response_to_document(document, admin_dept_abstract)
    generate_bar_chart(document, year, generate_data['admin_dept'], '责任单位任务承担情况')

    '''项目经费情况'''
    _add_heading(document, '项目经费执行情况', 1)

    '''项目经费只要生成'''
    amount_abstract = await generate_title_analysis('项目经费执行情况', generate_data['project_amount'], year)
    _add_llm_response_to_document(document, amount_abstract)
    # transform_markdown2docx(document, amount_abstract)

    '''总结'''
    _add_heading(document, '项目报告总结', 1)
    summary_abstract = await generate_title_analysis('项目报告总结', generate_data, year)
    _add_llm_response_to_document(document, summary_abstract)

    md5 = get_random_md5()

    # doc_static_path = r'rag/static/docx'
    doc_static_path = pathlib.Path(__file__).parent.parent / 'static' / 'docx'
    if not os.path.exists(doc_static_path):
        os.mkdir(doc_static_path)
    file_path = os.path.join(doc_static_path, str(year) + '年度报告' + '.docx')

    document.save(file_path)

    file_path = file_path.split('rag/')[1]

    return {'generate_docx': file_path}





async def generate_week_document(question, generate_data, year):
    # totol_desc_resp = await get_llm_response("# Option 分析某单位年度数据并生成总体概况，采用系统性、结构化的方法，重点突出关键指标和业务价值，同时规避数据缺失的影响。 # Rules 1. 只对存在的数据进行分析，不存在的数据忽略2. 生成支持word的文本格式，禁止使用markdown",str(generate_data['total_description']),False)
    totol_desc_resp = await generate_title_analysis("总体概况", str(generate_data), year)
    document = Document()

    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')

    '''报告首页'''
    generate_report_home_page(document)

    # report_picture_path = pathlib.Path(__file__).parent.parent / 'static' / 'chart.png'
    # document.add_picture(str(report_picture_path),width=Cm(15.2), height=Cm(22.9))

    _add_heading(document, '总体概况', 1)
    _add_llm_response_to_document(document, totol_desc_resp)

    # transform_markdown2docx(document, totol_desc_resp)


    '''总结'''
    _add_heading(document, '项目报告总结', 1)
    summary_abstract = await generate_title_analysis('项目报告总结', generate_data, year)
    _add_llm_response_to_document(document, summary_abstract)

    md5 = get_random_md5()

    # doc_static_path = r'rag/static/docx'
    doc_static_path = pathlib.Path(__file__).parent.parent / 'static' / 'docx'
    if not os.path.exists(doc_static_path):
        os.mkdir(doc_static_path)
    file_path = os.path.join(doc_static_path, str(year) + '年度报告' + '.docx')

    document.save(file_path)

    file_path = file_path.split('rag/')[1]

    return {'generate_docx': file_path}


async def generate_delay_document(question, generate_data, year):
    # totol_desc_resp = await get_llm_response("# Option 分析某单位年度数据并生成总体概况，采用系统性、结构化的方法，重点突出关键指标和业务价值，同时规避数据缺失的影响。 # Rules 1. 只对存在的数据进行分析，不存在的数据忽略2. 生成支持word的文本格式，禁止使用markdown",str(generate_data['total_description']),False)
    totol_desc_resp = await generate_delay_analysis("项目拖期分析报告", str(generate_data), year)

    return {'延期分析': totol_desc_resp}



# async def generate_template_document(all_list):
#
#
#     document = Document()
#
#     document.styles['Normal'].font.name = 'Times New Roman'
#     document.styles['Normal']._element.rPr.rFonts.set(oxml.ns.qn('w:eastAsia'), u'宋体')
#
#     '''报告首页'''
#     generate_report_home_page(document)
#
#
#     for item in all_list:
#
#
#     # report_picture_path = pathlib.Path(__file__).parent.parent / 'static' / 'chart.png'
#     # document.add_picture(str(report_picture_path),width=Cm(15.2), height=Cm(22.9))
#
#     _add_heading(document, '项目拖期分析', 1)
#     _add_llm_response_to_document(document, item)
#
#     # transform_markdown2docx(document, totol_desc_resp)
#
#
#     '''总结'''
#     _add_heading(document, '项目报告总结', 1)
#     summary_abstract = await generate_title_analysis('项目报告总结', all_list, year)
#     _add_llm_response_to_document(document, summary_abstract)
#
#     md5 = get_random_md5()
#
#     # doc_static_path = r'rag/static/docx'
#     doc_static_path = pathlib.Path(__file__).parent.parent / 'static' / 'docx'
#     if not os.path.exists(doc_static_path):
#         os.mkdir(doc_static_path)
#     file_path = os.path.join(doc_static_path, str(year) + '年度报告' + '.docx')
#
#     document.save(file_path)
#
#     file_path = file_path.split('rag/')[1]
#
#     return {'generate_docx': file_path}


def transfor_report_data2json(report_data):
    ret = {}
    """总体概况"""
    # ret['total_description'] = ('项目总体概况：项目状态为正在执行的数量：' + report_data['processing_count']
    #                             + '，型号研制项目数量：' + report_data['model_count'] + '，涉火项目数量：' +
    #                             report_data['is_fire_count'] + '，重点跟踪数量：' + report_data['importance_count']
    #                             + '，新增项目数量：' + report_data['new_project_count'] + '，拖期项目数量：' +
    #                             report_data['delay_project_count'] + '。' + str(report_data['projects_count']))

    ret['total_description'] = {
        '项目状态为正在执行的数量': report_data['processing_count'],
        '型号研制项目数量': report_data['model_count'],
        '涉火项目数量': report_data['is_fire_count'],
        '重点跟踪数量': report_data['importance_count'],
        '新增项目数量': report_data['new_project_count'],
        '拖期项目数量': report_data['delay_project_count'],
        '项目来源情况': report_data['projects_count']
    }

    """各个项目来源的数量"""
    ret['project_sources'] = report_data['projects_count']

    '''责任单位任务承担'''
    ret['admin_dept'] = report_data['admin_dept_count']

    '''项目经费情况'''
    ret['project_amount'] = report_data['projects_source_amount']

    return ret


async def generate_report(task_id: str, api_name: str, question: str, history: List[str]):
    # if history:
    #     """根据history重新生成question"""
    #     question = await generate_question_by_history(question, history)
    #
    # """提取文本相关字段"""
    # project_info_params = await get_question_classification_from_question(question, api_name)

    add_task_status(task_id, 'mapping')

    # await mapping_project_info(project_info_params, api_name)

    # TODO： 根据question提取时间

    print('report status: extract date')
    date = await get_year_by_question(question)

    print('report date:', date)

    year = date[0].split('-')[0]
    # year = '2024'

    project_info_params = {'project_date': date}

    add_task_status(task_id, 'acquire')
    """生成查询"""
    report_data = generate_report_data(project_info_params)

    """查询结果生成数据"""
    generate_data = transfor_report_data2json(report_data)

    """data生成图表，生成document"""
    file_path = await generate_document(question, generate_data, year)

    print('document_file_path:', file_path)

    generate_data['责任单位任务承担'] = generate_data['admin_dept']
    generate_data['项目经费情况'] = generate_data['project_amount']

    # del generate_data['admin_dept']
    # del generate_data['project_amount']

    return generate_data, file_path, year




async def generate_week_report(task_id: str, api_name: str, question: str, history: List[str]):
    # if history:
    #     """根据history重新生成question"""
    #     question = await generate_question_by_history(question, history)
    #
    # """提取文本相关字段"""
    # project_info_params = await get_question_classification_from_question(question, api_name)

    add_task_status(task_id, 'mapping')

    # await mapping_project_info(project_info_params, api_name)

    # TODO： 根据question提取时间

    print('report status: extract date')
    date = await get_year_by_question(question)

    print('report date:', date)

    year = date[0].split('-')[0]
    # year = '2024'

    project_info_params = {'project_date': date}

    add_task_status(task_id, 'acquire')
    """生成查询"""
#    report_data = generate_week_process_report_data(project_info_params)

#    report_data = generate_week_problem_report_data(project_info_params)
    """风险查询"""
#    report_data = generate_week_risk_report_data(project_info_params)
    """延迟查询"""
    report_data = generate_week_delay_report_data(project_info_params)

    """查询结果生成数据"""
    # generate_data = transfor_report_data2json(report_data)


    """data生成图表，生成document"""
    file_path = await generate_week_document(question, report_data, year)

    print('document_file_path:', file_path)


    # del generate_data['admin_dept']
    # del generate_data['project_amount']

    return report_data, file_path, year

class Report:
    async def analyze(self, prompt,chart_prompt,para):
        result_init=get_week_nomal_processing_count(para)
        print("result_init:", repr(result_init), flush=True)
        json_end = await generate_normal_result(result_init, "请处理返回结果json只包含  项目名称  风险名称  风险类型  风险严重性  风险可能性  风险状态  风险说明  识别时间   风险值  风险等级  责任单位  识别人")
        print("json_end:", repr(json_end), flush=True)

        markdown_result_end = await generate_markdown_normal_result(json_end, prompt)
        print("markdown_result_end:", repr(markdown_result_end), flush=True)
        print("chart_prompt:", repr(chart_prompt), flush=True)
        chart_result_end = await generate_graph_normal_result(json_end, chart_prompt)
        chart_result_end = re.sub(r"```(?:python)?\n", "", chart_result_end)
        chart_result_end = re.sub(r"```", "", chart_result_end)
        print("chart_clean_code:"+chart_result_end, flush=True)
        return markdown_result_end,chart_result_end

class Risk:
    async def analyze(self, prompt):
        await asyncio.sleep(0.1)
        risk_data=generate_week_risk_report_data()

        get_week_delay_processing_count()
#        print("risk_data:", risk_data, flush=True)
        risk_data_result =await generate_normal_result(risk_data,prompt)
        print("risk_data_result:", repr(risk_data_result), flush=True)
        return risk_data_result


class Delay:
    async def analyze(self, prompt):
        await asyncio.sleep(0.1)
        delay_data=generate_week_delay_report_data()
        delay_data_result =await generate_normal_result(delay_data,prompt)
        return delay_data_result


##to do  delete  验证完了就删除
# 核心函数：读取 JSON，动态调用方法，汇总 Markdown 报告
# async def generate_report_from_json(json_path):
#     print('generate report from json:', json_path)
#     print(f"Current working directory: {os.getcwd()}")
#     with open(json_path, 'r', encoding='utf-8') as f:
#         config = json.load(f)
#
#     tasks = []
#     module_names = []
#
#     for module_cfg in config["modules"]:
#         class_name = module_cfg["class"]
#         method_name = module_cfg["method"]
#         prompt = module_cfg["prompt"]
#         module_display_name = module_cfg["name"]
#
#         # 反射获取类
#         cls = globals().get(class_name)
#         if cls is None:
#             raise ValueError(f"找不到类：{class_name}")
#
#         # 创建类实例
#         instance = cls()
#
#         # 反射获取方法
#         method = getattr(instance, method_name, None)
#         if method is None:
#             raise ValueError(f"类 {class_name} 中找不到方法：{method_name}")
#
#         # 加入异步任务队列
#         tasks.append(method(prompt))
#         module_names.append(module_display_name)
#
#     # 并发执行所有任务
#     results = await asyncio.gather(*tasks)
#
#     # 构建 Markdown 报告
#     markdown_report = "# 自动生成报告\n\n"
#     for name, result in zip(module_names, results):
#         markdown_report += f"## {name}\n\n{result}\n\n"
#     document = Document()
#     _add_llm_response_to_document(document,markdown_report)
#
#     doc_static_path = pathlib.Path(__file__).parent.parent / 'static' / 'docx'
#     if not os.path.exists(doc_static_path):
#         os.mkdir(doc_static_path)
#     file_path = os.path.join(doc_static_path, '报告' + '.docx')
#
#     document.save(file_path)
#
#     file_path = file_path.split('rag/')[1]
#     return markdown_report, file_path, ""
#

def html_to_docx(html, doc, img_paths=[]):
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'h4':
            doc.add_heading(element.text, level=4)
        elif element.name == 'h5':
            doc.add_heading(element.text, level=5)
        elif element.name == 'p':
            para = doc.add_paragraph()
            for item in element.children:
                if item.name == 'strong':
                    run = para.add_run(item.text)
                    run.bold = True
                elif item.name == 'em':
                    run = para.add_run(item.text)
                    run.italic = True
                else:
                    if isinstance(item, str):
                        para.add_run(item)
                    else:
                        para.add_run(item.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        else:
            if hasattr(element, 'text') and element.text:
                doc.add_paragraph(element.text)
        # 插入图片
    for img_path in img_paths:
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(5.5))

# 新增的辅助函数：md 转 docx
def markdown_to_docx(md_text, doc, img_paths=[]):
    html = markdown.markdown(md_text)
    html_to_docx(html, doc, img_paths)

def execute_chart_code(code, chart_dir, module_name):
    print("图表code:", code, flush=True)
    os.makedirs(chart_dir, exist_ok=True)
    # 给每个模块生成唯一文件名，避免覆盖
    img_filename = f"{module_name}_{uuid.uuid4().hex}.png"
    img_path = chart_dir / img_filename
    print("img_path:", img_path, flush=True)
    code = f"{code}\n\nplt.savefig('{img_path}')"
    # 提供全局变量：允许生成代码使用 img_path
    global_vars = {"__name__": "__main__", "img_path": str(img_path)}
    # print("global_vars"+global_vars)
    local_vars = {}

    try:
        exec(code, global_vars, local_vars)
        if os.path.exists(img_path):
            return str(img_path)
        else:
            print(f"[WARN] 图表代码执行完成，但未生成图片: {module_name}")
            return None
    except Exception:
        print(f"[ERROR] 执行模块 {module_name} 图表代码出错：")
        traceback.print_exc()
        return None
# -------------------------------
# 异步主流程
# -------------------------------
async def generate_report_from_json(json_str:str):
    # with open(json_path, 'r', encoding='utf-8') as f:
    #     config = json.load(f)
    print("来处理jso",json_str,flush=True)
    config = json.loads(json_str)

    tasks = []
    module_names = []
    chart_prompts = []
    paras = []


    for module_cfg in config["modules"]:
        cls = globals().get(module_cfg["class"])
        if cls is None:
            raise ValueError(f"找不到类：{module_cfg['class']}")
        instance = cls()
        method = getattr(instance, module_cfg["method"], None)
        if method is None:
            raise ValueError(f"类 {module_cfg['class']} 中找不到方法：{module_cfg['method']}")

        module_names.append(module_cfg["name"])
        chart_prompts.append(module_cfg.get("chart_prompt"))
        paras.append(copy.deepcopy(module_cfg.get("para",{})))
        tasks.append(method(module_cfg["prompt"],module_cfg.get("chart_prompt"),copy.deepcopy(module_cfg.get("para",{}))))


    results = await asyncio.gather(*tasks)

    markdown_report = "# 自动生成报告\n\n"
    doc = Document()
    doc_static_path = pathlib.Path(__file__).parent.parent / 'static'
    chart_dir = doc_static_path / 'charts'
    os.makedirs(chart_dir, exist_ok=True)

    for name, (md_text, py_code) in zip(module_names, results):
        md_section = f"## {name}\n\n{md_text}\n\n"
        img_paths = []

        if py_code:
            img_path = execute_chart_code(py_code, chart_dir, name)
            if img_path:
                md_section += f"![图表]({img_path})\n\n"
                img_paths.append(img_path)
        markdown_report += md_section
        markdown_to_docx(md_section, doc, img_paths)

    docx_path = doc_static_path / 'docx'
    os.makedirs(docx_path, exist_ok=True)
    document_name = "月报" + datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + ".docx"
    file_path = os.path.join(docx_path, document_name)

    # file_path_obj = {'generate_docx': file_path}
    doc.save(file_path)

    file_path = file_path.split('rag/')[1] if 'rag/' in file_path else file_path
    file_path_obj = {'generate_docx':file_path}

    return markdown_report, file_path_obj, ""





# 你原始的异步函数，改动部分在最后写文档时   原文档
# async def generate_report_from_json(json_path):
#     with open(json_path, 'r', encoding='utf-8') as f:
#         config = json.load(f)
#
#     tasks = []
#     module_names = []
#
#     class_name = "Report"
#     method_name = "analyze"
#
#     for module_cfg in config["modules"]:
#         # class_name = module_cfg["class"]
#         # method_name = module_cfg["method"]
#         prompt = module_cfg["prompt"]
#         chart_prompt = module_cfg["chart_prompt"]
#         module_display_name = module_cfg["name"]
#         para = module_cfg.get("para",{})
#         # para_json = json.loads(para)
#         #
#         para_process_sql = copy.deepcopy(para)
#
#         cls = globals().get(class_name)
#         # class_path = f"{cls.__module__}.{cls.__qualname__}"
#         # print("class_path:", class_path)
#         if cls is None:
#             raise ValueError(f"找不到类：{class_name}")
#
#         instance = cls()
#
#         method = getattr(instance, method_name, None)
#         if method is None:
#             raise ValueError(f"类 {class_name} 中找不到方法：{method_name}")
#
#         tasks.append(method(prompt,para_process_sql))
#         module_names.append(module_display_name)
#
#     results = await asyncio.gather(*tasks)
#
#     markdown_report = "# 自动生成报告\n\n"
#     for name, result in zip(module_names, results):
#         markdown_report += f"## {name}\n\n{result}\n\n"
#
#     document = Document()
#     # 这里调用我们的 md->docx 转换函数，传入 document
#     markdown_to_docx(markdown_report, document)
#
#     doc_static_path = pathlib.Path(__file__).parent.parent / 'static'
#     chart_dir = doc_static_path / 'charts'
#     os.makedirs(chart_dir, exist_ok=True)
#
#
#     if not os.path.exists(doc_static_path):
#         os.mkdir(doc_static_path)
#     file_path = os.path.join(doc_static_path, '报告.docx')
#
#     document.save(file_path)
#
#     # 这里你之前代码做了路径拆分，确认路径规则合理
#     file_path = file_path.split('rag/')[1] if 'rag/' in file_path else file_path
#     file_path_obj = {'generate_docx':file_path}
#
#     return markdown_report, file_path_obj, ""





async def generate_template_report(task_id: str, api_name: str, question: str, history: List[str],config_id:str):
    print("进来模板service")
    add_task_status(task_id, 'mapping')
    # save_file_dir = str(pathlib.Path(__file__).parent/"modules.json")
    # print("save_file_dir"+save_file_dir)
    #通过config_id找到相关的配置的字符串，然后把字符串传过去
    status_code, resp, detail   = await get_user_config(config_id)
    print("resp", resp, flush=True)
    json = resp.get("config_json",str)
    print("json"+json,flush=True)
    markdown_report, file_path, extra = await generate_report_from_json(json)
    print("markdown_report:", markdown_report,flush=True)
    print("file_path:", file_path,flush=True)

    # 保持返回值结构
    return markdown_report, file_path, extra


#需求管理部分
async def generate_request_report(task_id: str,  question: str ,document_id:str ,token:str):

    #首先将问题传递给大模型，然后拿到大模型进行处理拿到markdown  ，然后响应接口，返回200即是成功。调用大模型
    print("进来需求处理部分",flush=True)
    add_task_status(task_id, 'mapping')
    total_llm_resp = await generate_request_normal_result(question,"", thing_pattern=False)
    print("大模型返回值",total_llm_resp,flush=True)
    doc = Document()
    markdown_to_docx(total_llm_resp, doc, [])
    markdown_report = "# 自动生成报告\n\n"
    doc_static_path = pathlib.Path(__file__).parent.parent / 'static'
    docx_path = doc_static_path / 'docx'
    os.makedirs(docx_path, exist_ok=True)
    document_name = "需求" + datetime.now().strftime("%Y-%m-%d_%H-%M-%S") + ".docx"
    file_path = os.path.join(docx_path, document_name)
    # file_path_obj = {'generate_docx': file_path}
    doc.save(file_path)

    print("file_path",file_path,flush=True)
    #
    # if not os.path.exists(file_path):
    #     print("文件路径不存在")
    # elif not os.path.isfile(file_path):
    #     print("路径不是文件")
    # else:
    #     print("文件存在哦")
    #
    # try:
    #     with open(file_path, 'rb') as f:
    #         f.read(10)
    #     print("文件可读，可以上传")
    # except Exception as e:
    #     print("文件不可读",e)
    url = REQUEST_URL

    with open(file_path,"rb") as f:
        # f.seek(0)
        files = {
            "splitFile": (file_path.split("/")[-1], f, "application/octet-stream"),
        }
        # 需要传递的其他表单参数
        data = {
            "documentId": str(document_id).encode("utf-8"),
            "splitFlag": str(random.random()).encode("utf-8"),
            "tokenId": token+",2"
        }

        # 请求头（如果需要）
        headers = {
            "Authorization": "Bearer "+token,
            "User-agent" : "python-requests"
        }

        # 发送 POST 请求
        response = requests.post(url, headers=headers, data=data, files=files)

        # 查看响应结果
    print("Response Content:", response.text)

        # 关闭文件
    # files["splitFile"].close()
        # resp = requests.post("http://192.168.0.63:9034/sysorm/api/word/importLocal?_=1763627615807", json=sql_query, cookies={'access_token': TOKEN}).json()




    # file_path = file_path.split('rag/')[1] if 'rag/' in file_path else file_path
    # file_path_obj = {'generate_docx': file_path}


    #
    # print("markdown_report:", markdown_report,flush=True)
    # print("file_path:", file_path,flush=True)
    #
    # # 保持返回值结构
    # return markdown_report, file_path, extra
#画图助手
async def generate_mermaid_report( question: str ,thing_pattern:str ):

    ocr_result = await generate_mermaid_result(question,thing_pattern)
    return ocr_result




#合同管理
async def generate_contract_report(task_id: str,  question: str ,fileinput:UploadFile):
    ocr_text = await qa_ocr_file2txt(task_id,"","","","",fileinput)

    # print("进来合同处理部分", flush=True)
    # add_task_status(task_id, 'mapping')
    # total_llm_resp = await generate_contract_result(question, ocr_text, thing_pattern=False)
    # #处理返回来的数据
    # print("大模型返回值", total_llm_resp, flush=True)
    #
    # return total_llm_resp,[],{}

    print("进来合同处理部分", flush=True)
    add_task_status(task_id, 'mapping')
    ocr_result = await generate_contract_result(question, ocr_text, thing_pattern=False)
    content = ocr_result.choices[0].message.content
    cleand_content= content.strip()
    if cleand_content.startswith("```"):
        cleand_content = re.sub(r"```(?:json)?\n", "", cleand_content)
        cleand_content = re.sub(r"```", "", cleand_content)
        cleand_content = json.loads(cleand_content)
    inputData = cleand_content
    print("ocr返回值如下", inputData, flush=True)
    # 处理返回来的数据,要有源文件返回来的数据，还要有对应的数据（map对应值，比如合同名称对应name，密级对应secret,有一个专门的文件，然后还有一个要输出的文件类型json）
    fieldMapping={
        "合同名称":"name",
        "合同密级":"scret",
        "合同编号":"contract_inside_code",
        "委托方法定代表人": "party_a_legal_person",
        "承研方法定代表人": "party_b_legal_person",
        "委托单位": "party_a_org_guid",
        "承研单位": "party_b_org_guid",
        "签订日期": "sign_date",
        "生效日期": "start_date",
        "终止日期": "end_date",
        "总金额":"total_amount_new",
        "委托方投资金额":"invest_amount",
        "自筹金额": "rise_amount",
        "外协金额":"waixiejingfei"

    }

    # inputData={
    #     "合同名称":"采购合同",
    #     "密级":"机密",
    #     "合同编号":"1111",
    #     "金额":22222
    # }
    result = await generate_contract_end_json_result(inputData, fieldMapping, thing_pattern=False)
    content = result.choices[0].message.content
    cleand = content.strip()
    if cleand.startswith("```"):
        cleand = re.sub(r"```(?:json)?\n", "", cleand)
        cleand = re.sub(r"```", "", cleand)
        json_cleand = json.loads(cleand)
    ##这里要添加一种信息，需要调用150的接口，拿到和单位名称完全匹配的那条记录，然后替换这个对象中
    #首先拿到委托单位的名称  然后响应接口，拿到返回值，替换json值中的value值

    print("初步处理完的json"+cleand,flush=True)
    party_a_org_name = json_cleand.get("party_a_org_guid")
    party_b_org_name = json_cleand.get("party_b_org_guid")
    #
    # party_a_org_name = "北方石油国际有限公司"
    # party_b_org_name = "北方石油国际有限公司"

    BQ_FIND_ALL_BY_ORGTYPE = "http://192.168.1.150/gateway/ebpAppSys/api/sys-orgs/findAllByOrgType"
    headers = {
        "Authorization": "Bearer eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9.eyJwcmluY2lwYWwiOiJ7XCJhY2NvdW50Tm9uRXhwaXJlZFwiOnRydWUsXCJhY2NvdW50Tm9uTG9ja2VkXCI6dHJ1ZSxcImF1dGhvcml0aWVzXCI6W10sXCJhdXRob3JpdHlMaXN0XCI6W10sXCJjcmVkZW50aWFsc05vbkV4cGlyZWRcIjp0cnVlLFwiZW5hYmxlZFwiOnRydWUsXCJndWlkXCI6XCI3NWUzYjE3Mi0zOWRhLTQ1MzctOTU0NC05NTExOGQ0ZjdjZDJcIixcInBhc3N3b3JkXCI6XCIkMmEkMTAkUU9jNW5FMU1SbnhmbTV3YzQuLlZnZUFkT1lLMnZabGJ0TkJLcEZxbFZDLi85bUNXcHFPNzZcIixcInVzZXJuYW1lXCI6XCJzeXNjb25maWdcIn0iLCJ1c2VyX25hbWUiOiJzeXNjb25maWciLCJzY29wZSI6WyJvcGVuaWQiXSwiZXhwIjoxNzY2NDcwOTcwLCJpYXQiOjE3NjU4NjYxNzAsImp0aSI6IjQ5ZGEzZWM1LWNkOTAtNDZlMS1iOTA0LWE5MDM1ZmFjMDQxMiIsImNsaWVudF9pZCI6IndlYl9hcHAifQ.a7P6-DyPCd32E52_sX4XQnsf2cWwSxbuUXqslt2W-qdwwr14uZGW1L5ofcARIjac4UKHZwi6pDazS0Ycp92x6vdWfWe9U0nKuMfqCtOfSuYQa5dzNRpg0MrOgep8cdS6KMB0IexNOFo39DWpozcOFNFEnJ5FnyQMD23n8pCZ70Zm-wePIkHxGOTZQ6BZ4IU8yw33YVpFetGXo4myGGmXDlGppJcIsJKYcnLcY7XPjwfarc70TnR1Unys7RjBcDGCElA_gjU7natsEieyNS9GVj-Vd2za3yd6n00ECHDerBZCzmzh02Jz6AT4JHZrzmFU30rfEVR5N6kO5g3Kth9g6g",
    }

    paramas = {
        "orgTypeGuid": 3,
        "query": party_a_org_name,
    }
    respa = requests.get(BQ_FIND_ALL_BY_ORGTYPE, params=paramas, headers = headers)
    dataA = respa.json().get("data")[0]

    # print("委托单位"+dataA)

    parambs = {
        "orgTypeGuid": 3,
        "query": party_b_org_name,
    }
    respb = requests.get(BQ_FIND_ALL_BY_ORGTYPE, params=parambs,  headers = headers)
    dataB = respb.json().get("data")[0]

    # print("承研单位"+dataB)


    json_cleand["party_a_org_guid"] = dataA
    json_cleand["party_b_org_guid"] = dataB

    all_money = json_cleand.get("total_amount_new")
    total_amount_new={
        "property_value_1": all_money,
        "unit_name": "万元",
        "unit_sign": "万元",
        "value_style": "equal"
    }
    json_cleand["total_amount_new"] = total_amount_new

    invest_amount_money = json_cleand.get("invest_amount")
    invest_amount = {
        "property_value_1": invest_amount_money,
        "unit_name": "万元",
        "unit_sign": "万元",
        "value_style": "equal"
    }
    json_cleand["invest_amount"] = invest_amount

    rise_amount_money = json_cleand.get("rise_amount")
    rise_amount = {
        "property_value_1": rise_amount_money,
        "unit_name": "万元",
        "unit_sign": "万元",
        "value_style": "equal"
    }
    json_cleand["rise_amount"] = rise_amount


    waixiejingfei_money = json_cleand.get("waixiejingfei")
    waixiejingfei = {
        "property_value_1": waixiejingfei_money,
        "unit_name": "万元",
        "unit_sign": "万元",
        "value_style": "equal"
    }
    json_cleand["waixiejingfei"] = waixiejingfei


    json_result = {"contractStandard":json_cleand}
    print("大模型返回值", json_result, flush=True)
    return  json_result

   # return result, [], {}





async def generate_jira_week_report(task_id: str, api_name: str, question: str, history: List[str],config_id:str):
    print("进来jira 周报模板")
    add_task_status(task_id, 'mapping')

    await generate_jira_week_report_detial()

    # 保持返回值结构
    return 1, 2, 3




async def week_task_format(jira_json_init):
    rows = jira_json_init
    sorted_data = sorted(rows, key=lambda row: row[2])
    sorted_data_end = [[i] + row for i, row in enumerate(sorted_data, start=1)]
    return sorted_data_end


async def  process_column(matrix,col_idx):
    for row in matrix:
        if col_idx < len(row):
            value = row[col_idx]
            if re.search("上期工作综述",value):
                await extract_summary(value)


async def extract_previous_summary(text: str) -> str:
    if not text:
        return  ""
        # 提取 阶段目标、上期工作综述、本期重点工作
    pattern = re.compile(
        r"\[阶段目标\](.*?)\[上期工作综述\](.*?)\[本期重点工作\](.*)",
        re.S
    )
    match = pattern.search(text)
    if match:
        return  match.group(2).strip()

    else:
        return text


async def extract_next_summary(text: str) -> str:
    if not text:
        return  ""
        # 提取 阶段目标、上期工作综述、本期重点工作
    pattern = re.compile(
        r"\[阶段目标\](.*?)\[上期工作综述\](.*?)\[本期重点工作\](.*)",
        re.S
    )
    match = pattern.search(text)
    if match:
        return  match.group(3).strip()

    else:
        return text



async def generate_jira_week_report_all():


# ---------------- 拿到当前周和下一周的日期 ----------------
    today = datetime.today()
    monday = today - timedelta(days=today.weekday())
    friday = monday + timedelta(days=4)

    week_range = f"{monday.strftime('%Y-%m-%d')}-{friday.strftime('%Y-%m-%d')}"


    monday = today - timedelta(days=today.weekday())
    next_monday = monday + timedelta(days=7)
    next_friday = next_monday + timedelta(days=4)

    next_week_range = f"{next_monday.strftime('%Y-%m-%d')}-{next_friday.strftime('%Y-%m-%d')}"

# ---------------- 拿到所有数据 ----------------
    jira_json_init = await generate_jira_week_report_detial()

    comment_report_pre =await batch_extract(jira_json_init)
    current_task = await  current_task_issue()
    comment_report_next = await batch_extract(current_task)

    print("comment_report_pre", comment_report_pre, flush=True)
    print("comment_report_next", comment_report_next, flush=True)

    # jira_json = await week_task_format(jira_json_init)
    # current_task_json = await week_task_format(current_task)


    jira_json = jira_json_init
    current_task_json =current_task
    print(current_task_json)


    # rows = jira_json_init[1:]
    # sorted_data = sorted(rows, key=lambda row: row[2])
    # sorted_data_end = [[i]+row for i,row in enumerate(sorted_data,start=1)]
    # jira_json = sorted_data_end




    # prompt = "请生成阶段目标，分条列出，使用简洁的条目形式，不超过300字,数据源如下{data_info}，最终结果类似如下 只是举例①系统使用维护。②知识产权管理模块更新③两级功能开发"
    # stage_goal = await generate_jira_normal_result_report(jira_json, prompt)
    # stage_goal_end = re.sub(r"<think>.*?</think>", "", stage_goal, flags=re.DOTALL)
    # ---------------- 调用 DeepSeek API ----------------
    # print("111"+prompt.format(data_info=jira_json))
    # print("prompt",prompt,flush=True)

    # print("总结" + stage_goal_end, flush=True)
    #这里调用冲刺面板，获取面板中的目标
    stage_goal_end=await get_jira_sprint_detail()

#     # ---------------- 生成工作综述 ----------------
    prompt = "请生成上期工作概述，请保证最终结果足够抽象，忽略具体细节和操作性的内容，只有一层（可以忽略不必要的内容），分条列出，使用简洁的条目形式，不超过300字,数据源如下{data_info}，只提取数据源中的所属阶段目标和上期工作综述字段及任务名称字段，仅产生数据源相关数据勿产生脏数据，最终结果以下列形式展示  如"+stage_goal_end
    # ---------------- 调用 DeepSeek API ----------------
    # print("111"+prompt.format(data_info=jira_json))
    # print("prompt",prompt,flush=True)
    summary_last = await generate_jira_normal_result_report(comment_report_pre, prompt)
    summary_last_end  = re.sub(r"<think>.*?</think>","",summary_last,flags=re.DOTALL)
    # print("总结"+summary_last_end, flush=True)


    prompt = "请生成本期重点工作，请保证最终结果足够抽象，忽略具体细节和操作性的内容，只有一层（可以忽略不必要的内容）分条列出，使用简洁的条目形式，不超过300字,数据源如下{data_info}，只提取数据源中的所属阶段目标和本期重点工作字段及任务名称字段，仅产生数据源相关数据勿产生脏数据，最终结果以下列形式展示 如"+stage_goal_end
    # ---------------- 调用 DeepSeek API ----------------
    # print("111"+prompt.format(data_info=jira_json))
    # print("prompt",prompt,flush=True)
    key_tasks_of_this_period = await generate_jira_normal_result_report(comment_report_next, prompt)
    key_tasks_of_this_period_end = re.sub(r"<think>.*?</think>", "", key_tasks_of_this_period, flags=re.DOTALL)
    # print("总结" + key_tasks_of_this_period, flush=True)


#
#     table_data  = re.sub(r"<think>.*?</think>","",matches,flags=re.DOTALL)
#     # table_data = re.search(table_data, matches,flags=re.DOTALL)
#
#     # pattern = r" json\s*(.*?)\s* "
#     # table_data = re.findall(pattern, matches,re.DOTALL)
#     table_data = re.sub(r"```(?:json)?\n", "", table_data)
#     table_data = re.sub(r"```", "", table_data)
#     table_data = json.loads(table_data)
    table_data = jira_json
    table_data_next = current_task_json
    print("table_data", flush=True)
    print(table_data, flush=True)
    print("table_data_next", flush=True)
    print(table_data_next, flush=True)

#     ---------------- 打开 Excel 并追加 ----------------

    # ---------------- 打开 Excel 并追加 ----------------
    wb = load_workbook("PM项目周报_STD+GP+BKY+IC+SYS_20250825-0829.xlsx")
    ws = wb["周报_SYS（周一）"]
    table_name = "pre_summary"
    #首先要删掉两个任务备份和bug
    filter_data = [row for  row in table_data if row[1] not in  ['PLTSOM-417','PLTSOM-673'] ]
    filter_data = await week_task_format(filter_data)

    #把评论里边包含关键字的东西提出来重置，拿出来这一列的东西，所有符合的都只剩下上期进展。

    for row in filter_data:
        if len(row) > 6:
            row[6] = await extract_previous_summary(row[6])
    append_table_data_with_style(ws, table_name, filter_data)

    table_name_next = "now_week_plan"


    filter_data_next = [row for row in table_data_next if row[1] not in ['PLTSOM-417', 'PLTSOM-673']]
    filter_data_next = await week_task_format(filter_data_next)

    for row in filter_data_next:
        if len(row) > 6:
            row[6] = await extract_next_summary(row[6])
    append_table_data_with_style(ws, table_name_next, filter_data_next)


    for row in ws.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                if "{{week_range}}" in cell.value:
                    cell.value = cell.value.replace("{{week_range}}", week_range)
                if "{{next_week_range}}" in cell.value:
                    cell.value = cell.value.replace("{{next_week_range}}", next_week_range)
                # if "{{summary_of_previous_work}}" in cell.value:
                #     cell.value = cell.value.replace("{{summary_of_previous_work}}", summary_last_end)
                if "{{stage_goal}}" in cell.value:
                    cell.value = cell.value.replace("{{stage_goal}}", stage_goal_end)
                # if "{{key_tasks_of_this_period}}" in cell.value:
                #     cell.value = cell.value.replace("{{key_tasks_of_this_period}}", key_tasks_of_this_period_end)





# ---------------- 运维周报 ----------------

    ws_bak = wb["周报_质量（周一）"]
    table_name_bak = "pre_bak"
    #下边是备份情况问题
    bak_data = [row for  row in table_data if row[1] == "PLTSOM-673"]
    print("bake_data_1")
    print(bak_data, flush=True)

    bak_data = await week_task_format(bak_data)
    print("bake_data")
    print(bak_data, flush=True)
    #拿到详细说明，使用大模型生成json，然后再生成二维数组，然后插入，拿到di
    bak_detial = bak_data[0][6]
    prompt = ("请将下列数据处理为json格式的二维数组，只输出结果，数据源为 {data_info}")
    bak = await generate_jira_normal_result(bak_detial, prompt)
    print("备份大模型返回", flush=True)
    print(type(bak), flush=True)
    bak = re.sub(r"```(?:json)?\n", "", bak)
    bak = re.sub(r"```", "", bak)
    bak = json.loads(bak)
    print(bak, flush=True)
    append_table_data_with_style(ws_bak, table_name_bak, bak)





# ---------------- 关于运维上期工作总结 ----------------

    table_name_wjc  = "wjc_data"
    wjc_data = [row for  row in table_data if row[2] in ["王佳程","周昊"]]
    wjc_data = await week_task_format(wjc_data)
    print(wjc_data, flush=True)
    new_wjc_data = []
    for i,row in enumerate(wjc_data,start=1):
        single_wjc_data = [i,row[1],"","",row[3],row[4],row[5],'',row[8]]
        new_wjc_data.append(single_wjc_data)
    # new_wjc_data = json.loads(new_wjc_data)
    # append_table_data_with_style(ws_bak, table_name_wjc, wjc_data)
    # print("wjc筛选")
    # print(type(wjc_data), flush=True)
    # prompt = ("请将下列数据处理为json格式的二维数组，只输出结果不带表头,但是表头为序号、名称、空一列、空一列、负责人、配合、时间段（只包含时间，不包含别的）、""、完成情况，一共九段，没有的用空字符串代替。序号从1开始，数据源为 {data_info}  的删除说明那一列")
    # wjc = await generate_jira_normal_result(wjc_data, prompt)
    #
    #
    # print("wjc大模型")
    # print(type(wjc), flush=True)
    # wjc = re.sub(r"```(?:json)?\n", "", wjc)
    # wjc = re.sub(r"```", "", wjc)
    # print("wjc去除json")
    # print(type(wjc), flush=True)
    # print("返回结果wjc" + wjc)


    # wjc = json.dumps(wjc,ensure_ascii=False)
    # print("wjc转json1")
    # print(type(wjc), flush=True)
    # wjc_result = json.loads(wjc)
    print("wjc转json2")
    print(type(new_wjc_data), flush=True)
    print(new_wjc_data,flush=True)
    append_table_data_with_style(ws_bak, table_name_wjc, new_wjc_data)

# ---------------- 关于运维本期工作进展----------------
    table_name_wjc_next = "wjc_data_next"
    wjc_data_next = [row for row in table_data_next if row[2] in ["王佳程","周昊"]]
    wjc_data_next = await week_task_format(wjc_data_next)
    print(wjc_data_next, flush=True)
    new_wjc_data_next = []
    for i, row in enumerate(wjc_data_next, start=1):
        single_wjc_data_next = [i, row[1], "", "", row[3], row[4], row[5], '', row[8]]
        new_wjc_data_next.append(single_wjc_data_next)
    print("wjc转json2111")
    print(type(new_wjc_data_next), flush=True)
    print(new_wjc_data_next, flush=True)
    append_table_data_with_style(ws_bak, table_name_wjc_next, new_wjc_data_next)

# ---------------- 上期bug总结 ----------------
    table_name_bug = "pre_bug_summary"
    # 下边是备份情况问题
    bug_data = [row for row in table_data if row[1] == "PLTSOM-417"]
    bug_data = await week_task_format(bug_data)
    print(bug_data, flush=True)
    # 拿到详细说明，使用大模型生成json，然后再生成二维数组，然后插入，拿到di
    bug_detial = bug_data[0][6]
    # prompt = ("请将下列数据处理为json格式的二维数组。只显示内容不显示表头，最左侧添加一列序号从1开始往后顺延，组别这一列的右侧加一列空行列放置空字符串，最右侧添加两列放置空字符串，只输出结果，数据源为 {data_info}")
    prompt = ("请将下列数据处理为json格式的二维数组。只显示内容不显示表头,只输出结果，数据源为 {data_info}")
    bug = await generate_jira_normal_result(bug_detial, prompt)
    bug = re.sub(r"```(?:json)?\n", "", bug)
    bug = re.sub(r"```", "", bug)
    print("返回结果bug"+bug)
    bug = json.loads(bug)

    new_bug = []
    for i, row in enumerate(bug, start=1):
        bug = [i, row[0], "", row[1], row[2], row[3], row[4], "",""]
        new_bug.append(bug)
    print(new_bug)
    append_table_data_with_style(ws_bak, table_name_bug, new_bug)


    for row in ws_bak.iter_rows():
        for cell in row:
            if cell.value and isinstance(cell.value, str):
                if "{{week_range}}" in cell.value:
                    cell.value = cell.value.replace("{{week_range}}", week_range)
                if "{{next_week_range}}" in cell.value:
                    cell.value = cell.value.replace("{{next_week_range}}", next_week_range)


    #这里是为了显示状态的颜色

    fill_dict = {
        "完成" : PatternFill("solid", fgColor="FF92D050"),
        "处理中" : PatternFill("solid", fgColor="FF00FF00"),
        "待办" : PatternFill("solid", fgColor="FFFFFF00")
    }
    for  ws in  wb.worksheets:
        for  row  in   ws.iter_rows():
            for  cell  in  row:
                text = str(cell.value).strip() if cell.value else ""
                if text in  fill_dict:
                    cell.fill   = fill_dict[text]


    doc_static_path = pathlib.Path(__file__).parent.parent / 'static'
    week_report_dir = doc_static_path / 'week_report'
    os.makedirs(week_report_dir, exist_ok=True)
    report_name = "PM项目周报"+next_week_range+".xlsx"
    file_path = os.path.join(week_report_dir, report_name)
    # file_path_obj = {'generate_docx': file_path}
    wb.save(file_path)
    print("✅ DeepSeek 输出已追加到 Excel，样式保留")
    file_path = file_path.split('rag/')[1] if 'rag/' in file_path else file_path
    file_path_obj = {'generate_docx':file_path}

    total_llm_resp = await generate_stream_result("报告内容如下：", thing_pattern=False)

    return total_llm_resp, "", file_path_obj



# ---------------- 插入表格数据并保持样式 ----------------
def append_table_data_with_style(ws, table_name, new_data):
    table_obj = ws.tables[table_name]
    table_range = table_obj.ref
    start_cell, end_cell = table_range.split(':')
    min_col_letter, min_row = '', 0
    max_col_letter, max_row = '', 0
    for ch in start_cell:
        if ch.isalpha():
            min_col_letter += ch
        else:
            min_row = int(start_cell[len(min_col_letter):])
    for ch in end_cell:
        if ch.isalpha():
            max_col_letter += ch
        else:
            max_row = int(end_cell[len(max_col_letter):])

    num_cols = ord(max_col_letter) - ord(min_col_letter) + 1
    # print(type(new_data), flush=True)
    # print(new_data[0], flush=True)
    # print(len(new_data), flush=True)
    # print(table_obj.ref)
    # print(num_cols)
    # for i,row in enumerate(new_data):
    #     print(len(row))
    #     print(row)
    if any(len(row) != num_cols for row in new_data):
        raise ValueError("每行数据列数必须和表格列数一致")

    border = Border(
        left=Side(border_style="thin", color="000000"),
        right=Side(border_style="thin", color="000000"),
        top=Side(border_style="thin", color="000000"),
        bottom=Side(border_style="thin", color="000000")
    )
    fill_even = PatternFill("solid", fgColor="FFFFFF")
    fill_odd = PatternFill("solid", fgColor="FFFFFF")

    start_insert_row = max_row + 1
    # ws.insert_rows(start_insert_row, amount = len(new_data))

    for r_idx, row in enumerate(new_data):
        ws.row_dimensions[start_insert_row + r_idx].height = 100
        fill = fill_odd if (start_insert_row + r_idx) % 2 == 1 else fill_even
        for c_idx, value in enumerate(row):
            col_letter = get_column_letter(ord(min_col_letter) - 64 + c_idx)
            cell = ws[f"{col_letter}{start_insert_row + r_idx}"]
            cell_ref = f"{col_letter}{start_insert_row + r_idx}"
            if isinstance(cell,MergedCell):
                print(cell_ref)
            cell.value = value
            cell.border = border
#            cell.fill = fill
            cell.font = Font(name="SimHei", size=10.5)
            cell.alignment = Alignment(horizontal="left", vertical="center",wrap_text=True)

    new_max_row = max_row + len(new_data)
    table_obj.ref = f"{min_col_letter}{min_row}:{max_col_letter}{new_max_row}"





    # start_insert_row = max_row + 1
    # ws.insert_rows(start_insert_row, amount = len(new_data))
    # for r_idx, row in enumerate(new_data):
    #     insert_row = start_insert_row + r_idx
    #     for c_idx, value in enumerate(row):
    #         col_letter = get_column_letter(ord(min_col_letter) - 64 + c_idx)
    #         cell = ws[f"{col_letter}{insert_row}"]
    #         cell_ref = f"{col_letter}{start_insert_row + r_idx}"
    #         if isinstance(cell, MergedCell):
    #             print(cell_ref)
    #         cell.value = value
    #         cell.border = border
    # #            cell.fill = fill
    # new_max_row = start_insert_row + len(new_data) - 1
    # table_obj.ref = f"{min_col_letter}{start_cell[len(min_col_letter)]}:{max_col_letter}{new_max_row}"
    #
    #

# ---------------- 拿到评论进行划分 ----------------

async def extract_summary(comment: str):
    # 提取 阶段目标、上期工作综述、本期重点工作
    pattern = re.compile(
        r"\[阶段目标\](.*?)\[上期工作综述\](.*?)\[本期重点工作\](.*)",
        re.S
    )
    match = pattern.search(comment)
    # print("match",flush=True)
    # print(match,flush=True)
    if match:
        return {
            "所属阶段目标": match.group(1).strip(),
            "上期工作综述": match.group(2).strip(),
            "本期重点工作": match.group(3).strip()
        }
    else:
        return None

async def batch_extract(result):
    extracted = []
    for row in result:
        task_name = row[0]   # 任务名称
        comment = row[5]     # 评论
        print(task_name,flush=True)
        print(comment,flush=True)

        summary = await extract_summary(comment)
        if summary:
            summary["任务名称"] = task_name  # 在结果里加上任务名称
            extracted.append(summary)
    return json.dumps(extracted, ensure_ascii=False, indent=2)



# ---------------- 追加数据 ----------------

#获取当前面板中状态为待办或者处理中的issue
async def current_task_issue():
    today = datetime.today()
    # monday = today - timedelta(days=today.weekday())
    # friday = monday + timedelta(days=4)
    #
    # week_range = f"{monday.strftime('%Y-%m-%d')}-{friday.strftime('%Y-%m-%d')}"
    monday = today - timedelta(days=today.weekday())
    next_monday = monday + timedelta(days=7)
    next_friday = next_monday + timedelta(days=4)

    next_week_range = f"{next_monday.strftime('%Y-%m-%d')}-{next_friday.strftime('%Y-%m-%d')}"
    # Jira 信息
    jira_url = JIRA_URL
    user = JIRA_USER
    token = JIRA_TOKEN
    board_id = 8

    # 调用接口
    url = f"{jira_url}/rest/agile/1.0/board/{board_id}/issue"
    params = {
        "jql": "issuetype in (故事 , 子任务, 任务) and status in(\"To Do\",\"In Progress\")",
        "fields": "key,summary,status,assignee,updated,comment,description",
        "maxResults": 1000
    }
    resp = requests.get(url, params=params, auth=HTTPBasicAuth(user, token))
    data = resp.json()

    records = []
    sort = 0
    for issue in data.get("issues", []):
        sort += 1
        key = issue["key"]
        fields = issue["fields"]
        summary = fields.get("summary", "")
        assignee = fields.get("assignee", {}).get("displayName", "")
        status = fields.get("status", {}).get("name", "")
        updated = fields.get("updated", "")
        description = fields.get("description", )
        # print("描述",description,flush=True)
        cooperate = await get_description_mentions(description)
        # print("cooperate",cooperate,flush=True)


        # 最新评论
        comments = fields.get("comment", {}).get("comments", [])
        # if comments:
        #     latest_comment = max(comments, key=lambda x: x["updated"])
        #     latest_comment_body = latest_comment["body"]
        #     if "192.168.1.46" in latest_comment_body:
        #         latest_comment_body = ""
        # else:
        #     latest_comment_body = ""
        latest_comment_body = ''
        if not comments:
            latest_comment_body = ""

        comments_sorted = sorted(comments, key=lambda c: c["updated"])
        for comment in comments_sorted:
            body = comment["body"]
            if "192.168.1.46" not in body:
                latest_comment_body = body

        records.append({
            "名称": summary,
            "任务编号": key,
            "负责人": assignee,
            '配合': cooperate,
            "时间": next_week_range,
            "任务交付": latest_comment_body,
            "说明": "",
            "完成情况": status,
        })
    result = [list(item.values()) for item in records]
    return result


async def generate_jira_week_report_detial():


    today = datetime.today()
    monday = today - timedelta(days=today.weekday())
    friday = monday + timedelta(days=4)

    week_range = f"{monday.strftime('%Y-%m-%d')}-{friday.strftime('%Y-%m-%d')}"
    # Jira 信息
    jira_url = JIRA_URL
    user = JIRA_USER
    token = JIRA_TOKEN
    board_id = JIRA_BOARD

    # 调用接口
    url = f"{jira_url}/rest/agile/1.0/board/{board_id}/issue"
    params = {
        "jql": "updated >= -7d and issuetype in (故事 , 子任务, 任务)",
        "fields": "key,summary,status,assignee,updated,comment,description",
        "maxResults": 1000
    }
    resp = requests.get(url, params=params, auth=HTTPBasicAuth(user, token))
    data = resp.json()

    records = []
    sort = 0
    for issue in data.get("issues", []):
        sort += 1
        key = issue["key"]
        fields = issue["fields"]
        summary = fields.get("summary", "")
        assignee = fields.get("assignee", {}).get("displayName", "")
        status = fields.get("status", {}).get("name", "")
        updated = fields.get("updated", "")
        description = fields.get("description", )
        cooperate = await get_description_mentions(description)
        # 最新评论
        comments = fields.get("comment", {}).get("comments", [])
        # if comments:
        #     latest_comment = max(comments, key=lambda x: x["updated"])
        #     latest_comment_body = latest_comment["body"]
        #     if  "192.168.1.46" in  latest_comment_body:
        #         latest_comment_body = ""
        # else:
        #     latest_comment_body = ""

        latest_comment_body = ''
        if not comments:
            latest_comment_body = ''

        comments_sorted = sorted(comments, key=lambda c: c["updated"])
        for comment in comments_sorted:
            body = comment["body"]
            if "192.168.1.46" not in body:
                latest_comment_body = body
        records.append({
            "名称": summary,
            "任务编号":key,
            "负责人": assignee,
            '配合':cooperate,
            "时间": week_range,
            "任务交付": latest_comment_body,
            "说明":"",
            "完成情况": status,
        })
    result = [list(item.values()) for item in records]
    # print("上周1", flush=True)
    # print(records, flush=True)
    return result


async def get_jira_sprint_detail():
    sprint = await get_jira_current_sprint()
    if not sprint:
        print("当前面板没有冲刺")
        return

    sprint_id = sprint.get("id")
    goal =  await get_sprint_goal(sprint_id)
    print(f"当前sprint_id为 {sprint_id}的目标是{goal}")
    return goal

#获取冲刺的目标
async def get_jira_current_sprint():

    # Jira 信息
    jira_url = JIRA_URL
    user = JIRA_USER
    token = JIRA_TOKEN
    board_id = JIRA_BOARD

    # 调用接口
    url = f"{jira_url}/rest/agile/1.0/board/{board_id}/sprint?state=active"
    resp = requests.get(url,  auth=HTTPBasicAuth(user, token))
    resp.raise_for_status()
    data = resp.json()
    sprints = data.get("values", [])
    if not sprints:
        return  None
    return sprints[0]



async def get_sprint_goal(sprint_id):
    jira_url = JIRA_URL
    user = JIRA_USER
    token = JIRA_TOKEN
    board_id = JIRA_BOARD

    url = f"{jira_url}/rest/agile/1.0/sprint/{sprint_id}"
    resp = requests.get(url, auth=HTTPBasicAuth(user, token))
    resp.raise_for_status()
    sprint_info = resp.json()
    return sprint_info.get("goal")







async def extract_mentions(description):
    """提取 mention 出来的用户名或 accountId"""
    # 兼容 Jira Server (~username) 和 Jira Cloud ([~accountid:xxx])
    if description is not None:
        matches = re.findall(r"~([a-zA-Z0-9._-]+)|~accountid:([a-zA-Z0-9:-]+)", description)
        # print("match",matches,flush=True)
        result = []
        if matches:
            for m in matches:
                if m[0]:  # username
                    result.append(("username", m[0]))
                elif m[1]:  # accountId
                    result.append(("accountId", m[1]))
            return result

        else:
            return []

async def resolve_user(user_refs):
    jira_url = JIRA_URL
    user = JIRA_USER
    token = JIRA_TOKEN
    board_id = JIRA_BOARD

    """调用 Jira API，把 username/accountId 转换成中文 displayName"""
    display_names = []
    for ref_type, ref_value in user_refs:
        if ref_type == "username":  # Jira Server/DC
            url = f"{JIRA_URL}/rest/api/2/user"
            params = {"username": ref_value}
        else:  # Jira Cloud
            url = f"{JIRA_URL}/rest/api/3/user"
            params = {"accountId": ref_value}

        resp = requests.get(url, auth=HTTPBasicAuth(user, token), params=params)
        if resp.status_code == 200:
            data = resp.json()
            display_names.append(data.get("displayName", ref_value))
        else:
            display_names.append(ref_value)  # 找不到就用原值
    return display_names


async   def get_description_mentions(description):
    """主函数：返回描述中所有 @人 的中文名，用顿号隔开"""
    refs = await extract_mentions(description)
    if  refs:
        names =  await resolve_user(refs)
        return "、".join(names)
    else:
        return ""







def transfer_sql_query(sql_query):
    """
    "queryList": [
    {
      "queryKey": "project_source",
      "queryValue": ${queryValue},
      "operateType": "in"
    }
    ]
    """
    ret = {}
    query_list = sql_query.get('queryList')
    if query_list:
        for query in query_list:
            ret.update({query['queryKey']: query['queryValue']})
        ret.update({'viewKey': "projectAllList"})
    return ret


def get_table_info_by_sql_query(sql_query, page, page_size):
    """根据sql_query获取数据查询详细"""
    new_query = transfer_sql_query(sql_query)
    new_query.update({'page': page, 'pageSize': page_size})
    data = get_table_info(new_query)
    return data
