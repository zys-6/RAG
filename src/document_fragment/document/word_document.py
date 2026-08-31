import copy
import io
import logging
import os
import re
import time
from collections import defaultdict, Counter
from pathlib import Path
from typing import Union, List, Dict
from xml.dom.minidom import parseString
from xml.etree import ElementTree

import tqdm
from docx import Document as PyDocument
from docx.document import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree

from document_fragment.document.base import BaseDocument
from document_fragment.document.fragment import PictureFragment, TableFragment, ContentFragment, BaseFragment
from document_fragment.document.fragment.chart_fragment import ChartFragment
from document_fragment.document.fragment.content_fragment import HtmlTag
from document_fragment.document.fragment.formula_fragment import FormulaFragment
from document_fragment.document.fragment.shape_fragment import ShapeFragment
from document_fragment.document.utils import (crop_image, add_namespaces_onto, add_namespaces_into, get_val_from_dom,
                                              get_text_from_run_dom, get_crop_shape, clean_wildcard)
from document_fragment.ooxml_to_latex.parser import load_string
from document_fragment.utils import is_windows, get_uuid_string

logger = logging.getLogger(__name__)
oxml_parser = etree.XMLParser(remove_blank_text=True, resolve_entities=False, huge_tree=True)

if is_windows():
    import pythoncom
    import win32com.client as win32
    from win32com.client import constants


class WordDocument(BaseDocument):

    def __init__(self, emf_to_png: bool = True, font_name_sort: Dict = None, **kwargs):
        super().__init__(**kwargs)
        """进行word解析"""
        self._doc = self._convert_to_docx()
        """获取整个文档的全局信息"""
        self._refs = self.get_refs(self._doc)
        self._rels = self._doc.part.rels
        self._numberings = self.get_numberings(self._doc)
        self._style_outlines, self._style2numid = self.get_style_outlines(self._doc)
        self._auto_mappings = defaultdict(int)
        self._numbering_stack = defaultdict(list)
        """转换为碎片"""
        self._font_name_sort = font_name_sort if font_name_sort else {"黑体": 2, "宋体": 1}
        self._emf_to_png = emf_to_png
        self._default_fmt = str
        self._support_fmt = {
            "chineseCountingThousand": lambda x: "一二三四五六七八九十"[x - 1],
            "chineseCounting": lambda x: "一二三四五六七八九十"[x - 1],
            "japaneseCounting": lambda x: "一二三四五六七八九十"[x - 1],
            "decimal": self._default_fmt,
            "lowerLetter": lambda x: "abcdefghijklmnopqrstuvwxyz"[x - 1],
            "lowerRoman": lambda x: 'ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ'[x - 1].lower()
        }
        self._tag_names = ("w:t",
                           "w:drawing",
                           "v:shape",
                           "c:chart",
                           "m:oMath",
                           "m:oMathPara",
                           "w:endnoteReference",
                           "w:footnoteReference",
                           "w:noBreakHyphen")
        if not self.lazy:
            self.fragments += self._convert_to_fragments()

    def _convert_to_docx(self):
        try:
            document = PyDocument(io.BytesIO(self._content))
        except PackageNotFoundError as e:
            """进行doc2docx"""
            """尝试进行doc -> docx"""
            try:
                self._content = self.convert_doc_to_docx(self._content, self.temp_directory)
                document = PyDocument(io.BytesIO(self._content))
            except Exception as e:
                logger.error("Convert doc to docx failed.")
                raise e
        except Exception as e:
            try:
                self._content = self.convert_doc_to_docx(self._content, self.temp_directory)
                document = PyDocument(io.BytesIO(self._content))
            except Exception as e:
                logger.error("Convert doc to docx failed.")
                raise e
        return document

    @classmethod
    def convert_doc_to_docx(cls, content: Union[bytes, Path, str], temp_directory=None) -> bytes:
        temp_directory = temp_directory if temp_directory else Path("static/tmp")
        temp_directory = Path(temp_directory)
        temp_directory.mkdir(parents=True, exist_ok=True)
        if isinstance(content, bytes):
            doc_filepath = temp_directory.joinpath(get_uuid_string() + ".doc")
            with open(doc_filepath, "wb") as fout:
                fout.write(content)
        elif isinstance(content, str) or isinstance(content, Path):
            doc_filepath = Path(content)
        else:
            raise TypeError("`content` must be `bytes`, `str`, `Path`")

        if is_windows():
            doc_filepath = Path(doc_filepath).resolve()
            try:
                pythoncom.CoInitialize()
                word = win32.gencache.EnsureDispatch('Word.Application')
                doc = word.Documents.Open(str(doc_filepath))
                doc.Activate()
                docx_filepath = doc_filepath.with_suffix(".docx")
                word.ActiveDocument.SaveAs(
                    str(docx_filepath), FileFormat=constants.wdFormatXMLDocument
                )
                doc.Close(False)
                word.Quit()
                with open(docx_filepath, "rb") as fin:
                    content = fin.read()
            except Exception as e:
                logger.error(f"_convert_doc_to_docx: {e}")
                raise e
            finally:
                doc_filepath.unlink(missing_ok=True)
                if doc_filepath.with_suffix(".docx").exists():
                    try:
                        doc_filepath.with_suffix(".docx").unlink(missing_ok=True)
                    except PermissionError:
                        logger.warning("未删除 {} 该文件".format(doc_filepath.with_suffix(".docx")))

                pythoncom.CoUninitialize()
            return content
        else:
            os.system(f"unoconv -d document --format=docx {doc_filepath}")
            time.sleep(3)
            if doc_filepath.with_suffix(".docx").exists():
                with open(doc_filepath.with_suffix(".docx"), "rb") as fin:
                    content = fin.read()
            else:
                logger.error(f"_convert_doc_to_docx: convert failed")
                raise FileNotFoundError(f"Can't find {doc_filepath.with_suffix('.docx').name}")
            return content

    @classmethod
    def get_refs(cls, doc: PyDocument):
        """存在三种参考文献：footnotes，endnotes，catalog==0但是存在自动编号的且自动编号以[]方式存在）"""
        ret = {
            "footnotes": {},
            "endnotes": {},
            "numberings": {}
        }
        for name in ret.keys():
            idx = 1
            for key, val in doc.part.rels.items():
                if val._reltype.endswith(name):
                    xml = val._target.blob
                    xml = etree.fromstring(xml, parser=oxml_parser)
                    dom = parseString(etree.tounicode(xml))
                    notes = dom.getElementsByTagName("w:{}".format(name[:-1]))
                    for note in notes:
                        _id = note.getAttribute("w:id")
                        if _id == "": continue
                        text = ""
                        for t in note.getElementsByTagName("w:t"):
                            for n in t.childNodes:
                                text = text + n.data
                        if not text.strip(): continue
                        if text.strip().startswith("[]"):
                            text = text[2:]
                        text = f'[{idx}]' + text.strip()
                        idx += 1
                        ret[name][str(_id)] = text
        return ret

    @classmethod
    def get_numberings(cls, doc: PyDocument):
        try:
            dom = parseString(
                etree.tounicode(etree.fromstring(doc.part.numbering_part._element.xml, parser=oxml_parser)))
        except Exception as e:
            logger.warning("get_numberings: {}".format(e))
            return {}
        num_id2num = defaultdict(dict)
        for child in dom.firstChild.childNodes:
            if hasattr(child, "tagName") and child.tagName == 'w:abstractNum':
                abstract_num_id = child.getAttribute("w:abstractNumId")
                for lvl in child.getElementsByTagName("w:lvl"):
                    ilvl = lvl.getAttribute("w:ilvl")
                    num_id2num[abstract_num_id][ilvl] = {
                        "start": lvl.getElementsByTagName("w:start")[0].getAttribute(
                            "w:val") if lvl.getElementsByTagName(
                            "w:start") else 0,
                        "numFmt": lvl.getElementsByTagName("w:numFmt")[0].getAttribute("w:val"),
                        "lvlText": lvl.getElementsByTagName("w:lvlText")[0].getAttribute("w:val")
                    }
        numberings = {}
        for w_num in dom.getElementsByTagName("w:num"):
            num_id = w_num.getAttribute("w:numId")
            abstract_num_id = w_num.getElementsByTagName("w:abstractNumId")[0].getAttribute("w:val")
            numberings[num_id] = copy.deepcopy(num_id2num[abstract_num_id])
            """更新重新编号部分"""
            for w_lvlOverride in w_num.getElementsByTagName("w:lvlOverride"):
                _ilvl = w_lvlOverride.getAttribute("w:ilvl")
                for w_startOverride in w_lvlOverride.getElementsByTagName("w:startOverride"):
                    _start = w_startOverride.getAttribute("w:val")
                    if int(_ilvl) >= int(_start):
                        numberings[num_id][_ilvl]['startOverride'] = _start

        return numberings

    @classmethod
    def get_style_outlines(cls, doc: PyDocument):
        mappings = {}
        style2numid = {}
        dom = parseString(etree.tounicode(etree.fromstring(doc.styles.element.xml, parser=oxml_parser)))
        outlineLvls = set()
        for element in dom.getElementsByTagName("w:outlineLvl"):
            try:
                outlineLvl = int(element.getAttribute("w:val").strip())
                outlineLvl = outlineLvl + 1
            except Exception:
                outlineLvl = 0
            try:
                node = element.parentNode.parentNode
                style_id = node.getAttribute("w:styleId")
                for numpr in node.getElementsByTagName("w:numPr"):
                    _num_id = None
                    _ivl = "0"
                    for numid in numpr.getElementsByTagName("w:numId"):
                        _num_id = numid.getAttribute("w:val")
                    for ivl in numpr.getElementsByTagName("w:ilvl"):
                        _ivl = ivl.getAttribute("w:val")
                    if _num_id:
                        style2numid[style_id] = (_num_id, _ivl)
            except Exception as e:
                logger.warning("get_style_outlines: {}".format(e))
                continue
            outlineLvls.add(outlineLvl)
            mappings[style_id] = outlineLvl
        """还有一些没有outlineLvl的元素"""
        for element in dom.getElementsByTagName("w:style"):
            style_id = element.getAttribute("w:styleId")
            if style_id and style_id not in style2numid:
                for numpr in element.getElementsByTagName("w:numPr"):
                    _num_id = None
                    _ivl = "0"
                    for numid in numpr.getElementsByTagName("w:numId"):
                        _num_id = numid.getAttribute("w:val")
                    for ivl in numpr.getElementsByTagName("w:ilvl"):
                        _ivl = ivl.getAttribute("w:val")
                    if _num_id:
                        style2numid[style_id] = (_num_id, _ivl)
        return mappings, style2numid

    @classmethod
    def block_generator(cls, parent: Document):
        if isinstance(parent, Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @classmethod
    def get_dom_from_xml(cls, xml: str):
        return parseString(etree.tounicode(etree.fromstring(xml, parser=oxml_parser)))

    @classmethod
    def is_picture_block_dom(cls, dom):
        try:
            return len(dom.getElementsByTagName("w:drawing")) != 0 and len(dom.getElementsByTagName("a:blip")) != 0
        except Exception as e:
            logger.warning("is_picture_block_dom: {}".format(e))
            return False

    @classmethod
    def is_table_block_dom(cls, dom):
        try:
            return len(dom.getElementsByTagName("w:tbl")) != 0
        except Exception as e:
            logger.warning("is_table_block_dom: {}".format(e))
            return False

    @classmethod
    def is_formula_block_dom(cls, dom):
        try:
            return len(dom.getElementsByTagName("m:oMathPara")) != 0
        except Exception as e:
            logger.warning("is_formula_block_dom: {}".format(e))
            return False

    @classmethod
    def is_chart_block_dom(cls, dom):
        try:
            return len(dom.getElementsByTagName("c:chart")) != 0
        except Exception as e:
            logger.warning("is_chart_block_dom: {}".format(e))
            return False

    @classmethod
    def is_shape_block_dom(cls, dom):
        try:
            return len(dom.getElementsByTagName("v:shape")) != 0
        except Exception as e:
            logger.warning("is_shape_block_dom: {}".format(e))
            return False

    def generate_picture_fragments_from_dom(self, dom) -> List[PictureFragment]:

        def get_picture_from_drawings(drawings):
            ret = []
            for drawing in drawings:
                for a_blip in drawing.getElementsByTagName("a:blip"):
                    rect = a_blip.parentNode.getElementsByTagName("a:srcRect")
                    if rect:
                        l = rect[0].getAttribute("l")
                        r = rect[0].getAttribute("r")
                        t = rect[0].getAttribute("t")
                        b = rect[0].getAttribute("b")
                    else:
                        l, r, t, b = None, None, None, None
                    rel_id = a_blip.getAttribute("r:embed")
                    if rel_id in self._rels and 'image' in self._rels[rel_id]._target._content_type:
                        suffix = self._rels[rel_id]._target._content_type.split("/")[-1]
                        if suffix in ("x-emf", "x-wmf"):
                            suffix = "emf" if suffix == 'x-emf' else "wmf"
                            ret.append(PictureFragment(blob=self._rels[rel_id]._target.blob,
                                                       caption="",
                                                       suffix=f".{suffix}"))
                        else:
                            blob = crop_image(self._rels[rel_id]._target.blob, l, r, t, b)
                            ret.append(PictureFragment(blob=blob, caption="", suffix=f".{suffix}"))
            return ret

        fragments = []

        if "".join([get_text_from_run_dom(r) for r in dom.getElementsByTagName("w:r")]).strip():
            content_fragments = []
            for run in dom.getElementsByTagName("w:r"):
                drawings = run.getElementsByTagName("w:drawing")
                if drawings:
                    if content_fragments:
                        fragments.append(
                            ContentFragment(text="".join([_fragment.text for _fragment in content_fragments])))
                        content_fragments = []
                    fragments.extend(get_picture_from_drawings(drawings))
                elif get_text_from_run_dom(run).strip():
                    content_fragments.append(ContentFragment(text=get_text_from_run_dom(run).strip()))
            if content_fragments:
                fragments.append(ContentFragment(text="".join([_fragment.text for _fragment in content_fragments])))
                content_fragments = []
        else:
            fragments.extend(get_picture_from_drawings(dom.getElementsByTagName("w:drawing")))
        """最好将fragments最后的ContentFragment与平常的Fragment进行分开"""
        if fragments and isinstance(fragments[-1], ContentFragment):
            fragments[-1].connect_next = False
        return fragments

    def generate_formula_fragments_from_dom(self, dom) -> List[FormulaFragment]:

        def construct_xml(xml: str):
            xml = xml.replace("<m:oMathPara",
                              "<m:oMathPara xmlns:a=\"http://schemas.openxmlformats.org/drawingml/2006/main\" xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\"")
            xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n""" + xml
            return xml

        def get_o_math_xml(math_para):
            try:
                xml = math_para.toxml()
                xml = construct_xml(xml)
                xml = ElementTree.fromstring(xml.encode("utf-8"))
            except Exception as e:
                logger.warning(f"generate_formula_fragments_from_dom: {e}")
                xml = None
            return xml

        fragments = []
        if "".join([get_text_from_run_dom(r) for r in dom.getElementsByTagName("w:r")]).strip():
            """存在文字需要按照run结构进行，组织成一个行内ContentFragment"""
            content_fragments = []
            for run in dom.getElementsByTagName("w:r"):
                run_math = run.getElementsByTagName("m:oMathPara")
                if run_math:
                    if content_fragments:
                        fragments.append(
                            ContentFragment(text="".join([_fragment.text for _fragment in content_fragments])))
                        content_fragments = []
                    for math_para in run_math:
                        xml = get_o_math_xml(math_para)
                        if xml is None:
                            continue
                        xml_str = ElementTree.tostring(xml, encoding="unicode")
                        for o_math in load_string(xml_str):
                            fragments.append(FormulaFragment(latex=o_math.latex,
                                                             tag=FormulaFragment.convert_latex_to_mathml(o_math.latex)))
                elif get_text_from_run_dom(run).strip():
                    content_fragments.append(ContentFragment(text=get_text_from_run_dom(run).strip()))
            if content_fragments:
                fragments.append(ContentFragment(text="".join([_fragment.text for _fragment in content_fragments])))
                content_fragments = []
        else:
            for math_para in dom.getElementsByTagName("m:oMathPara"):
                xml = get_o_math_xml(math_para)
                if xml is None:
                    continue
                xml_str = ElementTree.tostring(xml, encoding='unicode')
                for o_math in load_string(xml_str):
                    fragments.append(FormulaFragment(latex=o_math.latex,
                                                     tag=FormulaFragment.convert_latex_to_mathml(o_math.latex)))
        if fragments and isinstance(fragments[-1], ContentFragment):
            fragments[-1].connect_next = False
        return fragments

    def generate_shape_fragments_from_dom(self, dom) -> List[Union[ShapeFragment, ContentFragment]]:

        def get_shape_image(shape):
            fragments = []
            for imagedata in shape.getElementsByTagName("v:imagedata"):
                rel_id = imagedata.getAttribute("r:id")
                l, r, t, b = get_crop_shape(imagedata)
                if rel_id in self._rels and "image" in self._rels[rel_id]._target._content_type:
                    blob = self._rels[rel_id]._target.blob
                    suffix = self._rels[rel_id]._target._content_type.split("/")[-1]
                    if 'emf' in suffix or 'wmf' in suffix:
                        suffix = 'emf' if 'emf' in suffix else 'wmf'
                        if blob is not None:
                            fragments.append(ShapeFragment(blob=blob, suffix="." + suffix, caption=""))
                        else:
                            fragments.append(ShapeFragment(blob=self._rels[rel_id]._target.blob,
                                                           suffix="." + suffix,
                                                           caption=""))
                    elif suffix.lower() in ("jpg", "jpeg", "png"):
                        """jpg png"""
                        fragments.append(ShapeFragment(blob=crop_image(blob, l, r, t, b),
                                                       suffix="." + suffix.lower(),
                                                       caption=""))
                    else:
                        logger.warning("Unrecognized suffix {}".format(suffix))
                        fragments.append(ShapeFragment(blob=None,
                                                       suffix="." + suffix.lower(),
                                                       caption=""))
            return fragments

        fragments = []
        if "".join([get_text_from_run_dom(r) for r in dom.getElementsByTagName("w:r")]).strip():
            """"""
            content_fragments = []
            for run in dom.getElementsByTagName("w:r"):
                shapes = run.getElementsByTagName("v:shape")
                if shapes:
                    if content_fragments:
                        fragments.append(
                            ContentFragment(text="".join([_fragment.text for _fragment in content_fragments])))
                        content_fragments = []
                    for shape in shapes:
                        fragments += get_shape_image(shape)
                elif get_text_from_run_dom(run).strip():
                    content_fragments.append(ContentFragment(text=get_text_from_run_dom(run).strip()))
            if content_fragments:
                fragments.append(ContentFragment(text="".join([_fragment.text for _fragment in content_fragments])))
                content_fragments = []
        else:
            for shape in dom.getElementsByTagName("v:shape"):
                fragments += get_shape_image(shape)
        if fragments and isinstance(fragments[-1], ContentFragment):
            fragments[-1].connect_next = False
        return fragments

    def generate_chart_fragments_from_dom(self, dom) -> List[ChartFragment]:
        fragments = []
        for chart in dom.getElementsByTagName("c:chart"):
            rel_id = chart.getAttribute("r:id")
            xml = self._rels[rel_id]._target.blob.decode("utf-8")
            fragments.append(ChartFragment(xml=xml))
        if fragments and isinstance(fragments[-1], ContentFragment):
            fragments[-1].connect_next = False
        return fragments

    def generate_picture_fragments_from_cell(self, cell):
        dom = parseString(etree.tounicode(etree.fromstring(cell._element.xml, parser=oxml_parser)))
        return self.generate_picture_fragments_from_dom(dom)

    def generate_shape_fragments_from_cell(self, cell):
        dom = parseString(etree.tounicode(etree.fromstring(cell._element.xml, parser=oxml_parser)))
        return self.generate_shape_fragments_from_dom(dom)

    def generate_chars_fragments_from_cell(self, cell):
        pass

    @classmethod
    def is_picture_cell(cls, cell):
        dom = parseString(etree.tounicode(etree.fromstring(cell._element.xml, parser=oxml_parser)))
        return cls.is_picture_block_dom(dom)

    def generate_table_cell_data(self, table):
        fragments = []
        table_data = []
        for row in table.rows:
            """在这里判断cell里面是文本还是图片"""
            row_data = []
            for _cell in row.cells:
                cell_dom = parseString(etree.tounicode(etree.fromstring(_cell._element.xml, parser=oxml_parser)))
                if self.is_picture_block_dom(cell_dom):
                    """在这里获取图片的二进制数据"""
                    fragments += self.generate_picture_fragments_from_dom(cell_dom)
                elif self.is_shape_block_dom(cell_dom):
                    fragments += self.generate_shape_fragments_from_dom(cell_dom)
                elif self.is_chart_block_dom(cell_dom):
                    fragments += self.generate_chart_fragments_from_dom(cell_dom)
                elif self.is_formula_block_dom(cell_dom):
                    fragments += self.generate_formula_fragments_from_dom(cell_dom)
                if _cell.tables:
                    for _table in _cell.tables:
                        fragments += self.generate_table_fragments_from_block(_table)
                row_data.append(_cell.text)
            table_data.append(row_data)
        return table_data, fragments

    def generate_table_fragments_from_block(self, table) -> List[Union[TableFragment, PictureFragment]]:
        fragments = []
        cells = table._cells
        cols = table._column_count
        length = len(cells)
        merge_info = []
        del_info = []
        for i, cell in enumerate(cells):
            if cell in cells[:i]:  # 如果该单元格不是在表中第一次出现则跳过
                continue
            for j in range(length - 1, -1, -1):  # 倒序查找
                if cell is cells[j]:  # 找到"相同"的单元格，如果没有"合并"单元格，则会倒序找到"自己"
                    break
            if i != j:
                r1, c1 = divmod(i, cols)  # 合并单元格区域的"起始"位置，同时也是左上角单元格的行列坐标
                r2, c2 = divmod(j, cols)  # 合并单元格区域的"结束"位置，同时也是右下角单元格的行列信息
                merge_info.append({
                    "row_index": r1,
                    "col_index": c1,
                    "row_span": (r2 - r1) + 1,
                    "col_span": (c2 - c1) + 1
                })
        for _merge_info in merge_info:
            for col_index in range(_merge_info["col_index"] + 1, _merge_info["col_index"] + _merge_info["col_span"]):
                for row_index in range(_merge_info['row_index'], _merge_info['row_index'] + _merge_info['row_span']):
                    del_info.extend([{"row_index": row_index,
                                      "col_index": col_index}])
            for row_index in range(_merge_info['row_index'] + 1, _merge_info['row_index'] + _merge_info['row_span']):
                for col_index in range(_merge_info["col_index"], _merge_info["col_index"] + _merge_info["col_span"]):
                    del_info.extend([{"row_index": row_index,
                                      "col_index": col_index}])
        table_data, _fragments = self.generate_table_cell_data(table)
        fragments += _fragments

        width_list = []
        dom = parseString(etree.tounicode(etree.fromstring(table.table._element.xml, parser=oxml_parser)))
        for element in dom.getElementsByTagName("w:tblGrid"):
            for item in element.getElementsByTagName("w:gridCol"):
                width_list.append(int(item.getAttribute("w:w")))
            if width_list:
                """转换为比例"""
                total = sum(width_list)
                if total != 0:
                    width_list = list(map(lambda x: x / total, width_list))
                else:
                    width_list = []
            break

        fragments.append(TableFragment(data=table_data, merge_cell=merge_info, delete_cell=del_info, width=width_list))
        return fragments

    def get_auto_numbering_fmt_from_dom(self, dom):
        num_prs = dom.getElementsByTagName("w:numPr")
        fmt = None
        if num_prs:
            num_id = num_prs[0].getElementsByTagName("w:numId")[0].getAttribute("w:val")
            if num_id != '0':
                i_lvl = num_prs[0].getElementsByTagName("w:ilvl")[0].getAttribute("w:val")
                fmt = self._numberings[num_id][i_lvl]
                fmt['num_id'] = num_id
                fmt['i_lvl'] = i_lvl
        else:
            for p_style in dom.getElementsByTagName("w:pStyle"):
                p_style_id = p_style.getAttribute("w:val")
                num_id, i_lvl = self._style2numid.get(p_style_id, (None, None))
                if not num_id or num_id == '0':
                    continue
                if dom.getElementsByTagName("w:ilvl"):
                    i_lvl = dom.getElementsByTagName("w:ilvl")[0].getAttribute("w:val")
                fmt = self._numberings[num_id].get(i_lvl, None)
                if fmt:
                    fmt['num_id'] = num_id
                    fmt['i_lvl'] = i_lvl
                    break
        return fmt

    def get_auto_numbering_from_dom(self, dom, outline):
        default_fmt = self._default_fmt
        support_fmt = self._support_fmt
        numprs = dom.getElementsByTagName("w:numPr")
        if numprs:
            numid = numprs[0].getElementsByTagName("w:numId")[0].getAttribute("w:val")
            if numid == '0': return ""
            ilvl = numprs[0].getElementsByTagName("w:ilvl")[0].getAttribute("w:val")
            fmt = self._numberings[numid][ilvl]
        else:
            """读取style中的编号"""
            fmt = None
            for p_style in dom.getElementsByTagName("w:pStyle"):
                p_style_id = p_style.getAttribute("w:val")
                numid, ilvl = self._style2numid.get(p_style_id, (None, None))
                if not numid or numid == '0':
                    continue
                if dom.getElementsByTagName("w:ilvl"):
                    ilvl = dom.getElementsByTagName("w:ilvl")[0].getAttribute("w:val")
                fmt = self._numberings[numid].get(ilvl, None)
                if fmt:
                    break
            if fmt is None or not numid:
                return ""
        # catalog = int(ilvl) + 1  # fmt['lvlText'].count('%')
        catalog = int(outline)
        pattern = '[0-9a-zA-Z一二三四五六七八九十ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]'
        if self._numbering_stack["numid"]:
            if self._numbering_stack["numid"][-1][0] < catalog:
                _string, _format = self._numbering_stack["numid"][-1][1], self._numbering_stack["numid"][-1][2]
                ret = re.sub("%[0-9]", support_fmt.get(fmt['numFmt'], default_fmt)(int(fmt['start'])),
                             re.sub(_format, _string, fmt['lvlText']))
                self._numbering_stack["numid"].append((catalog, ret, fmt['lvlText'], int(fmt['start'])))
            else:
                while self._numbering_stack["numid"] and self._numbering_stack["numid"][-1][0] > catalog:
                    self._numbering_stack["numid"].pop(-1)
                if self._numbering_stack["numid"]:
                    _string, _format, _num = self._numbering_stack["numid"][-1][1:]

                    __format = clean_wildcard(_format, sp='()[]')
                    eng = re.sub(f'%{pattern}+?', f"({pattern}+?)", __format)
                    rst = {_k: _v for _k, _v in zip(re.findall(f'%{pattern}+?', _format),
                                                    re.findall(eng, _string)[0])}
                    if 'startOverride' in fmt:
                        _num = int(fmt['startOverride'])
                        del fmt['startOverride']
                    else:
                        _num += 1
                    if rst:
                        lvl_text = fmt['lvlText']
                        for _f, _i in rst.items():
                            if _f in re.findall(f'%{pattern}+?', fmt['lvlText'])[:-1]:
                                lvl_text = re.sub(_f, _i, lvl_text)
                        ret = re.sub(f'%{pattern}+?', '{}', lvl_text).format(support_fmt.get(fmt['numFmt'],
                                                                                             default_fmt)(_num))
                        # ret = re.sub('%[0-9]+?', '{}', fmt['lvlText']).format(*(list(rst[0])[:-1] +
                        #                                                         [support_fmt.get(fmt['numFmt'],
                        #                                                                          default_fmt)(_num)]))
                    else:
                        raise ValueError("Error")
                else:
                    ret = re.sub(f'%{pattern}+?', support_fmt.get(fmt['numFmt'], default_fmt)(int(fmt['start'])),
                                 fmt['lvlText'])
                    _num = int(fmt['start'])
                self._numbering_stack["numid"].append((catalog, ret, fmt['lvlText'], _num))
        else:
            """第一个"""
            ret = re.sub('%[0-9]+?', support_fmt.get(fmt['numFmt'], default_fmt)(int(fmt['start'])),
                         fmt['lvlText'])
            _num = int(fmt['start'])
            self._numbering_stack["numid"].append((catalog, ret, fmt['lvlText'], _num))
        return ret + " "

    def get_text_from_block(self, block, outline: int):

        def item_generator(dom, names):
            for child in dom.childNodes:
                if not child: continue
                if hasattr(child, "tagName") and child.tagName in names:
                    yield child
                elif hasattr(child, "childNodes"):
                    yield from item_generator(child, names)

        def get_text_from_element(element, tag_name="w:t"):
            text = ""
            for tag in element.getElementsByTagName(tag_name):
                if hasattr(tag, "childNodes"):
                    text += "".join([child.data for child in tag.childNodes])
            return text

        def get_inline_formula(dom):
            for o_math in load_string(dom.toxml()):
                return o_math.latex
            return ""

        text = ""
        dom = parseString(etree.tounicode(etree.fromstring(block._element.xml, parser=oxml_parser)))
        namespaces = [f"xmlns:{k}='{v}'" for k, v in block._element.nsmap.items()]
        for item in item_generator(dom, ("w:r", "m:oMath")):
            if hasattr(item, "tagName") and item.tagName == 'w:r':
                xml = add_namespaces_into(item.toxml(), namespaces, "<w:r ")  # 需要加入命名空间
                run_dom = parseString(etree.tounicode(etree.fromstring(xml, parser=oxml_parser)))
                for ref in run_dom.getElementsByTagName("w:endnoteReference"):
                    try:
                        matched = re.search('^\[([0-9]+)\]', self._refs['endnotes'][ref.getAttribute('w:id')])
                    except Exception:
                        logger.warning("get_text_from_block: 未找到endnotes: {}".format(ref.getAttribute("w:id")))
                        matched = None
                    if matched:
                        text = text + matched.group(1)
                for ref in run_dom.getElementsByTagName("w:footnoteReference"):
                    try:
                        matched = re.search('^\[([0-9]+)\]', self._refs['footnotes'][ref.getAttribute('w:id')])
                    except Exception:
                        logger.warning("get_text_from_block: 未找到footnotes: {}".format(ref.getAttribute("w:id")))
                        matched = None
                    if matched:
                        text = text + matched.group(1)
                text = text + get_text_from_run_dom(item)
            elif hasattr(item, "tagName") and item.tagName == "m:oMath":
                """内联公式, 需转换成latex"""
                xml = add_namespaces_onto(item.toxml(), namespaces, "<m:oMathPara ")  # 需要加入命名空间
                run_dom = parseString(etree.tounicode(etree.fromstring(xml, parser=oxml_parser)))
                try:
                    formula_text = get_inline_formula(run_dom)
                except Exception as e:
                    logger.error("get_text_from_block: {}".format(e))
                    formula_text = " "
                text = text + "${}$".format(formula_text)
        if text.strip():
            try:
                _number = self.get_auto_numbering_from_dom(dom, outline)
            except Exception as e:
                logger.error("get_auto_number_from_dom: {}".format(e))
                _number = ""
            text = _number + text.strip()
        return text

    def get_style_id_from_dom(self, dom):
        style_ids = dom.getElementsByTagName("w:pStyle")
        if style_ids:
            return Counter([style_id.getAttribute("w:val") for style_id in style_ids]).most_common(1)[0][0]
        return ""

    def get_outline_from_block(self, block):
        dom = parseString(etree.tounicode(etree.fromstring(block._element.xml, parser=oxml_parser)))
        lvls = dom.getElementsByTagName("w:outlineLvl")
        outline = None
        if lvls:
            try:
                outline = int(lvls[0].getAttribute("w:val")) + 1
            except Exception as e:
                logger.error(f"get_outline_from_block: {e}")
        if outline is None:
            """尝试获取"""
            style_id = self.get_style_id_from_dom(dom)
            outline = self._style_outlines.get(style_id, 0)
        return outline

    def get_text_fields_from_block(self, block):
        fields = {}
        _xml = etree.fromstring(block.style.element.xml, parser=oxml_parser)
        _dom = parseString(etree.tounicode(_xml))
        font_size_counter = defaultdict(int)
        font_name_counter = defaultdict(int)
        bold = False
        for run in block.runs:
            xml = etree.fromstring(run.element.xml, parser=oxml_parser)
            dom = parseString(etree.tounicode(xml))
            w_ts = dom.getElementsByTagName("w:t")
            texts = []
            for w_t in w_ts:
                for child in w_t.childNodes:
                    texts.append(child.data)
            if not texts:
                continue
            """尝试获取加粗"""
            bold = get_val_from_dom(dom, "w:b", "w:val")
            if not bold:
                bold = get_val_from_dom(_dom, "w:b", "w:val") or get_val_from_dom(_dom, "w:bCs", "w:val")
            if bold is None or bold == '0' or bold == False:
                bold = False
            else:
                bold = True
            """尝试获取字体大小"""
            if run.font.size is not None:
                font_size_counter[str(int(run.font.size.pt * 2))] += 1
            else:
                w_szs = dom.getElementsByTagName("w:sz")
                if not w_szs:
                    """尝试获取style的字体"""
                    w_szs = _dom.getElementsByTagName("w:sz")
                if w_szs:
                    font_size = w_szs[0].getAttribute("w:val")
                    font_size_counter[font_size] += 1
            """尝试获取字体"""
            w_rfonts = dom.getElementsByTagName('w:rFonts')
            if w_rfonts:
                for w_rfont in w_rfonts:
                    font_name = w_rfont.getAttribute('w:eastAsia')
                    font_name_counter[font_name] += 1
        if font_size_counter:
            font_size = sorted(font_size_counter.items(), key=lambda x: x[1], reverse=True)[0][0]
        else:
            font_size = None
        if font_name_counter:
            font_name = sorted(font_name_counter.items(), key=lambda x: x[1], reverse=True)[0][0]
        else:
            font_name = None
        """尝试获取style的字体大小和名称用来弥补runs"""
        xml = etree.fromstring(block.style.element.xml, parser=oxml_parser)
        dom = parseString(etree.tounicode(xml))
        if not font_size:
            w_szs = dom.getElementsByTagName("w:sz")
            if not w_szs:
                w_szs = dom.getElementsByTagName("w:szCs")
            if w_szs:
                font_size = w_szs[0].getAttribute("w:val")
        if not font_name:
            w_rfonts = dom.getElementsByTagName('w:rFonts')
            if w_rfonts:
                for w_rfont in w_rfonts:
                    font_name = w_rfont.getAttribute('w:eastAsia')
        if not font_size:
            font_size = '14'
        if not font_name:
            font_name = ""
        fields['font_size'] = font_size
        fields['font_name'] = font_name
        fields['bold'] = bold
        """alignment"""
        if block.style:
            try:
                alignment = block.style.paragraph_format.alignment \
                    if block.alignment is None else block.alignment
            except Exception as e:
                alignment = None
        else:
            alignment = None
        if alignment is None:
            alignment = "LEFT"
        elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
            alignment = "CENTER"
        elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
            alignment = "RIGHT"
        else:
            alignment = "LEFT"
        fields['alignment'] = alignment
        fields['outline'] = self.get_outline_from_block(block)
        return fields

    def generate_content_fragments_from_block(self, block):
        fragments = []
        fields = self.get_text_fields_from_block(block)
        text = self.get_text_from_block(block, fields['outline'])
        if not text.strip():
            # logger.warning("generate_content_fragments_from_block: Empty paragraph")
            return fragments
        """extract inline formula"""
        texts = re.split("(\$.*?\$)", text)
        _texts = []
        tags = []
        for _text in texts:
            if _text.startswith("$") and _text.endswith("$"):
                tags.append(HtmlTag(start_offset=len("".join(_texts)),
                                    end_offset=len("".join(_texts)),
                                    type="math",
                                    content=FormulaFragment.convert_latex_to_mathml(_text)))
            else:
                _texts.append(_text)
        text = "".join(_texts)
        fragments.append(ContentFragment(text=text, **fields, tags=tags))
        return fragments

    @classmethod
    def _match_fragment_captions(cls, fragments: List[BaseFragment]):
        idx = 0
        ret = []
        while idx < len(fragments):
            if isinstance(fragments[idx], TableFragment):
                """获取上一个"""
                if ret and isinstance(ret[-1], ContentFragment) and \
                        (ret[-1].text.strip().startswith("表") or
                         ret[-1].text.strip().startswith("tab")) and \
                        len(ret[-1].text.strip()) < 100:
                    fragments[idx].caption = ret[-1].text.replace('\t', '')
                    ret.pop(-1)
                elif len(ret) > 1 and isinstance(ret[-2], ContentFragment) and \
                        (ret[-2].text.strip().startswith("表") or
                         ret[-2].text.strip().startswith("tab")) and \
                        len(ret[-2].text.strip()) < 100:
                    fragments[idx].caption = ret[-2].text.replace('\t', '')
                    ret.pop(-2)
                ret.append(fragments[idx])
            elif isinstance(fragments[idx], PictureFragment) \
                    or isinstance(fragments[idx], ShapeFragment) \
                    or isinstance(fragments[idx], ChartFragment):
                if idx + 1 < len(fragments) and \
                        isinstance(fragments[idx + 1], ContentFragment) and \
                        (fragments[idx + 1].text.strip().startswith("图") or
                         fragments[idx + 1].text.strip().lower().startswith("fig")) and \
                        len(fragments[idx + 1].text.strip()) < 100:
                    """获取下一个或者下下一个"""
                    fragments[idx].caption = fragments[idx + 1].text.replace("/", "-").replace('\t', '')
                    ret.append(fragments[idx])
                    idx += 1
                elif idx + 2 < len(fragments) and isinstance(fragments[idx + 2], ContentFragment) and \
                        isinstance(fragments[idx + 1], ContentFragment) and \
                        (fragments[idx + 2].text.strip().startswith("图") or
                         fragments[idx + 2].text.strip().lower().startswith("fig")) and \
                        len(fragments[idx + 2].text.strip()) < 100:
                    # if len(fragments[idx + 1].text.strip() + fragments[idx + 2].text.strip()) < 100:
                    #     fragments[idx].caption = fragments[idx + 2].text.replace("/", "-").replace('\t', '') + \
                    #                              fragments[idx + 1].text.replace("/", "-").replace('\t', '')
                    #     ret.append(fragments[idx])
                    #     idx += 2
                    # else:
                    fragments[idx].caption = fragments[idx + 2].text.replace("/", "-").replace("\t", '')
                    ret.append(fragments[idx])
                    ret.append(fragments[idx + 1])
                    idx += 2

                elif ret and isinstance(ret[-1], ContentFragment) and \
                        (ret[-1].text.strip().startswith("图") or
                         ret[-1].text.strip().lower().startswith("fig")) and \
                        len(ret[-1].text.strip()) < 100:
                    """如果下一个没有尝试获取上一个"""
                    fragments[idx].caption = ret[-1].text.replace("/", "-").replace('\t', '')
                    ret.pop(-1)
                    ret.append(fragments[idx])
                else:
                    ret.append(fragments[idx])
            else:
                ret.append(fragments[idx])
            idx += 1
        return ret

    def _recognize_outlines_by_chars(self, fragments: List[BaseFragment]):

        def get_font_name_sort_value(font_name):
            for _key in self._font_name_sort:
                if font_name in _key:
                    return self._font_name_sort[_key]
            return sum(font_name.encode('utf-8'))

        need = True
        for fragment in fragments:
            if isinstance(fragment, ContentFragment) and fragment.outline != 0:
                need = False
                break

        if need:

            if not fragments:
                return
            clusters = defaultdict(list)
            for fragment in fragments:
                if isinstance(fragment, ContentFragment):
                    key = (
                        int(fragment.font_size), fragment.bold, get_font_name_sort_value(fragment.font_name))
                    clusters[key].append(fragment)
            if len(clusters) == 0:
                return
            counter = defaultdict(float)
            for key in clusters:
                _font_size, _bold, _font_name = key
                _fragments = clusters[key]
                _count = 0
                for _fragment in _fragments:
                    _count += len(_fragment.text)
                counter[key] = _count / len(_fragments)
            """默认counter 最大的那个作为正文"""
            counter = sorted(counter.items(), key=lambda x: x[1], reverse=True)
            max_key = counter[0][0]
            for _fragment in clusters[max_key]:
                _fragment.outline = 0
            clusters = sorted(clusters.items(), key=lambda x: x[0], reverse=True)
            """最后一级别默认为正文"""
            level = 0
            for index in range(min(len(clusters), 6)):
                if clusters[index][0][0] >= max_key[0] and clusters[index][0] != max_key:
                    for fragment in clusters[index][1]:
                        fragment.outline = level + 1
                    level += 1

    def add_refs_into_fragments(self, fragments):
        """首先过滤一下需要加入的fragments"""

        def _can(ref):
            for fragment in fragments[::-1]:
                if isinstance(fragment, ContentFragment) and ref == fragment.text:
                    return False

            return True

        for name in ("endnotes", "footnotes"):
            for _, ref in self._refs[name].items():
                if _can(ref):
                    fragments.append(ContentFragment(text=ref))

            _refs = {}
            for ref in list(self._refs[name].values()):
                matched = re.search("^(\[[0-9]+\])", ref)
                if matched:
                    _refs[matched.group(1)] = ref
            ref2fragment = {}
            for fragment in fragments:
                if isinstance(fragment, ContentFragment):
                    for ref_idx, ref in _refs.items():
                        if ref == fragment.text:
                            ref2fragment[ref_idx] = fragment
                            break
            for fragment in fragments:
                if isinstance(fragment, ContentFragment):
                    for idx, matched in enumerate(re.finditer("(\[[0-9]+\])", fragment.text)):
                        if matched.group(1) in ref2fragment:
                            tag = HtmlTag(start_offset=matched.start(),
                                          end_offset=matched.end(),
                                          type="a",
                                          content=ref2fragment[matched.group(1)].text)
                            fragment.tags.append(tag)
        # for name in ("endnotes", "footnotes"):

    # def node_generator(self, dom, tag_names):
    #     if not dom or not hasattr(dom, "childNodes"):
    #         raise StopIteration
    #     for child in dom.childNodes:
    #         if child and hasattr(child, "tagName") and child.tagName in tag_names:
    #             yield child
    #         yield from self.node_generator(child, tag_names)

    def node_generator(self, dom, tag_names, level):
        def in_fullback_dom(_dom):
            """去掉mc:Fallback"""
            while hasattr(_dom, "tagName"):
                if _dom.tagName == 'mc:Fallback':
                    return True
                if hasattr(_dom, "parentNode"):
                    _dom = _dom.parentNode
                else:
                    break
            return False

        def get_run_parent(n):
            if not n or not hasattr(n, "tagName"):
                return None
            if n.tagName == 'w:r':
                return n
            return get_run_parent(n.parentNode)

        if not in_fullback_dom(dom):

            if hasattr(dom, "tagName"):

                if dom.tagName in tag_names:
                    if dom.tagName == 'w:t':
                        """返回其父结构（w:r）"""
                        parent_node = get_run_parent(dom)
                        if parent_node:
                            yield parent_node, level
                        else:
                            yield dom, level
                    elif dom.tagName == 'm:oMath':
                        if dom.parentNode.tagName != 'm:oMathPara':
                            yield dom, level
                    elif dom.tagName == 'w:endnoteReference' or dom.tagName == 'w:footnoteReference':
                        parent_node = get_run_parent(dom)
                        if parent_node:
                            yield parent_node, level
                        else:
                            yield dom, level
                    else:
                        yield dom, level

                for child in dom.childNodes:
                    yield from self.node_generator(child, tag_names, level + 1)

    def generate_picture_fragments_from_dom(self, dom):
        ret = []
        for a_blip in dom.getElementsByTagName("a:blip"):
            rect = a_blip.parentNode.getElementsByTagName("a:srcRect")
            if rect:
                l = rect[0].getAttribute("l")
                r = rect[0].getAttribute("r")
                t = rect[0].getAttribute("t")
                b = rect[0].getAttribute("b")
            else:
                l, r, t, b = None, None, None, None
            rel_id = a_blip.getAttribute("r:embed")
            if rel_id in self._rels and 'image' in self._rels[rel_id]._target._content_type:
                suffix = self._rels[rel_id]._target._content_type.split("/")[-1]
                if suffix in ("x-emf", "x-wmf"):
                    suffix = "emf" if suffix == 'x-emf' else "wmf"
                    ret.append(PictureFragment(blob=self._rels[rel_id]._target.blob,
                                               caption="",
                                               suffix=f".{suffix}"))
                else:
                    blob = crop_image(self._rels[rel_id]._target.blob, l, r, t, b)
                    ret.append(PictureFragment(blob=blob, caption="", suffix=f".{suffix}"))
        if not ret:
            ret.append(PictureFragment(blob=None,
                                       caption="",
                                       suffix=".png"))
        return ret

    def generate_formula_fragments_from_dom(self, dom):
        def construct_xml(xml: str):
            nms = "<m:oMathPara xmlns:a=\"http://schemas.openxmlformats.org/drawingml/2006/main\" xmlns:m=\"http://schemas.openxmlformats.org/officeDocument/2006/math\" xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\""
            xml = re.sub("<m:oMathPara([ >])", nms+"\g<1>", xml)
            # xml = xml.replace("<m:oMathPara",
            xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n""" + xml
            return xml

        def get_o_math_xml(math_para):
            try:
                xml = math_para.toxml()
                xml = construct_xml(xml)
                xml = ElementTree.fromstring(xml.encode("utf-8"))
            except Exception as e:
                logger.warning(f"generate_formula_fragments_from_dom: {e}")
                xml = None
            return xml

        def get_inline_formula(dom):
            for o_math in load_string(dom.toxml()):
                return o_math.latex
            return ""

        fragments = []
        if not hasattr(dom, "tagName"):
            dom = dom.getElementsByTagName("m:oMathPara")
            if dom:
                dom = dom[0]
            else:
                return []

        if dom.tagName == 'm:oMathPara':
            xml = get_o_math_xml(dom)
            if xml is None:
                return []
            xml_str = ElementTree.tostring(xml, encoding='unicode')
            for o_math in load_string(xml_str):
                fragments.append(FormulaFragment(latex=o_math.latex,
                                                 tag=FormulaFragment.convert_latex_to_mathml(
                                                     "${}$".format(o_math.latex))))
        else:
            xml = add_namespaces_onto(dom.toxml(), self._namespaces, "<m:oMathPara ")
            run_dom = parseString(etree.tounicode(etree.fromstring(xml, parser=oxml_parser)))
            try:
                formula_text = get_inline_formula(run_dom)
                fragments.append(FormulaFragment(latex=formula_text,
                                                 inline=True,
                                                 tag=FormulaFragment.convert_latex_to_mathml(
                                                     "${}$".format(formula_text))))
            except Exception as e:
                logger.error("get_text_from_block: {}".format(e))
        return fragments

    def generate_shape_fragments_from_dom(self, dom):
        fragments = []
        for imagedata in dom.getElementsByTagName("v:imagedata"):
            rel_id = imagedata.getAttribute("r:id")
            l, r, t, b = get_crop_shape(imagedata)
            if rel_id in self._rels and "image" in self._rels[rel_id]._target._content_type:
                blob = self._rels[rel_id]._target.blob
                suffix = self._rels[rel_id]._target._content_type.split("/")[-1]
                if 'emf' in suffix or 'wmf' in suffix:
                    suffix = 'emf' if 'emf' in suffix else 'wmf'
                    if blob is not None:
                        fragments.append(ShapeFragment(blob=blob, suffix="." + suffix, caption=""))
                    else:
                        fragments.append(ShapeFragment(blob=self._rels[rel_id]._target.blob,
                                                       suffix="." + suffix,
                                                       caption=""))
                elif suffix.lower() in ("jpg", "jpeg", "png"):
                    """jpg png"""
                    fragments.append(ShapeFragment(blob=crop_image(blob, l, r, t, b),
                                                   suffix="." + suffix.lower(),
                                                   caption=""))
                else:
                    logger.warning("Unrecognized suffix {}".format(suffix))
                    fragments.append(ShapeFragment(blob=None,
                                                   suffix="." + suffix.lower(),
                                                   caption=""))
        if not fragments:
            fragments.append(ShapeFragment(blob=None,
                                           suffix=".png",
                                           caption=""))
        return fragments

    def generate_chart_fragments_from_dom(self, dom):
        fragments = []
        rel_id = dom.getAttribute("r:id")
        xml = self._rels[rel_id]._target.blob.decode("utf-8")
        fragments.append(ChartFragment(xml=xml))
        return fragments

    def generate_content_fragments_from_dom(self, dom):
        if dom and hasattr(dom, "parentNode") \
                and hasattr(dom.parentNode, "tagName") \
                and dom.parentNode.tagName == 'w:hyperlink' \
                and dom.parentNode.getAttribute("w:anchor"):
            """去除图表目录"""
            return []
        text = ""
        note_type = ""
        for ref in dom.getElementsByTagName("w:endnoteReference"):
            try:
                matched = re.search('^\[([0-9]+)\]', self._refs['endnotes'][ref.getAttribute('w:id')])
            except Exception:
                logger.warning(
                    "generate_content_fragments_from_dom: 未找到endnotes: {}".format(ref.getAttribute("w:id")))
                matched = None
            if matched:
                note_type = 'endnotes'
                text = text + "[{}]".format(matched.group(1))
        for ref in dom.getElementsByTagName("w:footnoteReference"):
            try:
                matched = re.search('^\[([0-9]+)\]', self._refs['footnotes'][ref.getAttribute('w:id')])
            except Exception:
                logger.warning(
                    "generate_content_fragments_from_dom: 未找到footnotes: {}".format(ref.getAttribute("w:id")))
                matched = None
            if matched:
                note_type = 'footnotes'
                text = text + "[{}]".format(matched.group(1))
        for fld_char in dom.getElementsByTagName("w:fldChar"):
            if fld_char.getAttribute("w:fldCharType") == "separate":
                text += "-"
        text = text + get_text_from_run_dom(dom)
        text = text.replace("[[", "[").replace("]]", "]")
        if note_type == "":
            note_type = []
        else:
            note_type = [note_type]
        return [ContentFragment(text=text, note_type=note_type)]

    def generate_content_fragment_from_w_t_dom(self, dom):
        text = ""
        if hasattr(dom, "childNodes"):
            for child_node in dom.childNodes:
                if hasattr(child_node, "data"):
                    text += child_node.data
        return ContentFragment(text=text)

    def generate_content_fragment_from_w_ref_dom(self, dom, name):
        text = ""
        try:
            matched = re.search('^\[([0-9]+)\]', self._refs[name][dom.getAttribute('w:id')])
        except Exception:
            logger.warning("get_text_from_block: 未找到{}: {}".format(name, dom.getAttribute("w:id")))
            matched = None
        if matched:
            text = matched.group(1)
        return [ContentFragment(text=text, note_type=name)]

    def _add_auto_numbering_into_fragments(self, fragments):
        for fragment in fragments:
            pass

    def generate_fragments_from_node(self, dom):
        fragments = []
        if dom.tagName == 'w:drawing':
            fragments += self.generate_picture_fragments_from_dom(dom)
        elif dom.tagName == 'm:oMathPara':
            fragments += self.generate_formula_fragments_from_dom(dom)
        elif dom.tagName == 'v:shape':
            fragments += self.generate_shape_fragments_from_dom(dom)
        elif dom.tagName == 'c:chart':
            fragments += self.generate_chart_fragments_from_dom(dom)
        elif dom.tagName == 'm:oMath':
            fragments += self.generate_formula_fragments_from_dom(dom)
        elif dom.tagName == 'w:endnoteReference':
            fragments += self.generate_content_fragment_from_w_ref_dom(dom, name='endnotes')
        elif dom.tagName == 'w:footnoteReference':
            fragments += self.generate_content_fragment_from_w_ref_dom(dom, name='footnotes')
        elif dom.tagName == 'w:r':
            fragments += self.generate_content_fragments_from_dom(dom)
        elif dom.tagName == 'w:t':
            fragments += [self.generate_content_fragment_from_w_t_dom(dom)]
        elif dom.tagName == 'w:noBreakHyphen':
            fragments += [ContentFragment(text="-")]
        else:
            raise ValueError("Unknown dom: {}".format(dom.tagName))
        return fragments

    def merge_fragments(self, fragments, levels):
        if not fragments:
            return []
        ret = []
        """进行block级别的属性获取"""
        fields = self.get_text_fields_from_block(self._block)
        # try:
        #     _number = self.get_auto_numbering_from_dom(self._block_dom, fields['outline'])
        # except Exception as e:
        #     logger.error("merge_fragments: {}".format(e))
        #     _number = ""
        # fmt = self.get_auto_numbering_fmt_from_dom(self._block_dom)

        content = ContentFragment(text="", **fields)
        curr_level = None
        for idx, fragment in enumerate(fragments):
            if isinstance(fragment, ContentFragment):
                if curr_level is not None and curr_level != levels[idx]:
                    content.text = content.text.replace("[[", "[").replace("]]", "]")
                    ret.append(content)
                    content = ContentFragment(text=fragment.text, **fields)
                    curr_level = levels[idx]
                else:
                    content.text += fragment.text
                    if fragment.note_type:
                        content.note_type.extend(fragment.note_type)
            elif isinstance(fragment, FormulaFragment) and fragment.inline:
                content.tags.append(HtmlTag(start_offset=len(content.text),
                                            end_offset=len(content.text),
                                            content=fragment.tag,
                                            type="math"))
            elif isinstance(fragment, PictureFragment) or isinstance(fragment, ShapeFragment) \
                    or isinstance(fragment, ChartFragment) or (
                    isinstance(fragment, FormulaFragment) and not fragment.inline):
                if content.text.strip():
                    content.text = content.text.replace("[[", "[").replace("]]", "]")
                    ret.append(content)
                    content = ContentFragment(text="", **fields)
                ret.append(fragment)
            else:
                raise ValueError(fragment.__class__.__name__)
        if content.text.strip():
            content.text = content.text.replace("[[", "[").replace("]]", "]")
            ret.append(content)

        try:
            fmt = self.get_auto_numbering_fmt_from_dom(self._block_dom)
        except Exception as e:
            fmt = None
        if fmt is not None and re.search("^(表|图|Table|Figure)", fmt['lvlText'], re.IGNORECASE) is not None:
            _key = 'refs-{}-{}'.format(fmt.get("num_id", 0), fmt.get('i_lvl', 0))
            _num = int(fmt['start']) + len(self._numbering_stack[_key])
            prefix = re.sub(f'%[0-9a-zA-Z一二三四五六七八九十ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]+?',
                            self._support_fmt.get(fmt['numFmt'], self._default_fmt)(_num),
                            fmt['lvlText'])
            self._numbering_stack[_key].append((0, prefix, None, None, None))
            if ret:
                for _ret in ret:
                    if isinstance(_ret, ContentFragment):
                        _ret.text = prefix + _ret.text
                        break
        elif fmt is not None and fields['outline'] != 0:
            try:
                prefix = self.get_auto_numbering_prefix(fmt, int(fields['outline']))
            except Exception as e:
                logger.info("merge_fragments: {}".format(e))
                prefix = ""
            if ret:
                """获取第一个ContentFragment"""
                for _ret in ret:
                    if isinstance(_ret, ContentFragment):
                        _ret.text = prefix + _ret.text
                        break
        elif fields['outline'] != 0:
            self._numbering_stack['numid'].append((int(fields['outline']), None, None, None))
        elif fmt is not None and fields['outline'] == 0 and (
                re.search('\[[^\]]*?\]', fmt['lvlText']) is not None
                or re.search("表|图", fmt['lvlText']) is not None):
            """作为参考文献或者作为图表标题"""
            _key = 'refs-{}-{}'.format(fmt.get("num_id", 0), fmt.get('i_lvl', 0))
            _num = int(fmt['start']) + len(self._numbering_stack[_key])
            prefix = re.sub(f'%[0-9a-zA-Z一二三四五六七八九十ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]+?',
                            self._support_fmt.get(fmt['numFmt'], self._default_fmt)(_num),
                            fmt['lvlText'])
            self._numbering_stack[_key].append((0, prefix, None, None, None))
            if ret:
                for _ret in ret:
                    if isinstance(_ret, ContentFragment):
                        _ret.text = prefix + _ret.text
                        break

        return ret

    def get_auto_numbering_prefix(self, fmt, outline: int) -> str:
        pattern = '[0-9a-zA-Z一二三四五六七八九十ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫ]'
        if self._numbering_stack['numid']:
            if self._numbering_stack["numid"][-1][0] < outline:
                _string, _format = self._numbering_stack["numid"][-1][1], self._numbering_stack["numid"][-1][2]
                if _string is not None and _format is not None:
                    ret = re.sub("%[0-9]", self._support_fmt.get(fmt['numFmt'], self._default_fmt)(int(fmt['start'])),
                                 re.sub(_format, _string, fmt['lvlText']))
                    self._numbering_stack["numid"].append((outline, ret, fmt['lvlText'], int(fmt['start'])))
                else:
                    ret = re.sub('%[0-9]+?', self._support_fmt.get(fmt['numFmt'], self._default_fmt)(int(fmt['start'])),
                                 fmt['lvlText'])
                    self._numbering_stack["numid"].append((outline, ret, fmt['lvlText'], int(fmt['start'])))
            else:
                while self._numbering_stack["numid"] and self._numbering_stack["numid"][-1][0] > outline:
                    self._numbering_stack["numid"].pop(-1)
                if self._numbering_stack["numid"]:
                    _string, _format, _num = self._numbering_stack["numid"][-1][1:]
                    if not (_string is None or _format is None or _num is None):
                        __format = clean_wildcard(_format, sp='()[]')
                        eng = re.sub(f'%{pattern}+?', f"({pattern}+?)", __format)
                        rst = {_k: _v for _k, _v in zip(re.findall(f'%{pattern}+?', _format),
                                                        re.findall(eng, _string)[0])}
                        if 'startOverride' in fmt:
                            _num = int(fmt['startOverride'])
                            del fmt['startOverride']
                        else:
                            _num += 1
                        if rst:
                            lvl_text = fmt['lvlText']
                            for _f, _i in rst.items():
                                if _f in re.findall(f'%{pattern}+?', fmt['lvlText'])[:-1]:
                                    lvl_text = re.sub(_f, _i, lvl_text)
                            ret = re.sub(f'%{pattern}+?', '{}', lvl_text).format(self._support_fmt.get(fmt['numFmt'],
                                                                                                       self._default_fmt)(
                                _num))

                        else:
                            raise ValueError("Error")
                    else:
                        """同级，但是前面是None，说明前面那一个并没有自动编号，应该从1开始"""
                        ret = re.sub(f'%{pattern}+?',
                                     self._support_fmt.get(fmt['numFmt'], self._default_fmt)(int(fmt['start'])),
                                     fmt['lvlText'])
                        _num = int(fmt['start'])
                        self._numbering_stack["numid"].append((outline, ret, fmt['lvlText'], _num))
                else:
                    ret = re.sub(f'%{pattern}+?',
                                 self._support_fmt.get(fmt['numFmt'], self._default_fmt)(int(fmt['start'])),
                                 fmt['lvlText'])
                    _num = int(fmt['start'])
                self._numbering_stack["numid"].append((outline, ret, fmt['lvlText'], _num))
        else:
            """第一个"""
            ret = re.sub('%[0-9]+?', self._support_fmt.get(fmt['numFmt'], self._default_fmt)(int(fmt['start'])),
                         fmt['lvlText'])
            _num = int(fmt['start'])
            self._numbering_stack["numid"].append((outline, ret, fmt['lvlText'], _num))
        return ret + " "

    def _convert_to_fragments(self):
        """"""
        fragments = []

        _blocks = list(self.block_generator(self._doc))
        for block in tqdm.tqdm(_blocks, desc="文档抽取"):
            try:
                if isinstance(block, Paragraph):
                    """按照run结构进行"""
                    self._block = block
                    self._block_dom = self.get_dom_from_xml(block._element.xml).firstChild
                    self._block_style_dom = self.get_dom_from_xml(block.style.element.xml)
                    self._namespaces = [f"xmlns:{k}='{v}'" for k, v in block._element.nsmap.items()]
                    _levels = []
                    _fragments = []
                    for node, _level in self.node_generator(self._block_dom, self._tag_names, 0):
                        __fragments = self.generate_fragments_from_node(node)
                        _fragments.extend(__fragments)
                        _levels.extend([_level] * len(__fragments))
                    """对_fragments进行合并"""
                    fragments += self.merge_fragments(_fragments, _levels)
                elif self.table and isinstance(block, Table):
                    fragments += self.generate_table_fragments_from_block(block)
                else:
                    logger.warning(f"Unknown block type: {type(block)}")
            except Exception as e:
                logger.error("_convert_to_fragments error: {}".format(e))

        """图表Caption获取"""
        fragments = self._match_fragment_captions(fragments)
        """按照字体识别目录"""
        self._recognize_outlines_by_chars(fragments)
        """加入自动编号"""
        self._add_auto_numbering_into_fragments(fragments)
        """添加refs碎片"""
        self.add_refs_into_fragments(fragments)
        return fragments
