"""Generate API reference DOCX for document_fragment platform."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

DOCS_DIR = Path(__file__).resolve().parent.parent / "docs"
OUTPUT = DOCS_DIR / "API-Reference.docx"
OUTPUT_MD = DOCS_DIR / "API-Reference.md"

SWAGGER = [
    ("Document API", "12355", "http://localhost:12355/docs", "http://localhost:12355/openapi.json"),
    ("Embedding API", "12356", "http://localhost:12356/docs", "http://localhost:12356/openapi.json"),
    ("RAG / QA API", "12357", "http://localhost:12357/docs", "http://localhost:12357/openapi.json"),
]

APIS = [
    {
        "service": "Document API (port 12355)",
        "swagger": "http://localhost:12355/docs",
        "note": "FastAPI root_path=/api/v1 (for proxy/OpenAPI). Routes are mounted at paths below.",
        "groups": [
            {
                "name": "Word — /word",
                "rows": [
                    ("POST", "/word/sync", "multipart: file; body: package_id, user_id", "Upload Word/doc, convert & index in background"),
                    ("POST", "/word/doc2docx", "multipart: file", "Convert .doc to .docx; returns file download"),
                ],
            },
            {
                "name": "PDF — /pdf",
                "rows": [
                    ("POST", "/pdf/sync", "multipart: file; Form: max_threads (opt); body: package_id, user_id", "Upload PDF, parse & index in background"),
                ],
            },
            {
                "name": "OCR — /ocr",
                "rows": [
                    ("POST", "/ocr/sync", "multipart: file", "Upload for OCR; returns file_id (MD5)"),
                ],
            },
            {
                "name": "CAJ — /caj",
                "rows": [
                    ("POST", "/caj/sync/caj", "multipart: file; body: package_id, user_id", "Upload CAJ, process in background"),
                ],
            },
            {
                "name": "ZIP — /zip",
                "rows": [
                    ("POST", "/zip/sync", "multipart: file; body: package_id, user_id", "Upload ZIP, extract & process"),
                    ("POST", "/zip/ftp", "query: ftp_url", "Fetch ZIP from FTP (background)"),
                    ("POST", "/zip/upload_ftp", "query: ftp_url, package_id, user_id", "FTP ZIP to knowledge base"),
                ],
            },
            {
                "name": "Document manage — /manage",
                "rows": [
                    ("GET", "/manage/list", "query: page_no (1), page_size (10), sort_field, sort_type (desc)", "List document libraries"),
                    ("POST", "/manage/update", "body: task_id, kwargs (dict)", "Update library metadata"),
                    ("DELETE", "/manage/document_delete", "body: md5 (optional)", "Delete library by MD5"),
                ],
            },
        ],
    },
    {
        "service": "Embedding API (port 12356)",
        "swagger": "http://localhost:12356/docs",
        "note": "Vector embedding and reranking service.",
        "groups": [
            {
                "name": "Embeddings",
                "rows": [
                    ("POST", "/embeddings", 'body: { "input": string | string[] | int[][] }', "Text to embedding vectors (OpenAI-compatible)"),
                    ("POST", "/rerank", "body: query (str), texts (List[str])", "Rerank texts; returns scores & softmax_scores"),
                ],
            },
        ],
    },
    {
        "service": "RAG / QA API (port 12357)",
        "swagger": "http://localhost:12357/docs",
        "note": "Knowledge base, agents, Q&A, and API config management.",
        "groups": [
            {
                "name": "Agent — /agent",
                "rows": [
                    ("GET", "/agent/list", "—", "List all agents"),
                    ("GET", "/agent/get", "query: agent_id", "Get agent details"),
                    ("POST", "/agent/create", "body: name, agent_prompt, description, icon; opt: agent_type, agent_example, agent_temperature", "Create agent"),
                    ("DELETE", "/agent/delete", "body: agent_id", "Delete agent"),
                    ("POST", "/agent/update", "body: agent_id, agent_attr, agent_value", "Update agent attribute"),
                    ("POST", "/agent/stream", "body: agent_id, query; thing_pattern (default false)", "Agent chat (SSE stream)"),
                ],
            },
            {
                "name": "Q&A — /qa",
                "rows": [
                    ("POST", "/qa/qa", "body: task_id, query, history, thing_pattern, user_id", "General LLM chat (SSE)"),
                    ("POST", "/qa/ocr-chat", "body: task_id, query, history, thing_pattern, user_id, file", "OCR document chat (SSE)"),
                    ("POST", "/qa/ocr-org", "same as ocr-chat", "OCR team template chat (SSE)"),
                    ("POST", "/qa/agent/db_agent", "body: task_id, query, history, thing_pattern, user_id", "Database assistant (SSE)"),
                    ("POST", "/qa/agent/knowledge_agent", "body: package_id, task_id, query, history, thing_pattern, user_id", "Knowledge-base agent (SSE)"),
                    ("POST", "/qa/agent/report_agent", "body: task_id, query, thing_pattern, user_id", "Annual report agent (SSE)"),
                    ("POST", "/qa/agent/week_agent", "body: task_id, query, thing_pattern, user_id", "Weekly report agent (SSE)"),
                    ("POST", "/qa/agent/template_agent", "body: task_id, query, thing_pattern, user_id, config_id", "Template agent (SSE)"),
                    ("POST", "/qa/agent/jira_week_agent", "body: task_id, query, thing_pattern, user_id", "Jira weekly report agent (SSE)"),
                    ("POST", "/qa/qa_desc", "body: sql_query (dict), page (1), page_size (15)", "SQL query table data"),
                    ("POST", "/qa/get_status", "body: task_id", "Get QA task status"),
                    ("POST", "/qa/report", "body: query", "Generate report file (binary download)"),
                ],
            },
            {
                "name": "Knowledge base — /knowledge_manage",
                "rows": [
                    ("GET", "/knowledge_manage/tree", "query: user_id, group_id (opt)", "Knowledge base tree"),
                    ("GET", "/knowledge_manage/package/list", "query: user_id, group_id (opt)", "List packages"),
                    ("GET", "/knowledge_manage/package/get", "query: knowledge_id", "Package details"),
                    ("GET", "/knowledge_manage/package/recommend", "query: knowledge_id", "Recommended questions"),
                    ("POST", "/knowledge_manage/package/create", "body: package_name, description, user_id, group_id (opt)", "Create package"),
                    ("POST", "/knowledge_manage/package/update", "body: package_id, package_name, description", "Update package"),
                    ("DELETE", "/knowledge_manage/package/delete", "body: package_id", "Delete package"),
                    ("GET", "/knowledge_manage/file/list", "query: package_id (opt)", "List files in package"),
                    ("POST", "/knowledge_manage/file/create", "body: id, file_id, file_name, file_size, file_path, file_type, package_id, user_id", "Register file (internal)"),
                    ("POST", "/knowledge_manage/file/upload", "body: package_id, ftp_url, user_id", "Upload via FTP URL"),
                    ("POST", "/knowledge_manage/file/update", "body: package_id, file_id, attr, value", "Update file attribute"),
                    ("DELETE", "/knowledge_manage/file/delete", "body: file_id", "Delete file"),
                    ("DELETE", "/knowledge_manage/delete", "—", "Clear all Milvus vectors"),
                    ("GET", "/knowledge_manage/get", "—", "List files in Milvus"),
                ],
            },
            {
                "name": "API config — /api_manage",
                "rows": [
                    ("POST", "/api_manage/api_insert", "body: data (dict)", "Add API config"),
                    ("GET", "/api_manage/api_search", "query: api_name", "Get API config details"),
                    ("GET", "/api_manage/api_list", "—", "List all API configs"),
                    ("POST", "/api_manage/api_update", "body: api_name, info (dict)", "Update API metadata"),
                    ("GET", "/api_manage/field_list", "query: api_name", "List updatable fields"),
                    ("POST", "/api_manage/api_field_insert", "body: api_name, type, field_data (dict)", "Add field to config"),
                    ("POST", "/api_manage/api_data_insert", "body: api_name, field_name, field_data (dict)", "Add field data"),
                    ("POST", "/api_manage/api_data_update", "body: api_name, field_name, field_data (dict)", "Update field data"),
                    ("GET", "/api_manage/api_data_list", "query: page_on, page_size, api_name, field_name", "Paginated field data"),
                    ("DELETE", "/api_manage/api_data_delete", "body: api_name, field_name, data_list", "Delete field data items"),
                ],
            },
            {
                "name": "User config — /user_config_manage",
                "rows": [
                    ("GET", "/user_config_manage/list", "query: user_id", "List configs for user"),
                    ("GET", "/user_config_manage/get", "query: user_config_id", "Get config details"),
                    ("POST", "/user_config_manage/create", "body: user_id, config_json, config_name", "Create JSON template"),
                    ("DELETE", "/user_config_manage/delete", "body: user_config_id", "Delete template"),
                    ("POST", "/user_config_manage/update", "body: user_config_id, user_id, config_json, config_name", "Update template"),
                ],
            },
        ],
    },
]


def set_cell_shading(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_api_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Method", "Path", "Parameters", "Description"]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "D9E2F3")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for method, path, params, desc in rows:
        row = table.add_row().cells
        row[0].text = method
        row[1].text = path
        row[2].text = params
        row[3].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Document Fragment Platform — API Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Enterprise document processing and RAG backend. One codebase, three FastAPI services, "
        "one Docker image. All services expose interactive Swagger UI at /docs."
    )

    doc.add_heading("Swagger / OpenAPI Addresses", level=1)
    doc.add_paragraph(
        "Start services with docker compose up, then open Swagger in a browser. "
        "Replace localhost with your host IP when accessing remotely."
    )

    sw_table = doc.add_table(rows=1, cols=4)
    sw_table.style = "Table Grid"
    sw_hdr = sw_table.rows[0].cells
    for i, h in enumerate(["Service", "Port", "Swagger UI", "OpenAPI JSON"]):
        sw_hdr[i].text = h
        set_cell_shading(sw_hdr[i], "4472C4")
        for p in sw_hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    for name, port, docs_url, openapi_url in SWAGGER:
        row = sw_table.add_row().cells
        row[0].text = name
        row[1].text = port
        row[2].text = docs_url
        row[3].text = openapi_url
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_heading("How to use Swagger", level=2)
    bullets = [
        "Run: docker compose up -d (from document_fragment project root)",
        "Open the Swagger URL for the service you need (see table above)",
        "Expand an endpoint, click Try it out, fill parameters, and Execute",
        "Streaming endpoints (/qa/*, /agent/stream) return text/event-stream — use curl or a client that supports SSE",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Response formats", level=1)
    fmt_table = doc.add_table(rows=1, cols=2)
    fmt_table.style = "Table Grid"
    fmt_hdr = fmt_table.rows[0].cells
    fmt_hdr[0].text = "Service / route type"
    fmt_hdr[1].text = "Response shape"
    set_cell_shading(fmt_hdr[0], "D9E2F3")
    set_cell_shading(fmt_hdr[1], "D9E2F3")
    formats = [
        ("Document API", '{ "data", "detail", "status_code" }'),
        ("RAG API (most routes)", '{ "data", "detail", "status_code" }'),
        ("RAG /api_manage/*", '{ "status", "detail", "data" }'),
        ("QA streaming routes", "text/event-stream (SSE), ends with data: [DONE]"),
        ("Embedding API", "OpenAI-style embedding JSON / { scores, softmax_scores }"),
    ]
    for a, b in formats:
        r = fmt_table.add_row().cells
        r[0].text = a
        r[1].text = b

    doc.add_paragraph()

    for svc in APIS:
        doc.add_page_break()
        h = doc.add_heading(svc["service"], level=1)
        p = doc.add_paragraph()
        run = p.add_run("Swagger: ")
        run.bold = True
        p.add_run(svc["swagger"])
        if svc.get("note"):
            doc.add_paragraph(svc["note"])

        for group in svc["groups"]:
            doc.add_heading(group["name"], level=2)
            add_api_table(doc, group["rows"])

    doc.add_page_break()
    doc.add_heading("Appendix — Services not in docker-compose", level=1)
    doc.add_paragraph(
        "These modules exist in source but are not started by default in docker-compose.yaml:"
    )
    add_api_table(
        doc,
        [
            ("POST", "/convert (doc2docx)", "multipart: file", "Standalone .doc to .docx conversion"),
            ("POST", "/table, /structure/{lang}, /ocr/{lang}", "layout model internal", "Document layout/OCR model server"),
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Generated for document_fragment project. ").italic = True
    footer.add_run("See docs/README.md for full documentation index.").italic = True

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    write_markdown()
    print(f"Written: {OUTPUT}")
    print(f"Written: {OUTPUT_MD}")
    print("Open DOCX in Word (Cursor cannot preview .docx files).")


def write_markdown():
    lines = [
        "# Document Fragment Platform — API Reference",
        "",
        "Enterprise document processing and RAG backend. All services expose Swagger at `/docs`.",
        "",
        "> **Note:** Cursor cannot open `.docx` in the editor. Use this `.md` file in Cursor, or open",
        "> `API-Reference.docx` with Microsoft Word from File Explorer.",
        "",
        "## Swagger / OpenAPI",
        "",
        "| Service | Port | Swagger UI | OpenAPI JSON |",
        "|---------|------|------------|--------------|",
    ]
    for name, port, docs_url, openapi_url in SWAGGER:
        lines.append(f"| {name} | {port} | {docs_url} | {openapi_url} |")
    lines.extend([
        "",
        "## How to use Swagger",
        "",
        "1. Run `docker compose up -d` from the `document_fragment` project root",
        "2. Open a Swagger URL above in your browser",
        "3. Expand an endpoint → **Try it out** → **Execute**",
        "",
        "## Response formats",
        "",
        "| Service | Response shape |",
        "|---------|----------------|",
        '| Document API | `{ "data", "detail", "status_code" }` |',
        '| RAG (most routes) | `{ "data", "detail", "status_code" }` |',
        '| RAG `/api_manage/*` | `{ "status", "detail", "data" }` |',
        "| QA streaming | `text/event-stream` (SSE) |",
        "| Embedding API | OpenAI-style JSON / `{ scores, softmax_scores }` |",
        "",
    ])
    for svc in APIS:
        lines.extend(["---", "", f"## {svc['service']}", "", f"**Swagger:** {svc['swagger']}", ""])
        if svc.get("note"):
            lines.extend([svc["note"], ""])
        for group in svc["groups"]:
            lines.extend([f"### {group['name']}", ""])
            lines.append("| Method | Path | Parameters | Description |")
            lines.append("|--------|------|------------|-------------|")
            for method, path, params, desc in group["rows"]:
                lines.append(f"| {method} | `{path}` | {params} | {desc} |")
            lines.append("")
    lines.extend([
        "---",
        "",
        "## Appendix — not in docker-compose",
        "",
        "| Method | Path | Parameters | Description |",
        "|--------|------|------------|-------------|",
        "| POST | `/convert` | multipart: file | Standalone doc→docx |",
        "| POST | `/table`, `/structure/{lang}`, `/ocr/{lang}` | internal | Layout/OCR model |",
        "",
    ])
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
